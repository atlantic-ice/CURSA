#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
VKR MAIN to DOCX Converter
Converts VKR_MAIN.md to properly formatted DOCX file according to GOST 7.32-2017
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


class VKRConverter:
    def __init__(self):
        self.doc = Document()
        self.setup_document()
        self.heading_levels = {}
        
    def setup_document(self):
        """Setup document formatting according to GOST 7.32-2017"""
        # Page margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1.5)
        
        # Default style
        styles = self.doc.styles
        style = styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Cm(1.25)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)

    def add_paragraph_formatted(self, text, level=0, bold=False, centered=False, 
                               first_line_indent=True, space_before=0, space_after=0):
        """Add paragraph with specific formatting"""
        p = self.doc.add_paragraph()
        
        if level > 0:
            p.style = f'Heading {level}'
        
        p.paragraph_format.line_spacing = 1.5
        if first_line_indent and level == 0:
            p.paragraph_format.first_line_indent = Cm(1.25)
        else:
            p.paragraph_format.first_line_indent = Cm(0)
        
        if space_before > 0:
            p.paragraph_format.space_before = Pt(space_before)
        if space_after > 0:
            p.paragraph_format.space_after = Pt(space_after)
        
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        if bold:
            run.bold = True
        
        if centered:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return p

    def process_heading(self, text):
        """Process heading markers and return level and text"""
        level = 0
        while text.startswith('#'):
            level += 1
            text = text[1:].lstrip()
        return level, text.strip()

    def process_bold_text(self, text):
        """Process bold text (**text**)"""
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25)
        
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            elif part:
                run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        return p

    def process_table(self, lines, start_idx):
        """Process markdown table"""
        table_lines = [lines[start_idx]]
        i = start_idx + 1
        
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        
        rows = []
        for tbl_line in table_lines:
            if '---' in tbl_line or ':--:' in tbl_line:
                continue
            cells = [c.strip() for c in tbl_line.strip('|').split('|')]
            rows.append(cells)
        
        if len(rows) > 0:
            cols = len(rows[0])
            tbl = self.doc.add_table(rows=len(rows), cols=cols)
            tbl.style = 'Table Grid'
            
            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_text in enumerate(row_data):
                    cell = tbl.rows[row_idx].cells[col_idx]
                    
                    # Clear default paragraph
                    cell.text = ''
                    
                    # Add formatted text
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    
                    if row_idx == 0:
                        run.bold = True
                    
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            
            return i - 1
        
        return start_idx

    def convert(self, md_file: str, output_file: str):
        """Convert markdown file to DOCX"""
        # Read markdown file
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = content.split('\n')
        i = 0
        toc_entries = []
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Skip empty lines
            if not stripped:
                i += 1
                continue
            
            # Process heading
            if line.startswith('#'):
                level, heading_text = self.process_heading(line)
                
                if level == 1:
                    self.add_paragraph_formatted(heading_text, level=1, bold=True, 
                                               centered=True, first_line_indent=False,
                                               space_before=12, space_after=6)
                    toc_entries.append((level, heading_text))
                
                elif level == 2:
                    self.add_paragraph_formatted(heading_text, level=2, bold=True, 
                                               first_line_indent=False,
                                               space_before=12, space_after=6)
                    toc_entries.append((level, heading_text))
                
                elif level >= 3:
                    self.add_paragraph_formatted(heading_text, level=min(3, level), 
                                               bold=True, first_line_indent=False,
                                               space_before=8, space_after=4)
                    toc_entries.append((level, heading_text))
            
            # Process horizontal rule
            elif stripped.startswith('---'):
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            
            # Process table
            elif stripped.startswith('|'):
                i = self.process_table(lines, i)
            
            # Process bold text
            elif '**' in stripped:
                self.process_bold_text(stripped)
            
            # Process regular text
            else:
                text = stripped.replace('`', '')
                if text and text != '---':
                    self.add_paragraph_formatted(text)
            
            i += 1
        
        # Save document
        self.doc.save(output_file)
        return True

    def add_table_of_contents(self):
        """Add table of contents"""
        pass  # python-docx has limited TOC support


def main():
    """Main function"""
    script_dir = Path(__file__).parent
    input_file = script_dir / "VKR_MAIN.md"
    output_file = script_dir / "VKR_NIKITA_CURSA_2026.docx"
    
    if not input_file.exists():
        print(f"[ERROR] File not found: {input_file}")
        return False
    
    print(f"[INFO] Starting conversion of {input_file.name}")
    
    converter = VKRConverter()
    converter.convert(str(input_file), str(output_file))
    
    # Validate output
    if output_file.exists():
        file_size = output_file.stat().st_size
        print(f"[OK] Conversion completed successfully")
        print(f"[OK] Output file: {output_file.name}")
        print(f"[OK] File size: {file_size:,} bytes ({file_size / 1024:.2f} KB)")
        
        # Validate DOCX structure
        try:
            test_doc = Document(str(output_file))
            para_count = len(test_doc.paragraphs)
            table_count = len(test_doc.tables)
            print(f"[OK] Document validation: {para_count} paragraphs, {table_count} tables")
            print(f"[OK] Formatting applied: Times New Roman 14pt, 1.5 line spacing")
            print(f"[OK] Margins: Left 3cm, Right 1.5cm, Top 2cm, Bottom 2cm")
            return True
        except Exception as e:
            print(f"[ERROR] Validation failed: {e}")
            return False
    else:
        print(f"[ERROR] Output file was not created")
        return False


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
