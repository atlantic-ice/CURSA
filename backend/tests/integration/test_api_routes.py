"""Интеграционные тесты для API маршрутов."""
import os
import io
import pytest
import tempfile
from docx import Document


def test_api_document_upload_route_exists(client):
    """Тест существования маршрута для загрузки документа."""
    # Отправляем пустой запрос, ожидаем ошибку 400, но не 404
    response = client.post('/api/document/upload')
    assert response.status_code != 404, "Route /api/document/upload should exist"
    assert response.status_code == 400  # Нет файла


def test_document_upload_with_invalid_data(client):
    """Тест загрузки документа с неверными данными."""
    response = client.post('/api/document/upload', data={
        'wrong_key': 'wrong_value'
    })
    assert response.status_code == 400


def test_document_upload_with_valid_file(client):
    """Тест загрузки документа с корректным файлом."""
    # Создаём тестовый DOCX файл
    with tempfile.TemporaryDirectory() as tmpdir:
        test_path = os.path.join(tmpdir, 'test.docx')
        doc = Document()
        doc.add_paragraph("Тестовый документ")
        doc.save(test_path)

        with open(test_path, 'rb') as f:
            response = client.post(
                '/api/document/upload',
                data={'file': (io.BytesIO(f.read()), 'test.docx')},
                content_type='multipart/form-data'
            )

        assert response.status_code == 200
        assert 'check_results' in response.json
        assert 'temp_path' in response.json
        assert 'correction_success' in response.json
        assert 'quality_gate_passed' in response.json
        assert 'graduation_readiness' in response.json

        quality_metrics = response.json.get('quality_metrics')
        if quality_metrics:
            assert 'before_total_issues' in quality_metrics
            assert 'after_total_issues' in quality_metrics
            assert 'attempts' in quality_metrics
            if response.json.get('quality_gate_passed'):
                assert (
                    quality_metrics['after_total_issues']
                    <= quality_metrics['before_total_issues']
                )


def test_document_upload_quality_gate_never_degrades(client):
    """Upload-flow не должен ухудшать качество результата относительно исходного документа."""
    with tempfile.TemporaryDirectory() as tmpdir:
        test_path = os.path.join(tmpdir, 'complex_quality_gate.docx')
        doc = Document()

        # Формируем документ с заведомыми нарушениями форматирования.
        for i in range(40):
            p = doc.add_paragraph(
                f"Раздел {i + 1}. Текст с некорректным форматированием для проверки quality gate."
            )
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = None
        doc.save(test_path)

        with open(test_path, 'rb') as f:
            response = client.post(
                '/api/document/upload',
                data={'file': (io.BytesIO(f.read()), 'complex_quality_gate.docx')},
                content_type='multipart/form-data'
            )

        assert response.status_code == 200
        payload = response.json
        assert payload.get('correction_success') is True
        assert payload.get('quality_gate_passed') is True

        quality_metrics = payload.get('quality_metrics', {})
        assert quality_metrics.get('before_total_issues') is not None
        assert quality_metrics.get('after_total_issues') is not None
        assert quality_metrics.get('attempts')
        assert quality_metrics['after_total_issues'] <= quality_metrics['before_total_issues']
        assert 'fallback_applied' in quality_metrics

        readiness = payload.get('graduation_readiness', {})
        assert readiness.get('status') in {'ready', 'almost_ready', 'needs_revision', 'unknown'}


def test_document_upload_with_fake_docx_returns_400(client):
    """Файл с расширением .docx, но некорректным содержимым, должен отклоняться."""
    fake_docx = io.BytesIO(b"this is not a real docx archive")
    response = client.post(
        '/api/document/upload',
        data={'file': (fake_docx, 'fake.docx')},
        content_type='multipart/form-data',
    )

    assert response.status_code == 400
    assert 'error' in response.json


def test_document_analyze_returns_document_token(client):
    """Анализ документа возвращает токен сессии вместо temp_path для UI."""
    with tempfile.TemporaryDirectory() as tmpdir:
        test_path = os.path.join(tmpdir, 'analyze_test.docx')
        doc = Document()
        doc.add_paragraph("Тестовый документ для анализа")
        doc.save(test_path)

        with open(test_path, 'rb') as f:
            response = client.post(
                '/api/document/analyze',
                data={'file': (io.BytesIO(f.read()), 'analyze_test.docx')},
                content_type='multipart/form-data'
            )

        assert response.status_code == 200
        assert 'check_results' in response.json
        assert 'document_token' in response.json
        assert response.json['document_token']
        assert 'temp_path' not in response.json


def test_document_autocorrect_requires_document_token(client):
    """Новый маршрут автокоррекции требует document_token."""
    response = client.post('/api/document/autocorrect', json={})

    assert response.status_code == 400
    assert 'error' in response.json


def test_legacy_autocorrect_route_returns_not_found_without_session(client):
    """Legacy маршрут совместимости отвечает 404 при отсутствии активной сессии."""
    response = client.post('/api/document/missing-session/autocorrect', json={})

    assert response.status_code == 404
    assert 'error' in response.json


def test_document_autocorrect_returns_corrected_check_results(client):
    """Автокоррекция возвращает check_results уже исправленного документа."""
    with tempfile.TemporaryDirectory() as tmpdir:
        test_path = os.path.join(tmpdir, 'autocorrect_check_results.docx')
        doc = Document()
        paragraph = doc.add_paragraph("Тест для повторного анализа после коррекции")
        for run in paragraph.runs:
            run.font.name = 'Arial'
        doc.save(test_path)

        with open(test_path, 'rb') as f:
            analyze_response = client.post(
                '/api/document/analyze',
                data={'file': (io.BytesIO(f.read()), 'autocorrect_check_results.docx')},
                content_type='multipart/form-data'
            )

        assert analyze_response.status_code == 200
        document_token = analyze_response.json['document_token']
        initial_total_issues = analyze_response.json['check_results']['total_issues_count']

        autocorrect_response = client.post(
            '/api/document/autocorrect',
            json={
                'document_token': document_token,
                'original_filename': 'autocorrect_check_results.docx',
            }
        )

        assert autocorrect_response.status_code == 200
        assert 'check_results' in autocorrect_response.json
        assert 'original_preview_path' in autocorrect_response.json
        assert 'temp_path' not in autocorrect_response.json
        assert 'quality_gate_passed' in autocorrect_response.json
        assert autocorrect_response.json['quality_gate_passed'] is True
        assert 'quality_metrics' in autocorrect_response.json
        quality_metrics = autocorrect_response.json['quality_metrics']
        assert quality_metrics['after_total_issues'] <= quality_metrics['before_total_issues']
        assert 'fallback_applied' in quality_metrics
        assert 'graduation_readiness' in autocorrect_response.json
        assert (
            autocorrect_response.json['check_results']['total_issues_count']
            <= initial_total_issues
        )
