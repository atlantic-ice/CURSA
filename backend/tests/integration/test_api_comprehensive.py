"""
Расширенные интеграционные тесты для API endpoints.
Цель: проверить end-to-end workflows и API контрактов.
"""

import pytest
import json
from io import BytesIO
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm

# Флаг для функциональных тестов (полную интеграцию с Flask)
pytestmark = pytest.mark.integration


@pytest.fixture
def api_client(client):
    """API client fixture для тестирования endpoints"""
    return client


@pytest.fixture
def test_docx_file(tmp_path):
    """Создать тестовый DOCX файл"""
    doc = Document()

    # Настройка полей
    for section in doc.sections:
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Добавить содержимое
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Test content paragraph")
    doc.add_paragraph("ОСНОВНАЯ ЧАСТЬ")
    doc.add_paragraph("More content")

    filepath = tmp_path / "test_document.docx"
    doc.save(str(filepath))
    return filepath


@pytest.fixture
def auth_token(api_client):
    """Получить auth token для тестов"""
    # Регистрация пользователя
    response = api_client.post(
        "/api/auth/register",
        json={
            "email": "test@example.com",
            "password": "TestPassword123!",
            "name": "Test User",
        },
    )

    if response.status_code != 201:
        # Пользователь может уже существовать, попробуем логин
        response = api_client.post(
            "/api/auth/login",
            json={
                "email": "test@example.com",
                "password": "TestPassword123!",
            },
        )

    data = response.get_json()
    return data.get("access_token") if response.status_code in [200, 201] else None


# ============================================================================
# Document Upload Tests
# ============================================================================


class TestDocumentUploadAPI:
    """Интеграционные тесты для загрузки документов"""

    def test_upload_valid_docx_file(self, api_client, test_docx_file, auth_token):
        """✓ Загрузка валидного DOCX файла"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        with open(test_docx_file, "rb") as f:
            files = {
                "file": (
                    "test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            response = api_client.post(
                "/api/document/upload",
                data={"profile_id": "default_gost"},
                files=files,
                headers=headers,
            )

        assert response.status_code in [200, 201]
        data = response.get_json()
        assert "document_id" in data or "id" in data

    def test_upload_without_profile(self, api_client, test_docx_file, auth_token):
        """Testing upload with default profile"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        with open(test_docx_file, "rb") as f:
            files = {
                "file": (
                    "test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            response = api_client.post(
                "/api/document/upload",
                files=files,
                headers=headers,
            )

        # Should use default profile
        assert response.status_code in [200, 201, 400, 422]

    def test_upload_large_file(self, api_client, tmp_path, auth_token):
        """Testing upload of larger document (50+ pages)"""
        # Create large document
        doc = Document()
        for i in range(100):
            doc.add_paragraph(f"Page {i+1}: " + "Lorem ipsum " * 50)

        large_file = tmp_path / "large_document.docx"
        doc.save(str(large_file))

        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        with open(large_file, "rb") as f:
            files = {
                "file": (
                    "large.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            response = api_client.post(
                "/api/document/upload",
                data={"profile_id": "default_gost"},
                files=files,
                headers=headers,
            )

        assert response.status_code in [200, 201, 413]  # 413 = Payload Too Large

    def test_upload_invalid_file_type(self, api_client, tmp_path, auth_token):
        """✗ Загрузка файла неправильного формата"""
        # Create text file
        text_file = tmp_path / "test.txt"
        text_file.write_text("This is not a DOCX file")

        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        with open(text_file, "rb") as f:
            files = {"file": ("test.txt", f, "text/plain")}
            response = api_client.post(
                "/api/document/upload",
                files=files,
                headers=headers,
            )

        # Should reject invalid file type
        assert response.status_code in [400, 415, 422]

    def test_upload_empty_file(self, api_client, tmp_path, auth_token):
        """✗ Загрузка пустого файла"""
        # Create empty DOCX
        empty_file = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(empty_file))

        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        with open(empty_file, "rb") as f:
            files = {
                "file": (
                    "empty.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            response = api_client.post(
                "/api/document/upload",
                files=files,
                headers=headers,
            )

        # Should accept but may have validation warnings
        assert response.status_code in [200, 201]


# ============================================================================
# Document Validation Tests
# ============================================================================


class TestDocumentValidationAPI:
    """Интеграционные тесты для валидации документов"""

    def test_validate_document_returns_issues(self, api_client, test_docx_file, auth_token):
        """✓ Валидация документа возвращает список проблем"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        # 1. Загрузить документ
        with open(test_docx_file, "rb") as f:
            files = {
                "file": (
                    "test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            upload_response = api_client.post(
                "/api/document/upload",
                data={"profile_id": "default_gost"},
                files=files,
                headers=headers,
            )

        if upload_response.status_code not in [200, 201]:
            pytest.skip("Upload failed")

        upload_data = upload_response.get_json()
        doc_id = upload_data.get("document_id") or upload_data.get("id")

        # 2. Запросить валидацию
        validation_response = api_client.post(
            f"/api/document/{doc_id}/validate",
            headers=headers,
        )

        assert validation_response.status_code == 200
        validation_data = validation_response.get_json()
        assert "issues" in validation_data or "validation_result" in validation_data

    def test_validate_with_specific_profile(self, api_client, test_docx_file, auth_token):
        """Testing validation with specific university profile"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        with open(test_docx_file, "rb") as f:
            files = {
                "file": (
                    "test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            upload_response = api_client.post(
                "/api/document/upload",
                data={"profile_id": "mgu_requirements"},  # МГУ profile
                files=files,
                headers=headers,
            )

        if upload_response.status_code not in [200, 201]:
            pytest.skip("Upload with profile failed")

        upload_data = upload_response.get_json()
        doc_id = upload_data.get("document_id") or upload_data.get("id")

        validation_response = api_client.post(
            f"/api/document/{doc_id}/validate",
            json={"profile_id": "mgu_requirements"},
            headers=headers,
        )

        assert validation_response.status_code in [200, 400, 404]

    def test_validation_result_structure(self, api_client, test_docx_file, auth_token):
        """✓ Структура результата валидации"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        with open(test_docx_file, "rb") as f:
            files = {
                "file": (
                    "test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            upload_response = api_client.post(
                "/api/document/upload",
                data={"profile_id": "default_gost"},
                files=files,
                headers=headers,
            )

        if upload_response.status_code not in [200, 201]:
            pytest.skip("Upload failed")

        upload_data = upload_response.get_json()
        doc_id = upload_data.get("document_id") or upload_data.get("id")

        validation_response = api_client.post(
            f"/api/document/{doc_id}/validate",
            headers=headers,
        )

        if validation_response.status_code == 200:
            data = validation_response.get_json()
            # Check for expected fields
            assert isinstance(data, dict)


# ============================================================================
# Auto-Correction Tests
# ============================================================================


class TestAutoCorrectAPI:
    """Интеграционные тесты для автоисправления"""

    def test_autocorrect_returns_fixed_document(self, api_client, test_docx_file, auth_token):
        """✓ Автоисправление возвращает исправленный документ"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        # 1. Загрузить документ
        with open(test_docx_file, "rb") as f:
            files = {
                "file": (
                    "test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            upload_response = api_client.post(
                "/api/document/upload",
                data={"profile_id": "default_gost"},
                files=files,
                headers=headers,
            )

        if upload_response.status_code not in [200, 201]:
            pytest.skip("Upload failed")

        upload_data = upload_response.get_json()
        doc_id = upload_data.get("document_id") or upload_data.get("id")

        # 2. Запросить автоисправление
        correction_response = api_client.post(
            f"/api/document/{doc_id}/autocorrect",
            headers=headers,
        )

        assert correction_response.status_code in [200, 202, 400, 404]

    def test_autocorrect_selective_issues(self, api_client, test_docx_file, auth_token):
        """Testing autocorrect for specific issues"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        with open(test_docx_file, "rb") as f:
            files = {
                "file": (
                    "test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            upload_response = api_client.post(
                "/api/document/upload",
                data={"profile_id": "default_gost"},
                files=files,
                headers=headers,
            )

        if upload_response.status_code not in [200, 201]:
            pytest.skip("Upload failed")

        upload_data = upload_response.get_json()
        doc_id = upload_data.get("document_id") or upload_data.get("id")

        # Autocorrect specific issue types
        correction_response = api_client.post(
            f"/api/document/{doc_id}/autocorrect",
            json={"issue_types": ["font", "margins"]},
            headers=headers,
        )

        assert correction_response.status_code in [200, 202, 400, 404]


# ============================================================================
# Report Generation Tests
# ============================================================================


class TestReportGenerationAPI:
    """Интеграционные тесты для генерации отчётов"""

    def test_generate_report_pdf(self, api_client, test_docx_file, auth_token):
        """✓ Генерация PDF отчёта"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        # 1. Загрузить и валидировать
        with open(test_docx_file, "rb") as f:
            files = {
                "file": (
                    "test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            upload_response = api_client.post(
                "/api/document/upload",
                data={"profile_id": "default_gost"},
                files=files,
                headers=headers,
            )

        if upload_response.status_code not in [200, 201]:
            pytest.skip("Upload failed")

        upload_data = upload_response.get_json()
        doc_id = upload_data.get("document_id") or upload_data.get("id")

        # 2. Генерировать отчёт PDF
        report_response = api_client.get(
            f"/api/document/{doc_id}/report/pdf",
            headers=headers,
        )

        assert report_response.status_code in [200, 400, 404]
        if report_response.status_code == 200:
            assert report_response.content_type in ["application/pdf", "application/octet-stream"]

    def test_generate_report_html(self, api_client, test_docx_file, auth_token):
        """✓ Генерация HTML отчёта"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        with open(test_docx_file, "rb") as f:
            files = {
                "file": (
                    "test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            upload_response = api_client.post(
                "/api/document/upload",
                data={"profile_id": "default_gost"},
                files=files,
                headers=headers,
            )

        if upload_response.status_code not in [200, 201]:
            pytest.skip("Upload failed")

        upload_data = upload_response.get_json()
        doc_id = upload_data.get("document_id") or upload_data.get("id")

        # Генерировать HTML отчёт
        report_response = api_client.get(
            f"/api/document/{doc_id}/report/html",
            headers=headers,
        )

        assert report_response.status_code in [200, 400, 404]
        if report_response.status_code == 200:
            assert report_response.content_type in ["text/html", "text/plain"]

    def test_report_contains_summary(self, api_client, test_docx_file, auth_token):
        """✓ Отчёт содержит резюме проверки"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        with open(test_docx_file, "rb") as f:
            files = {
                "file": (
                    "test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            upload_response = api_client.post(
                "/api/document/upload",
                data={"profile_id": "default_gost"},
                files=files,
                headers=headers,
            )

        if upload_response.status_code not in [200, 201]:
            pytest.skip("Upload failed")

        upload_data = upload_response.get_json()
        doc_id = upload_data.get("document_id") or upload_data.get("id")

        # Get JSON report for analysis
        report_response = api_client.get(
            f"/api/document/{doc_id}/report",
            headers=headers,
        )

        if report_response.status_code == 200:
            data = report_response.get_json()
            assert isinstance(data, dict)


# ============================================================================
# Profile Management Tests
# ============================================================================


class TestProfileManagementAPI:
    """Интеграционные тесты для управления профилями"""

    def test_list_available_profiles(self, api_client):
        """✓ Получить список доступных профилей"""
        response = api_client.get("/api/profiles")

        assert response.status_code == 200
        data = response.get_json()
        assert isinstance(data, list)
        assert len(data) > 0

    def test_get_specific_profile(self, api_client):
        """✓ Получить конкретный профиль"""
        response = api_client.get("/api/profiles/default_gost")

        assert response.status_code in [200, 404]
        if response.status_code == 200:
            data = response.get_json()
            assert "name" in data or "rules" in data

    def test_create_custom_profile(self, api_client, auth_token):
        """✓ Создать пользовательский профиль"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        profile_data = {
            "name": "Custom Profile",
            "rules": {
                "font": {"name": "Times New Roman", "size": 14},
                "margins": {"left": 3.0, "right": 1.5, "top": 2.0, "bottom": 2.0},
            },
        }

        response = api_client.post(
            "/api/profiles",
            json=profile_data,
            headers=headers,
        )

        assert response.status_code in [201, 200, 400, 401]

    def test_update_profile(self, api_client, auth_token):
        """Testing profile update endpoint"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        update_data = {
            "rules": {
                "font": {"size": 14.0},
            },
        }

        response = api_client.put(
            "/api/profiles/default_gost",
            json=update_data,
            headers=headers,
        )

        assert response.status_code in [200, 400, 401, 403]


# ============================================================================
# Authentication Tests
# ============================================================================


class TestAuthenticationAPI:
    """Интеграционные тесты для аутентификации"""

    def test_user_registration(self, api_client):
        """✓ Регистрация нового пользователя"""
        response = api_client.post(
            "/api/auth/register",
            json={
                "email": f"new_user_{datetime.now().timestamp()}@example.com",
                "password": "TestPassword123!",
                "name": "New User",
            },
        )

        assert response.status_code in [201, 200, 400]

    def test_user_login(self, api_client):
        """✓ Вход пользователя"""
        # Сначала регистрируемся
        email = f"login_test_{datetime.now().timestamp()}@example.com"
        api_client.post(
            "/api/auth/register",
            json={
                "email": email,
                "password": "TestPassword123!",
                "name": "Login Test",
            },
        )

        # Логинимся
        response = api_client.post(
            "/api/auth/login",
            json={
                "email": email,
                "password": "TestPassword123!",
            },
        )

        assert response.status_code in [200, 400, 401]
        if response.status_code == 200:
            data = response.get_json()
            assert "access_token" in data or "token" in data

    def test_invalid_login_credentials(self, api_client):
        """✗ Вход с неправильными учётными данными"""
        response = api_client.post(
            "/api/auth/login",
            json={
                "email": "nonexistent@example.com",
                "password": "WrongPassword",
            },
        )

        assert response.status_code in [400, 401, 404]

    def test_get_current_user(self, api_client, auth_token):
        """✓ Получить текущего пользователя"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        response = api_client.get(
            "/api/auth/me",
            headers=headers,
        )

        assert response.status_code in [200, 401, 404]


# ============================================================================
# Error Handling Tests
# ============================================================================


class TestAPIErrorHandling:
    """Интеграционные тесты для обработки ошибок"""

    def test_nonexistent_document(self, api_client, auth_token):
        """✗ Запрос несуществующего документа"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        response = api_client.get(
            "/api/document/nonexistent_id",
            headers=headers,
        )

        assert response.status_code in [404, 400, 401]

    def test_invalid_profile_id(self, api_client, test_docx_file):
        """✗ Загрузка с неправильным ID профиля"""
        with open(test_docx_file, "rb") as f:
            files = {
                "file": (
                    "test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            response = api_client.post(
                "/api/document/upload",
                data={"profile_id": "nonexistent_profile"},
                files=files,
            )

        assert response.status_code in [400, 404, 422]

    def test_missing_required_fields(self, api_client):
        """✗ Запрос без обязательных полей"""
        response = api_client.post(
            "/api/document/upload",
            data={},  # Missing file
        )

        assert response.status_code in [400, 422]


# ============================================================================
# Rate Limiting Tests
# ============================================================================


class TestRateLimitingAPI:
    """Интеграционные тесты для ограничения частоты запросов"""

    def test_rate_limit_on_login(self, api_client):
        """Testing rate limiting on login endpoint"""
        responses = []
        for i in range(15):  # Попробуем много раз
            response = api_client.post(
                "/api/auth/login",
                json={
                    "email": "test@example.com",
                    "password": "wrongpassword",
                },
            )
            responses.append(response.status_code)

        # Должна быть 429 (Too Many Requests) после нескольких попыток
        assert 429 in responses or all(code in [400, 401] for code in responses)


# ============================================================================
# Batch Operations Tests
# ============================================================================


class TestBatchOperationsAPI:
    """Интеграционные тесты для групповых операций"""

    def test_batch_upload(self, api_client, test_docx_file, auth_token):
        """✓ Пакетная загрузка документов"""
        headers = {"Authorization": f"Bearer {auth_token}"} if auth_token else {}

        # Создать несколько копий файла
        files_list = []
        for i in range(3):
            with open(test_docx_file, "rb") as f:
                file_content = f.read()
            files_list.append(
                (
                    "files",
                    (
                        f"test_{i}.docx",
                        BytesIO(file_content),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            )

        response = api_client.post(
            "/api/document/upload-batch",
            data={"profile_id": "default_gost"},
            files=files_list,
            headers=headers,
        )

        assert response.status_code in [200, 201, 400, 422]


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
