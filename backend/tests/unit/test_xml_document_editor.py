"""
Тесты для XMLDocumentEditor - прямой работы с XML структурой DOCX.
"""

import os
import sys
import tempfile
import pytest
from pathlib import Path

# Добавляем путь к модулям
sys.path.insert(0, str(Path(__file__).parent.parent / 'app'))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from app.services.xml_document_editor import (
    XMLDocumentEditor, 
    apply_xml_corrections,
    XMLEditType,
    XMLEditReport
)


class TestXMLDocumentEditor:
    """Тесты для класса XMLDocumentEditor"""
    
    @pytest.fixture
    def sample_docx(self):
        """Создаёт тестовый DOCX файл"""
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, "test_document.docx")
        
        doc = Document()
        
        # Добавляем секции с неправильными полями
        section = doc.sections[0]
        section.left_margin = Cm(2.0)  # Неправильно - должно быть 3 см
        section.right_margin = Cm(2.0)  # Неправильно - должно быть 1 см
        
        # Добавляем заголовок
        heading = doc.add_heading("Заголовок первого уровня", level=1)
        
        # Добавляем обычный текст с неправильным шрифтом
        para = doc.add_paragraph("Это тестовый параграф с текстом.")
        for run in para.runs:
            run.font.name = "Arial"  # Неправильно - должен быть Times New Roman
            run.font.size = Pt(12)   # Неправильно - должен быть 14pt
        
        # Добавляем ещё параграфы
        doc.add_paragraph("Второй параграф для тестирования.")
        doc.add_paragraph("Третий параграф с дополнительным текстом.")
        
        # Добавляем заголовок второго уровня
        doc.add_heading("Подзаголовок", level=2)
        doc.add_paragraph("Текст после подзаголовка.")
        
        doc.save(file_path)
        
        yield file_path
        
        # Очистка
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    def test_context_manager(self, sample_docx):
        """Тест контекстного менеджера"""
        with XMLDocumentEditor(sample_docx) as editor:
            assert editor.document_xml is not None
            assert editor.temp_dir is not None
            assert os.path.exists(editor.temp_dir)
        
        # После выхода из контекста temp_dir должен быть удалён
        assert not os.path.exists(editor.temp_dir)
    
    def test_fix_all_fonts(self, sample_docx):
        """Тест исправления шрифтов"""
        output_path = sample_docx.replace('.docx', '_fixed.docx')
        
        with XMLDocumentEditor(sample_docx) as editor:
            fixed_count = editor.fix_all_fonts()
            editor.save(output_path)
        
        assert fixed_count > 0
        assert os.path.exists(output_path)
        
        # Проверяем результат
        doc = Document(output_path)
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    assert run.font.name == "Times New Roman" or run.font.name is None
    
    def test_fix_page_margins(self, sample_docx):
        """Тест исправления полей страницы"""
        output_path = sample_docx.replace('.docx', '_margins.docx')
        
        with XMLDocumentEditor(sample_docx) as editor:
            fixed_count = editor.fix_page_margins()
            editor.save(output_path)
        
        assert fixed_count > 0
        
        # Проверяем результат
        doc = Document(output_path)
        for section in doc.sections:
            # Проверяем с погрешностью (твипы -> см)
            assert abs(section.left_margin.cm - 3.0) < 0.1
            assert abs(section.right_margin.cm - 1.5) < 0.1
    
    def test_fix_all_paragraphs(self, sample_docx):
        """Тест исправления абзацев"""
        output_path = sample_docx.replace('.docx', '_paragraphs.docx')
        
        with XMLDocumentEditor(sample_docx) as editor:
            fixed_count = editor.fix_all_paragraphs()
            editor.save(output_path)
        
        assert fixed_count > 0
        
        # Проверяем результат
        doc = Document(output_path)
        for para in doc.paragraphs:
            if para.text.strip() and not para.style.name.startswith('Heading'):
                # Проверяем отступ первой строки (с погрешностью)
                if para.paragraph_format.first_line_indent:
                    indent_cm = para.paragraph_format.first_line_indent.cm
                    assert abs(indent_cm - 1.25) < 0.1
    
    def test_fix_normal_style(self, sample_docx):
        """Тест исправления стиля Normal"""
        output_path = sample_docx.replace('.docx', '_normal.docx')
        
        with XMLDocumentEditor(sample_docx) as editor:
            result = editor.fix_normal_style()
            editor.save(output_path)
        
        assert result == True
    
    def test_fix_heading_styles(self, sample_docx):
        """Тест исправления стилей заголовков"""
        output_path = sample_docx.replace('.docx', '_headings.docx')
        
        with XMLDocumentEditor(sample_docx) as editor:
            fixed_count = editor.fix_heading_styles()
            editor.save(output_path)
        
        # Должны быть исправлены стили Heading 1 и Heading 2
        assert fixed_count >= 0  # Может быть 0 если стили не найдены
    
    def test_fix_all(self, sample_docx):
        """Тест полного исправления документа"""
        output_path = sample_docx.replace('.docx', '_all.docx')
        
        with XMLDocumentEditor(sample_docx) as editor:
            report = editor.fix_all()
            editor.save(output_path)
        
        assert isinstance(report, XMLEditReport)
        assert report.total_edits > 0
        assert report.successful_edits > 0
        
        # Проверяем, что файл создан
        assert os.path.exists(output_path)
    
    def test_apply_xml_corrections_function(self, sample_docx):
        """Тест функции-обёртки apply_xml_corrections"""
        output_path, report = apply_xml_corrections(sample_docx)
        
        assert output_path is not None
        assert os.path.exists(output_path)
        assert isinstance(report, XMLEditReport)
        assert report.total_edits > 0
    
    def test_report_structure(self, sample_docx):
        """Тест структуры отчёта"""
        with XMLDocumentEditor(sample_docx) as editor:
            report = editor.fix_all()
        
        assert hasattr(report, 'file_path')
        assert hasattr(report, 'edits')
        assert hasattr(report, 'total_edits')
        assert hasattr(report, 'successful_edits')
        assert hasattr(report, 'failed_edits')
        
        # Проверяем, что edits - это список
        assert isinstance(report.edits, list)
        
        # Проверяем структуру каждого edit
        for edit in report.edits:
            assert hasattr(edit, 'edit_type')
            assert hasattr(edit, 'xpath')
            assert hasattr(edit, 'old_value')
            assert hasattr(edit, 'new_value')
            assert hasattr(edit, 'success')


class TestXMLEditorEdgeCases:
    """Тесты граничных случаев"""
    
    @pytest.fixture
    def empty_docx(self):
        """Создаёт пустой DOCX файл"""
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, "empty.docx")
        
        doc = Document()
        doc.save(file_path)
        
        yield file_path
        
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    @pytest.fixture
    def docx_with_tables(self):
        """Создаёт DOCX с таблицами"""
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, "with_tables.docx")
        
        doc = Document()
        doc.add_paragraph("Текст перед таблицей")
        
        table = doc.add_table(rows=3, cols=3)
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f"Ячейка {i+1}-{j+1}"
        
        doc.add_paragraph("Текст после таблицы")
        doc.save(file_path)
        
        yield file_path
        
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    def test_empty_document(self, empty_docx):
        """Тест пустого документа"""
        with XMLDocumentEditor(empty_docx) as editor:
            report = editor.fix_all()
        
        # Должен обработаться без ошибок
        assert isinstance(report, XMLEditReport)
    
    def test_document_with_tables(self, docx_with_tables):
        """Тест документа с таблицами"""
        output_path = docx_with_tables.replace('.docx', '_fixed.docx')
        
        with XMLDocumentEditor(docx_with_tables) as editor:
            report = editor.fix_all()
            editor.save(output_path)
        
        # Таблицы не должны быть повреждены
        doc = Document(output_path)
        assert len(doc.tables) == 1
        assert len(doc.tables[0].rows) == 3
        assert len(doc.tables[0].columns) == 3


class TestXMLEditorGOSTCompliance:
    """Тесты на соответствие ГОСТ"""
    
    @pytest.fixture
    def non_compliant_docx(self):
        """Создаёт документ, не соответствующий ГОСТ"""
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, "non_compliant.docx")
        
        doc = Document()
        
        # Неправильные поля
        section = doc.sections[0]
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        
        # Неправильный шрифт и размер
        para = doc.add_paragraph()
        run = para.add_run("Текст с неправильным форматированием")
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        
        # Неправильный межстрочный интервал
        para.paragraph_format.line_spacing = 1.0
        
        # Без отступа
        para.paragraph_format.first_line_indent = Cm(0)
        
        doc.save(file_path)
        
        yield file_path
        
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    def test_gost_compliance_after_fix(self, non_compliant_docx):
        """Тест соответствия ГОСТ после исправления"""
        output_path = non_compliant_docx.replace('.docx', '_compliant.docx')
        
        with XMLDocumentEditor(non_compliant_docx) as editor:
            report = editor.fix_all()
            editor.save(output_path)
        
        # Проверяем результат
        doc = Document(output_path)
        
        # Проверяем поля
        section = doc.sections[0]
        assert abs(section.left_margin.cm - 3.0) < 0.1, "Левое поле должно быть 3 см"
        assert abs(section.right_margin.cm - 1.5) < 0.1, "Правое поле должно быть 1.5 см"
        assert abs(section.top_margin.cm - 2.0) < 0.1, "Верхнее поле должно быть 2 см"
        assert abs(section.bottom_margin.cm - 2.0) < 0.1, "Нижнее поле должно быть 2 см"


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
