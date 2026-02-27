import unittest
import os
import tempfile
import docx
from app.services.preview_service import PreviewService


class TestPreviewService(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()
        self.docx_path = os.path.join(self.temp_dir, "test.docx")

        # Create a simple docx file
        doc = docx.Document()
        doc.add_paragraph("Hello World")
        p = doc.add_paragraph("Bold Text")
        p.runs[0].bold = True
        doc.save(self.docx_path)

        self.service = PreviewService()

    def tearDown(self):
        if os.path.exists(self.docx_path):
            os.remove(self.docx_path)
        os.rmdir(self.temp_dir)

    def test_generate_preview(self):
        html_content = self.service.generate_preview(self.docx_path)
        self.assertIn("Hello World", html_content)
        self.assertIn("Bold Text", html_content)
        self.assertIn("<strong>", html_content)  # mammoth генерирует <strong> для жирного текста
        self.assertIn("<p>", html_content)  # Проверяем наличие параграфов


if __name__ == "__main__":
    unittest.main()
