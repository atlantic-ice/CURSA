"""Тесты для рефакторизованного DocumentCorrector.

Проверяет что новая модульная архитектура работает так же как старая.
"""

import pytest
import os
import tempfile
from docx import Document
from docx.shared import Pt, Cm

# Импортируем новый рефакторизованный корректор
from app.services.document_corrector_refactored import DocumentCorrector, CorrectionReport


@pytest.fixture
def sample_document():
    """Создает простой тестовый документ."""
    doc = Document()

    # Добавляем параграф с неправильным шрифтом
    para = doc.add_paragraph("Это тестовый документ")
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)

    # Неправильные поля
    for section in doc.sections:
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

    return doc


@pytest.fixture
def temp_docx_file(sample_document):
    """Сохраняет тестовый документ во временный файл."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        temp_path = f.name
        sample_document.save(temp_path)

    yield temp_path

    # Cleanup
    if os.path.exists(temp_path):
        os.remove(temp_path)


class TestDocumentCorrector:
    """Тесты для DocumentCorrector."""

    def test_initialization(self):
        """Проверяет инициализацию корректора."""
        corrector = DocumentCorrector()
        assert corrector is not None
        assert len(corrector.correctors) == 4

    def test_correct_document(self, temp_docx_file):
        """Проверяет основную функцию коррекции."""
        corrector = DocumentCorrector()

        # Создаем выходной файл
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            out_path = f.name

        try:
            report = corrector.correct_document(temp_docx_file, out_path=out_path)

            # Проверяем отчет
            assert report is not None
            assert isinstance(report, CorrectionReport)
            assert report.passes_completed > 0
            assert report.total_issues_fixed >= 0
            assert os.path.exists(out_path)

            # Проверяем что новый документ валиден
            corrected_doc = Document(out_path)
            assert corrected_doc is not None

        finally:
            if os.path.exists(out_path):
                os.remove(out_path)

    def test_analyze_document(self, temp_docx_file):
        """Проверяет функцию анализа."""
        corrector = DocumentCorrector()
        analysis = corrector.analyze_document(temp_docx_file)

        assert isinstance(analysis, dict)
        assert "StyleCorrector" in analysis
        assert "StructureCorrector" in analysis

    def test_correct_font(self, temp_docx_file):
        """Проверяет исправление шрифтов."""
        corrector = DocumentCorrector()

        # Создаем выходной файл
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            out_path = f.name

        try:
            report = corrector.correct_document(temp_docx_file, out_path=out_path)

            # Проверяем что шрифты исправлены в выходном файле
            doc = Document(out_path)

            for para in doc.paragraphs:
                if para.text.strip():
                    for run in para.runs:
                        # После коррекции должно быть Times New Roman
                        assert run.font.name == "Times New Roman" or run.font.name is None
        finally:
            if os.path.exists(out_path):
                os.remove(out_path)

    def test_report_summary(self, temp_docx_file):
        """Проверяет сводку отчета."""
        corrector = DocumentCorrector()
        report = corrector.correct_document(temp_docx_file)

        summary = report.get_summary()

        assert "file" in summary
        assert "duration_seconds" in summary
        assert "passes_completed" in summary
        assert "total_issues_fixed" in summary
        assert "success_rate" in summary


class TestStyleCorrector:
    """Тесты для StyleCorrector."""

    def test_style_corrector_initialization(self):
        """Проверяет инициализацию StyleCorrector."""
        from app.services.correctors import StyleCorrector

        corrector = StyleCorrector()
        assert corrector is not None
        assert corrector.rules is not None

    def test_style_corrector_actions(self):
        """Проверяет что корректор отслеживает действия."""
        from app.services.correctors import StyleCorrector

        corrector = StyleCorrector()
        doc = Document()

        # Коррекция
        corrected = corrector.correct(doc)

        # Проверяем действия
        actions = corrector.get_actions()
        assert isinstance(actions, list)

    def test_style_corrector_summary(self):
        """Проверяет сводку StyleCorrector."""
        from app.services.correctors import StyleCorrector

        corrector = StyleCorrector()
        doc = Document()

        corrector.correct(doc)
        summary = corrector.get_summary()

        assert "total_actions" in summary
        assert "successful" in summary
        assert "failed" in summary
        assert "success_rate" in summary


class TestIntegration:
    """Интеграционные тесты."""

    def test_old_api_compatibility(self, temp_docx_file):
        """Проверяет обратную совместимость с старым API."""
        from app.services.document_corrector_refactored import correct_document

        # Старый API должен работать
        report = correct_document(temp_docx_file)
        assert report is not None


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
