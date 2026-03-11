"""Функциональные тесты для потока работы с документами.
Тестирует полный процесс от загрузки документа до скачивания исправленного документа.
"""
import os
import io
import pytest
import tempfile
from pathlib import Path
from docx import Document

# Путь к тестовым данным
TEST_DATA_DIR = Path(__file__).parent.parent / "test_data"


def _count_font_issues(check_results):
    issues = check_results.get('issues', []) if isinstance(check_results, dict) else []
    font_issue_count = 0

    for issue in issues:
        issue_type = str(issue.get('type', '')).lower()
        rule_name = str(issue.get('rule_name', '')).lower()
        description = str(issue.get('description', '')).lower()
        if 'font' in issue_type or 'шрифт' in rule_name or 'шрифт' in description:
            font_issue_count += 1

    return font_issue_count


def test_complete_document_flow(client):
    """Тестирует полный поток работы с документом.

    1. Загрузка документа для анализа через /api/document/upload
    2. Получение результатов проверки
    3. Отправка запроса на исправление через /api/document/correct
    4. Скачивание исправленного документа через /api/document/download-corrected
    """
    # Создаём тестовый документ с ошибками
    with tempfile.TemporaryDirectory() as tmpdir:
        test_file_path = os.path.join(tmpdir, 'test_flow.docx')

        doc = Document()
        # Добавляем текст с неправильным шрифтом
        p = doc.add_paragraph("Тестовый документ для проверки полного потока")
        for run in p.runs:
            run.font.name = 'Arial'  # Неправильный шрифт
        doc.save(test_file_path)

        # 1. Загрузка документа для анализа
        with open(test_file_path, 'rb') as f:
            response_upload = client.post(
                '/api/document/upload',
                data={'file': (io.BytesIO(f.read()), 'test_flow.docx')},
                content_type='multipart/form-data'
            )

        assert response_upload.status_code == 200, f"Upload failed: {response_upload.json}"
        upload_result = response_upload.json

        # Проверяем результат загрузки
        assert 'check_results' in upload_result, "check_results not in response"
        assert 'temp_path' in upload_result, "temp_path not in response"

        temp_path = upload_result['temp_path']
        check_results = upload_result['check_results']

        # 2. Отправка запроса на исправление
        correction_data = {
            'file_path': temp_path,
            'original_filename': 'test_flow.docx',
            'errors_to_fix': check_results.get('issues', [])
        }

        response_correction = client.post('/api/document/correct', json=correction_data)

        assert response_correction.status_code == 200, f"Correction failed: {response_correction.json}"
        correction_result = response_correction.json
        assert 'corrected_file_path' in correction_result, "corrected_file_path not in response"

        corrected_file_path = correction_result['corrected_file_path']

        # 3. Скачивание исправленного документа
        download_url = f"/api/document/download-corrected?path={corrected_file_path}"
        response_download = client.get(download_url)

        assert response_download.status_code == 200, f"Download failed with status {response_download.status_code}"
        assert len(response_download.data) > 0, "Downloaded file is empty"

        # Проверяем что это валидный DOCX
        content_type = response_download.headers.get('Content-Type', '')
        assert 'application/vnd.openxmlformats-officedocument' in content_type or len(response_download.data) > 1000


def test_analyze_autocorrect_download_flow(client):
    """Тестирует новый поток analyze -> autocorrect(token) -> download."""
    with tempfile.TemporaryDirectory() as tmpdir:
        test_file_path = os.path.join(tmpdir, 'test_token_flow.docx')

        doc = Document()
        paragraph = doc.add_paragraph("Тестовый документ для токенного потока")
        for run in paragraph.runs:
            run.font.name = 'Arial'
        doc.save(test_file_path)

        with open(test_file_path, 'rb') as f:
            response_analyze = client.post(
                '/api/document/analyze',
                data={'file': (io.BytesIO(f.read()), 'test_token_flow.docx')},
                content_type='multipart/form-data'
            )

        assert response_analyze.status_code == 200, (
            f"Analyze failed: {response_analyze.json}"
        )
        analyze_result = response_analyze.json

        assert 'check_results' in analyze_result, 'check_results not in analyze response'
        assert 'document_token' in analyze_result, 'document_token not in analyze response'
        assert 'temp_path' not in analyze_result, 'temp_path leaked in analyze response'

        initial_check_results = analyze_result['check_results']
        initial_total_issues = initial_check_results.get('total_issues_count', 0)
        initial_font_issues = _count_font_issues(initial_check_results)
        assert initial_total_issues > 0
        assert initial_font_issues > 0

        document_token = analyze_result['document_token']

        response_autocorrect = client.post(
            '/api/document/autocorrect',
            json={
                'document_token': document_token,
                'original_filename': 'test_token_flow.docx',
            }
        )

        assert response_autocorrect.status_code == 200, (
            f"Autocorrect failed: {response_autocorrect.json}"
        )
        autocorrect_result = response_autocorrect.json
        assert autocorrect_result.get('success') is True
        assert 'corrected_file_path' in autocorrect_result
        assert 'improvement' in autocorrect_result
        assert 'check_results' in autocorrect_result
        assert 'original_preview_path' in autocorrect_result
        assert autocorrect_result['improvement']['resolved_font_issues'] > 0
        assert (
            autocorrect_result['check_results']['total_issues_count']
            == autocorrect_result['improvement']['after_total_issues']
        )

        corrected_file_path = autocorrect_result['corrected_file_path']
        download_url = f"/api/document/download-corrected?path={corrected_file_path}"
        response_download = client.get(download_url)

        assert response_download.status_code == 200, (
            f"Download failed with status {response_download.status_code}"
        )
        assert len(response_download.data) > 0, 'Downloaded corrected file is empty'

        preview_response = client.post(
            '/api/preview/generate',
            json={'path': autocorrect_result['original_preview_path']}
        )
        assert preview_response.status_code == 200
        assert 'html' in preview_response.json

        content_type = response_download.headers.get('Content-Type', '')
        assert (
            'application/vnd.openxmlformats-officedocument' in content_type
            or len(response_download.data) > 1000
        )

        reanalyze_path = os.path.join(tmpdir, 'reanalyze_corrected.docx')
        with open(reanalyze_path, 'wb') as corrected_file:
            corrected_file.write(response_download.data)

        with open(reanalyze_path, 'rb') as corrected_file:
            response_reanalyze = client.post(
                '/api/document/analyze',
                data={
                    'file': (io.BytesIO(corrected_file.read()), 'reanalyze_corrected.docx')
                },
                content_type='multipart/form-data'
            )

        assert response_reanalyze.status_code == 200, (
            f"Reanalyze failed: {response_reanalyze.json}"
        )
        reanalyze_result = response_reanalyze.json
        updated_check_results = reanalyze_result['check_results']
        updated_total_issues = updated_check_results.get('total_issues_count', 0)
        updated_font_issues = _count_font_issues(updated_check_results)

        assert autocorrect_result['check_results']['total_issues_count'] == updated_total_issues
        assert updated_total_issues <= initial_total_issues
        assert updated_font_issues < initial_font_issues
        assert autocorrect_result['improvement']['after_total_issues'] <= initial_total_issues
        assert autocorrect_result['improvement']['resolved_total_issues'] >= 0
        assert autocorrect_result['improvement']['after_font_issues'] <= initial_font_issues

        response_reuse = client.post(
            '/api/document/autocorrect',
            json={'document_token': document_token}
        )
        assert response_reuse.status_code == 404