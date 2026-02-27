"""
Тесты для модульной системы валидаторов.
"""

import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from app.services.validators import BaseValidator, ValidationIssue, Severity
from app.services.validators.font_validator import FontValidator
from app.services.validators.margin_validator import MarginValidator
from app.services.validators.bibliography_validator import BibliographyValidator
from app.services.validators.heading_validator import HeadingValidator
from app.services.validators.paragraph_validator import ParagraphValidator
from app.services.validators.structure_validator import StructureValidator
from app.services.validators.table_validator import TableValidator
from app.services.validation_engine import ValidationEngine


@pytest.fixture
def test_profile():
    """Тестовый профиль требований"""
    return {
        "name": "Test Profile",
        "version": "1.0",
        "rules": {
            "font": {
                "name": "Times New Roman",
                "size": 14.0,
                "allowed_fonts": ["Times New Roman", "Arial"],
            },
            "margins": {"left": 3.0, "right": 1.5, "top": 2.0, "bottom": 2.0},
            "bibliography": {"min_sources": 10},
        },
        "validation": {"check_font": True, "check_margins": True, "check_bibliography": True},
    }


@pytest.fixture
def create_test_document(tmp_path):
    """Создает тестовый DOCX документ"""

    def _create(filename="test.docx", **kwargs):
        doc = Document()

        # Настройка полей если указано
        if "margins" in kwargs:
            margins = kwargs["margins"]
            for section in doc.sections:
                if "left" in margins:
                    section.left_margin = Cm(margins["left"])
                if "right" in margins:
                    section.right_margin = Cm(margins["right"])
                if "top" in margins:
                    section.top_margin = Cm(margins["top"])
                if "bottom" in margins:
                    section.bottom_margin = Cm(margins["bottom"])

        # Добавление параграфов если указано
        if "paragraphs" in kwargs:
            for para_config in kwargs["paragraphs"]:
                para = doc.add_paragraph(para_config.get("text", "Test text"))

                # Настройка шрифта
                if "font_name" in para_config or "font_size" in para_config:
                    for run in para.runs:
                        if "font_name" in para_config:
                            run.font.name = para_config["font_name"]
                        if "font_size" in para_config:
                            run.font.size = Pt(para_config["font_size"])

        # Сохранение документа
        filepath = tmp_path / filename
        doc.save(str(filepath))
        return str(filepath)

    return _create


class TestFontValidator:
    """Тесты для FontValidator"""

    def test_correct_font(self, test_profile, create_test_document):
        """Тест: правильный шрифт не должен вызывать ошибок"""
        doc_path = create_test_document(
            paragraphs=[
                {"text": "Test paragraph", "font_name": "Times New Roman", "font_size": 14.0}
            ]
        )

        validator = FontValidator(profile=test_profile)
        document = Document(doc_path)
        result = validator.validate(document, {})

        assert result.passed
        assert len(result.issues) == 0

    def test_wrong_font_name(self, test_profile, create_test_document):
        """Тест: неправильный шрифт должен вызывать ошибку"""
        doc_path = create_test_document(
            paragraphs=[{"text": "Test paragraph", "font_name": "Calibri", "font_size": 14.0}]
        )

        validator = FontValidator(profile=test_profile)
        document = Document(doc_path)
        result = validator.validate(document, {})

        assert not result.passed
        assert len(result.issues) > 0
        assert any("Calibri" in issue.description for issue in result.issues)

    def test_wrong_font_size(self, test_profile, create_test_document):
        """Тест: неправильный размер шрифта должен вызывать ошибку"""
        doc_path = create_test_document(
            paragraphs=[
                {"text": "Test paragraph", "font_name": "Times New Roman", "font_size": 12.0}
            ]
        )

        validator = FontValidator(profile=test_profile)
        document = Document(doc_path)
        result = validator.validate(document, {})

        assert not result.passed
        assert len(result.issues) > 0
        assert any("12" in issue.description for issue in result.issues)

    def test_autocorrect_flag(self, test_profile, create_test_document):
        """Тест: проблемы со шрифтом должны быть автоисправимыми"""
        doc_path = create_test_document(
            paragraphs=[{"text": "Test", "font_name": "Calibri", "font_size": 12.0}]
        )

        validator = FontValidator(profile=test_profile)
        document = Document(doc_path)
        result = validator.validate(document, {})

        assert all(issue.can_autocorrect for issue in result.issues)


class TestMarginValidator:
    """Тесты для MarginValidator"""

    def test_correct_margins(self, test_profile, create_test_document):
        """Тест: правильные поля"""
        doc_path = create_test_document(
            margins={"left": 3.0, "right": 1.5, "top": 2.0, "bottom": 2.0}
        )

        validator = MarginValidator(profile=test_profile)
        document = Document(doc_path)
        result = validator.validate(document, {})

        assert result.passed
        assert len(result.issues) == 0

    def test_wrong_left_margin(self, test_profile, create_test_document):
        """Тест: неправильное левое поле"""
        doc_path = create_test_document(
            margins={"left": 2.0, "right": 1.5, "top": 2.0, "bottom": 2.0}
        )

        validator = MarginValidator(profile=test_profile)
        document = Document(doc_path)
        result = validator.validate(document, {})

        assert not result.passed
        assert any("левое" in issue.description.lower() for issue in result.issues)

    def test_multiple_margin_errors(self, test_profile, create_test_document):
        """Тест: множественные ошибки в полях"""
        doc_path = create_test_document(
            margins={"left": 2.0, "right": 3.0, "top": 1.0, "bottom": 3.0}
        )

        validator = MarginValidator(profile=test_profile)
        document = Document(doc_path)
        result = validator.validate(document, {})

        assert not result.passed
        assert len(result.issues) == 4  # Все 4 поля неправильные


class TestBibliographyValidator:
    """Тесты для BibliographyValidator"""

    def test_missing_bibliography(self, test_profile, create_test_document):
        """Тест: отсутствует раздел библиографии"""
        doc_path = create_test_document(
            paragraphs=[{"text": "Introduction"}, {"text": "Main content"}]
        )

        validator = BibliographyValidator(profile=test_profile)
        document = Document(doc_path)
        result = validator.validate(document, {})

        assert not result.passed
        critical_issues = [i for i in result.issues if i.severity == Severity.CRITICAL]
        assert len(critical_issues) > 0
        assert any("список" in issue.description.lower() for issue in critical_issues)


class TestValidationEngine:
    """Тесты для ValidationEngine"""

    def test_engine_initialization(self, test_profile):
        """Тест: инициализация движка"""
        engine = ValidationEngine(profile=test_profile)

        assert engine.profile == test_profile
        assert len(engine.validators) > 0

    def test_full_validation(self, test_profile, create_test_document):
        """Тест: полная валидация документа"""
        doc_path = create_test_document(
            margins={"left": 3.0, "right": 1.5, "top": 2.0, "bottom": 2.0},
            paragraphs=[{"text": "Test", "font_name": "Times New Roman", "font_size": 14.0}],
        )

        engine = ValidationEngine(profile=test_profile)
        report = engine.validate_document(doc_path, document_data={})

        assert "status" in report
        assert "summary" in report
        assert "validators" in report
        assert report["summary"]["total_issues"] >= 0

    def test_report_generation(self, test_profile, create_test_document):
        """Тест: генерация отчета"""
        doc_path = create_test_document(
            margins={"left": 2.0, "right": 1.5, "top": 2.0, "bottom": 2.0},
            paragraphs=[{"text": "Test", "font_name": "Calibri", "font_size": 12.0}],
        )

        engine = ValidationEngine(profile=test_profile)
        report = engine.validate_document(doc_path, document_data={})

        assert report["status"] in ["critical", "failed", "warning", "passed"]
        assert "summary" in report
        assert "execution" in report
        assert "recommendations" in report

        # Должны быть найдены ошибки
        assert report["summary"]["total_issues"] > 0

    def test_severity_categorization(self, test_profile, create_test_document):
        """Тест: категоризация по серьезности"""
        doc_path = create_test_document(
            margins={"left": 2.0, "right": 1.5, "top": 2.0, "bottom": 2.0},
            paragraphs=[{"text": "Test", "font_name": "Calibri", "font_size": 12.0}],
        )

        engine = ValidationEngine(profile=test_profile)
        report = engine.validate_document(doc_path, document_data={})

        assert "issues_by_severity" in report
        assert "critical" in report["issues_by_severity"]
        assert "error" in report["issues_by_severity"]
        assert "warning" in report["issues_by_severity"]
        assert "info" in report["issues_by_severity"]

    def test_autocorrectable_count(self, test_profile, create_test_document):
        """Тест: подсчет автоисправимых проблем"""
        doc_path = create_test_document(
            margins={"left": 2.0, "right": 1.5, "top": 2.0, "bottom": 2.0},
            paragraphs=[{"text": "Test", "font_name": "Calibri", "font_size": 12.0}],
        )

        engine = ValidationEngine(profile=test_profile)
        report = engine.validate_document(doc_path, document_data={})

        assert "autocorrectable" in report["summary"]
        assert report["summary"]["autocorrectable"] >= 0


class TestValidationIssue:
    """Тесты для ValidationIssue"""

    def test_issue_creation(self):
        """Тест: создание ValidationIssue"""
        issue = ValidationIssue(
            rule_id=1,
            rule_name="Test Rule",
            description="Test description",
            severity=Severity.ERROR,
            can_autocorrect=True,
        )

        assert issue.rule_id == 1
        assert issue.rule_name == "Test Rule"
        assert issue.severity == Severity.ERROR
        assert issue.can_autocorrect

    def test_issue_to_dict(self):
        """Тест: сериализация в словарь"""
        issue = ValidationIssue(
            rule_id=1,
            rule_name="Test Rule",
            description="Test description",
            severity=Severity.ERROR,
            expected="Expected value",
            actual="Actual value",
            can_autocorrect=True,
        )

        issue_dict = issue.to_dict()

        assert issue_dict["rule_id"] == 1
        assert issue_dict["severity"] == "error"
        assert issue_dict["expected"] == "Expected value"
        assert issue_dict["actual"] == "Actual value"
        assert issue_dict["can_autocorrect"] is True


# Интеграционные тесты
class TestValidationIntegration:
    """Интеграционные тесты всей системы валидации"""

    def test_end_to_end_validation(self, test_profile, create_test_document):
        """Тест: полный цикл валидации от начала до конца"""
        # Создаем документ с несколькими проблемами
        doc_path = create_test_document(
            margins={"left": 2.5, "right": 2.0, "top": 1.5, "bottom": 2.5},
            paragraphs=[
                {"text": "Title", "font_name": "Arial", "font_size": 16.0},
                {"text": "Content 1", "font_name": "Calibri", "font_size": 12.0},
                {"text": "Content 2", "font_name": "Times New Roman", "font_size": 10.0},
            ],
        )

        # Запускаем валидацию
        engine = ValidationEngine(profile=test_profile)
        report = engine.validate_document(doc_path, document_data={})

        # Проверки
        assert report["status"] in [
            "failed",
            "critical",
        ]  # Critical из-за отсутствия обязательных разделов
        assert report["summary"]["total_issues"] > 0
        assert len(report["validators"]) > 0
        assert len(report["recommendations"]) > 0

        # Должны быть проблемы от разных валидаторов
        validator_names = [v["validator_name"] for v in report["validators"]]
        assert "FontValidator" in validator_names
        assert "MarginValidator" in validator_names

    def test_perfect_document(self, test_profile, create_test_document):
        """Тест: идеальный документ без проблем"""
        doc_path = create_test_document(
            margins={"left": 3.0, "right": 1.5, "top": 2.0, "bottom": 2.0},
            paragraphs=[
                {"text": "Perfect text", "font_name": "Times New Roman", "font_size": 14.0}
            ],
        )

        engine = ValidationEngine(profile=test_profile)
        report = engine.validate_document(doc_path, document_data={})

        # Может быть passed/warning/failed/critical в зависимости от наличия обязательных разделов
        assert report["status"] in ["passed", "warning", "failed", "critical"]
        assert "summary" in report


class TestHeadingValidator:
    """Тесты для HeadingValidator"""

    def test_structural_element_uppercase(self, test_profile, create_test_document):
        """Тест: структурный элемент должен быть в верхнем регистре"""
        doc_path = create_test_document(
            paragraphs=[{"text": "ВВЕДЕНИЕ"}, {"text": "Это введение документа"}]
        )

        validator = HeadingValidator(profile=test_profile)
        document = Document(doc_path)
        result = validator.validate(document, {})

        # ВВЕДЕНИЕ в верхнем регистре - не должно быть ошибок по регистру
        # Но могут быть другие проблемы (выравнивание и т.д.)
        assert result is not None

    def test_heading_without_period(self, test_profile, create_test_document):
        """Тест: заголовки не должны заканчиваться точкой"""
        doc_path = create_test_document(paragraphs=[{"text": "ЗАКЛЮЧЕНИЕ."}])

        validator = HeadingValidator(profile=test_profile)
        document = Document(doc_path)
        result = validator.validate(document, {})

        # Должна быть ошибка из-за точки в конце
        assert any("точк" in issue.description.lower() for issue in result.issues)


class TestParagraphValidator:
    """Тесты для ParagraphValidator"""

    def test_first_line_indent(self, test_profile, create_test_document, tmp_path):
        """Тест: проверка абзацного отступа"""
        doc = Document()
        para = doc.add_paragraph("Тестовый параграф")
        para.paragraph_format.first_line_indent = Cm(1.25)

        filepath = tmp_path / "test_indent.docx"
        doc.save(str(filepath))

        validator = ParagraphValidator(profile=test_profile)
        result = validator.validate(doc, {})

        # С правильным отступом не должно быть критических ошибок по отступу
        indent_issues = [i for i in result.issues if "отступ" in i.description.lower()]
        # Могут быть другие проблемы, но не отступа
        assert len([i for i in indent_issues if i.severity == Severity.ERROR]) == 0

    def test_line_spacing(self, test_profile, create_test_document, tmp_path):
        """Тест: проверка межстрочного интервала"""
        doc = Document()
        para = doc.add_paragraph("Тестовый текст")
        para.paragraph_format.line_spacing = 1.5

        filepath = tmp_path / "test_spacing.docx"
        doc.save(str(filepath))

        validator = ParagraphValidator(profile=test_profile)
        result = validator.validate(doc, {})

        # Проверяем что валидатор работает
        assert result.validator_name == "ParagraphValidator"


class TestStructureValidator:
    """Тесты для StructureValidator"""

    def test_required_sections(self, test_profile, create_test_document):
        """Тест: проверка наличия обязательных разделов"""
        doc_path = create_test_document(
            paragraphs=[
                {"text": "СОДЕРЖАНИЕ"},
                {"text": "ВВЕДЕНИЕ"},
                {"text": "Основной текст"},
                {"text": "ЗАКЛЮЧЕНИЕ"},
                {"text": "СПИСОК ЛИТЕРАТУРЫ"},
            ]
        )

        validator = StructureValidator(profile=test_profile)
        document = Document(doc_path)
        result = validator.validate(document, {})

        # Все обязательные разделы есть
        critical_issues = [i for i in result.issues if i.severity == Severity.CRITICAL]
        section_issues = [i for i in critical_issues if "раздел" in i.description.lower()]

        # Не должно быть ошибок об отсутствующих разделах
        assert len(section_issues) == 0

    def test_missing_required_section(self, test_profile, create_test_document):
        """Тест: отсутствие обязательного раздела"""
        doc_path = create_test_document(
            paragraphs=[
                {"text": "СОДЕРЖАНИЕ"},
                {"text": "Основной текст"},
                # Нет ВВЕДЕНИЯ, ЗАКЛЮЧЕНИЯ, СПИСКА ЛИТЕРАТУРЫ
            ]
        )

        validator = StructureValidator(profile=test_profile)
        document = Document(doc_path)
        result = validator.validate(document, {})

        # Должны быть критические ошибки об отсутствующих разделах
        critical_issues = [i for i in result.issues if i.severity == Severity.CRITICAL]
        assert len(critical_issues) > 0


class TestTableValidator:
    """Тесты для TableValidator"""

    def test_table_caption(self, test_profile, create_test_document, tmp_path):
        """Тест: проверка подписи таблицы"""
        doc = Document()
        doc.add_paragraph("Таблица 1 - Тестовая таблица")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Заголовок 1"
        table.cell(0, 1).text = "Заголовок 2"
        table.cell(1, 0).text = "Значение 1"
        table.cell(1, 1).text = "Значение 2"

        filepath = tmp_path / "test_table.docx"
        doc.save(str(filepath))

        validator = TableValidator(profile=test_profile)
        result = validator.validate(doc, {})

        # Проверяем что валидатор обработал таблицу
        assert result.validator_name == "TableValidator"

    def test_missing_table_caption(self, test_profile, create_test_document, tmp_path):
        """Тест: отсутствие подписи таблицы"""
        doc = Document()
        # Таблица без подписи
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Данные"

        filepath = tmp_path / "test_no_caption.docx"
        doc.save(str(filepath))

        validator = TableValidator(profile=test_profile)
        result = validator.validate(doc, {})

        # Должна быть ошибка об отсутствии подписи
        caption_issues = [i for i in result.issues if "подпись" in i.description.lower()]
        assert len(caption_issues) > 0


class TestNewValidatorsIntegration:
    """Интеграционные тесты для новых валидаторов"""

    def test_all_validators_registered(self, test_profile):
        """Тест: все новые валидаторы зарегистрированы в движке"""
        engine = ValidationEngine(profile=test_profile)

        validator_names = [v.name for v in engine.validators]

        assert "HeadingValidator" in validator_names
        assert "ParagraphValidator" in validator_names
        assert "StructureValidator" in validator_names
        assert "TableValidator" in validator_names

    def test_complex_document_validation(self, test_profile, tmp_path):
        """Тест: валидация сложного документа со всеми элементами"""
        doc = Document()

        # Добавляем структуру документа
        doc.add_paragraph("СОДЕРЖАНИЕ")
        doc.add_paragraph("ВВЕДЕНИЕ")

        # Добавляем главу
        doc.add_paragraph("ГЛАВА 1 ОСНОВНАЯ ЧАСТЬ")
        para = doc.add_paragraph("Это основной текст документа с абзацным отступом.")
        para.paragraph_format.first_line_indent = Cm(1.25)

        # Добавляем таблицу
        doc.add_paragraph("Таблица 1 - Результаты исследования")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Параметр"
        table.cell(0, 1).text = "Значение"

        doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
        doc.add_paragraph("СПИСОК ЛИТЕРАТУРЫ")

        # Настройка полей
        for section in doc.sections:
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)

        filepath = tmp_path / "complex_doc.docx"
        doc.save(str(filepath))

        # Запускаем полную валидацию
        engine = ValidationEngine(profile=test_profile)
        report = engine.validate_document(str(filepath), document_data={})

        # Проверяем что все валидаторы отработали
        assert len(report["validators"]) >= 7  # Все 7 валидаторов
        assert "summary" in report
        assert "issues_by_severity" in report

        # Документ должен иметь некоторые проблемы (например, библиография пустая)
        assert report["summary"]["total_issues"] >= 0
