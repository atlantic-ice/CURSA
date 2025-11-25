import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from services.document_corrector import DocumentCorrector

def create_bad_document(filename):
    doc = Document()
    
    # print("Setting margins...")
    # # Wrong margins
    # section = doc.sections[0]
    # section.left_margin = Cm(1.0)
    # section.right_margin = Cm(3.0)
    
    # print("Setting style...")
    # # Wrong font and spacing in Normal style
    # style = doc.styles['Normal']
    # style.font.name = 'Arial'
    # style.font.size = Pt(10)
    # style.paragraph_format.line_spacing = 1.0
    
    # print("Adding paragraphs...")
    # # Title page text (should be detected but maybe not fixed if disabled)
    # doc.add_paragraph("МИНИСТЕРСТВО ОБРАЗОВАНИЯ")
    # doc.add_paragraph("КУРСОВАЯ РАБОТА")
    
    # print("Adding pseudo-heading...")
    # # Pseudo-heading
    # p = doc.add_paragraph("ВВЕДЕНИЕ")
    # p.runs[0].bold = True
    # p.runs[0].font.size = Pt(16)
    # print("Setting alignment...")
    # try:
    #     p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    #     print("Alignment set.")
    # except Exception as e:
    #     print(f"Error setting alignment: {e}")
    
    # print("Adding text...")
    # # Text with wrong font
    # p = doc.add_paragraph("This is some text with wrong font.")
    # p.runs[0].font.name = 'Calibri'
    # p.runs[0].font.size = Pt(10)
    
    # print("Adding list...")
    # # List with wrong indentation
    # p = doc.add_paragraph("Item 1")
    # p.style = 'List Bullet'
    # p.paragraph_format.left_indent = Cm(0.5)
    
    # print("Adding table...")
    # # Table with wrong formatting
    # table = doc.add_table(rows=2, cols=2)
    # cell = table.cell(0, 0)
    # cell.text = "Header 1"
    # cell.paragraphs[0].runs[0].font.name = 'Arial'
    
    print("Saving...")
    doc.save(filename)
    print("Saved.")
    return filename

def verify_document(filename):
    doc = Document(filename)
    errors = []
    
    # Check margins
    section = doc.sections[0]
    if section.left_margin.cm != 3.0:
        errors.append(f"Left margin is {section.left_margin.cm}, expected 3.0")
    if section.right_margin.cm != 1.0:
        errors.append(f"Right margin is {section.right_margin.cm}, expected 1.0")
        
    # Check font (sampling)
    # Note: we check runs because style defaults might be overridden or not applied to runs if they were manual
    # But our corrector sets style defaults AND run properties
    
    # Check normal text
    for para in doc.paragraphs:
        if "This is some text" in para.text:
            for run in para.runs:
                if run.font.name != 'Times New Roman':
                    errors.append(f"Font name is {run.font.name}, expected Times New Roman")
                if run.font.size != Pt(14):
                    errors.append(f"Font size is {run.font.size.pt}, expected 14.0")
    
    # Check heading (promoted)
    found_heading = False
    for para in doc.paragraphs:
        if "ВВЕДЕНИЕ" in para.text:
            found_heading = True
            if para.style.name != 'Heading 1':
                errors.append(f"Heading style is {para.style.name}, expected Heading 1")
            if para.paragraph_format.alignment != 1: # WD_PARAGRAPH_ALIGNMENT.CENTER
                errors.append("Heading alignment is not CENTER")
    
    if not found_heading:
        errors.append("Heading 'ВВЕДЕНИЕ' not found")
        
    if not errors:
        print("VERIFICATION PASSED!")
        return True
    else:
        print("VERIFICATION FAILED:")
        for e in errors:
            print(f"- {e}")
        return False

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))
    bad_doc_path = os.path.join(base_dir, "test_bad.docx")
    corrected_doc_path = os.path.join(base_dir, "test_corrected.docx")
    
    print("Creating bad document...")
    create_bad_document(bad_doc_path)
    
    print("Running corrector...")
    corrector = DocumentCorrector()
    corrector.correct_document(bad_doc_path, out_path=corrected_doc_path)
    
    print("Verifying result...")
    verify_document(corrected_doc_path)
    
    # Cleanup
    if os.path.exists(bad_doc_path):
        os.remove(bad_doc_path)
    # Keep corrected doc for inspection if needed, or remove
    # if os.path.exists(corrected_doc_path):
    #     os.remove(corrected_doc_path)
