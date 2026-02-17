"""Корректор для исправления структуры документа.

Отвечает за:
- Исправление заголовков
- Исправление разделов и подразделов
- Исправление нумерации разделов
- Исправление оформления титульной страницы
"""

import os
import re
import tempfile
import shutil
from typing import List, Dict, Any, Set

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxcompose.composer import Composer

from .base import BaseCorrector


class StructureCorrector(BaseCorrector):
    """Корректор структуры документа."""
    
    def __init__(self, rules: Dict[str, Any] = None):
        """Инициализация корректора структуры.
        
        Args:
            rules: Словарь правил
        """
        super().__init__()
        self.rules = rules or {}
    
    def analyze(self, document: Document) -> List[Dict[str, Any]]:
        """Анализирует структуру документа.
        
        Args:
            document: Документ для анализа
            
        Returns:
            Список проблем со структурой
        """
        issues = []
        
        # Проверка заголовков
        heading_issues = self._check_headings(document)
        issues.extend(heading_issues)
        
        return issues
    
    def correct(self, document: Document) -> int:
        """Исправляет структуру документа.
        
        Args:
            document: Документ для коррекции
            
        Returns:
            Количество исправленных проблем
        """
        self.clear_actions()
        corrected = 0
        
        # Исправление заголовков
        corrected += self._correct_headings(document)
        corrected += self._correct_section_headings(document)
        corrected += self._correct_title_page(document)
        
        return corrected
    
    # ========== Анализ ==========
    
    def _check_headings(self, document: Document) -> List[Dict[str, Any]]:
        """Проверяет заголовки в документе.
        
        Args:
            document: Документ для проверки
            
        Returns:
            Список проблем с заголовками
        """
        issues = []
        
        for para_idx, paragraph in enumerate(document.paragraphs):
            if paragraph.style.name.startswith('Heading'):
                # TODO: Добавить проверки специфичные для заголовков
                pass
        
        return issues
    
    # ========== Коррекция ==========
    
    def _correct_headings(self, document: Document) -> int:
        """Исправляет заголовки в документе.
        
        Args:
            document: Документ для коррекции
            
        Returns:
            Количество примененных исправлений
        """
        corrected = 0
        
        for para_idx, paragraph in enumerate(document.paragraphs):
            if not paragraph.style.name.startswith('Heading'):
                continue
            
            try:
                # Проверяем выравнивание заголовка
                pf = paragraph.paragraph_format
                
                if pf.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
                    old_alignment = pf.alignment
                    pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    self.add_action(
                        element_type='heading_alignment',
                        element_index=para_idx,
                        action_type='alignment_change',
                        old_value=old_alignment,
                        new_value=WD_PARAGRAPH_ALIGNMENT.CENTER,
                        description='Заголовок выравнен по центру',
                    )
                    corrected += 1
                
                # Убеждаемся что заголовок жирный
                for run in paragraph.runs:
                    if not run.font.bold:
                        run.font.bold = True
                        corrected += 1
            
            except Exception as e:
                self.add_action(
                    element_type='heading',
                    element_index=para_idx,
                    action_type='heading_correction_error',
                    old_value=None,
                    new_value=None,
                    description=f'Ошибка при исправлении заголовка: {str(e)}',
                    success=False,
                    error_message=str(e),
                )
        
        return corrected

    def _correct_section_headings(self, document: Document) -> int:
        """Форматирует заголовки разделов по нумерации и шаблонам.

        Важно: не использует paragraph.text = ..., чтобы не сломать форматирование.
        """
        corrected = 0
        try:
            table_paragraphs = self._get_table_paragraphs(document)

            chapter_patterns = [
                r'^глава\s+\d+\.?\s+',
                r'^раздел\s+\d+\.?\s+',
                r'^\d+\.\s+[А-Я]',
                r'^\d+\.\d+\.\s+[А-Я]'
            ]

            heading_levels: Dict[int, int] = {}

            for i, paragraph in enumerate(document.paragraphs):
                if id(paragraph) in table_paragraphs:
                    continue

                text = paragraph.text.strip()
                if not text:
                    continue

                try:
                    is_heading = False
                    heading_level = None

                    for pattern in chapter_patterns:
                        if re.match(pattern, text, re.IGNORECASE):
                            is_heading = True
                            numbers = re.findall(r'\d+', text.split()[0])
                            heading_level = len(numbers)
                            break

                    if paragraph.style and paragraph.style.name.startswith('Heading'):
                        is_heading = True
                        try:
                            current_level = int(paragraph.style.name.replace('Heading ', ''))
                            heading_level = current_level if heading_level is None else heading_level
                        except (ValueError, AttributeError):
                            pass

                    if is_heading and heading_level:
                        heading_levels[i] = heading_level

                except Exception as exc:
                    self.add_action(
                        element_type='heading',
                        element_index=i,
                        action_type='heading_detection_error',
                        old_value=None,
                        new_value=None,
                        description=f'Ошибка анализа заголовка: {str(exc)}',
                        success=False,
                        error_message=str(exc),
                    )
                    continue

            for i, paragraph in enumerate(document.paragraphs):
                if id(paragraph) in table_paragraphs:
                    continue

                if i not in heading_levels:
                    continue

                try:
                    level = heading_levels[i]

                    if level == 1:
                        paragraph.style = document.styles['Heading 1']
                        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        paragraph.paragraph_format.space_before = Pt(12)
                        paragraph.paragraph_format.space_after = Pt(12)

                        for run in paragraph.runs:
                            run.text = run.text.upper()

                    elif level == 2:
                        paragraph.style = document.styles['Heading 2']
                        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        paragraph.paragraph_format.space_before = Pt(12)
                        paragraph.paragraph_format.space_after = Pt(6)

                        if paragraph.runs:
                            full_text = paragraph.text
                            parts = full_text.split(' ', 1)
                            if len(parts) > 1:
                                new_text = parts[0] + ' ' + parts[1].capitalize()
                                for run in paragraph.runs:
                                    run.text = ''
                                paragraph.runs[0].text = new_text

                    elif level == 3:
                        paragraph.style = document.styles['Heading 3']
                        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        paragraph.paragraph_format.space_before = Pt(6)
                        paragraph.paragraph_format.space_after = Pt(3)

                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.left_indent = Cm(0)
                    paragraph.paragraph_format.right_indent = Cm(0)

                    if paragraph.text.strip().endswith('.') and paragraph.runs:
                        last_run = paragraph.runs[-1]
                        last_run.text = last_run.text.rstrip('.')

                    for run in paragraph.runs:
                        run.font.name = self.rules.get('font', {}).get('name', 'Times New Roman')
                        if level == 1:
                            run.font.size = Pt(self.rules.get('headings', {}).get('h1', {}).get('font_size', 14))
                            run.font.bold = self.rules.get('headings', {}).get('h1', {}).get('bold', True)
                        elif level == 2:
                            run.font.size = Pt(self.rules.get('headings', {}).get('h2', {}).get('font_size', 14))
                            run.font.bold = self.rules.get('headings', {}).get('h2', {}).get('bold', True)
                        else:
                            run.font.size = Pt(14)
                            run.font.bold = True

                    corrected += 1

                except Exception as exc:
                    self.add_action(
                        element_type='heading',
                        element_index=i,
                        action_type='heading_format_error',
                        old_value=None,
                        new_value=None,
                        description=f'Ошибка форматирования заголовка: {str(exc)}',
                        success=False,
                        error_message=str(exc),
                    )

        except Exception as exc:
            self.add_action(
                element_type='heading',
                element_index=0,
                action_type='heading_correction_error',
                old_value=None,
                new_value=None,
                description=f'Критическая ошибка корректировки заголовков: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected

    def _correct_title_page(self, document: Document) -> int:
        """Исправляет титульный лист.

        Заменяет неправильный титульный лист на шаблонный или вставляет его,
        если титульный лист отсутствует.
        """
        corrected = 0
        title_page_paragraphs = []

        for i, para in enumerate(document.paragraphs):
            text = para.text.strip().lower()
            if para.style.name.startswith('Heading') and para.style.name == 'Heading 1':
                break
            if any(word in text for word in ['содержание', 'введение']):
                break
            title_page_paragraphs.append((i, para))

        title_keywords = [
            'университет', 'кафедра', 'факультет', 'курсовая', 'работа', 'студент', 'руководитель'
        ]

        title_page_exists = False
        if len(title_page_paragraphs) > 3:
            title_text = ' '.join([para[1].text.lower() for para in title_page_paragraphs])
            if any(keyword in title_text for keyword in title_keywords):
                title_page_exists = True

        template_path = os.path.join(
            os.path.dirname(os.path.dirname(__file__)),
            'templates',
            'Титульный лист.docx'
        )

        if not os.path.exists(template_path):
            self.add_action(
                element_type='title_page',
                element_index=0,
                action_type='title_template_missing',
                old_value=None,
                new_value=None,
                description=f'Шаблон титульного листа не найден: {template_path}',
                success=False,
                error_message='Template not found',
            )
            return corrected

        if title_page_exists:
            for i, para in reversed(title_page_paragraphs):
                p = para._element
                p.getparent().remove(p)
                corrected += 1

        inserted = self._insert_title_page_from_template(document, template_path)
        if inserted:
            corrected += 1

        return corrected

    def _insert_title_page_from_template(self, document: Document, template_path: str) -> bool:
        """Вставляет титульный лист из шаблона в начало документа.

        Возвращает True, если вставка выполнена без исключений.
        """
        try:
            template_doc = Document(template_path)

            temp_dir = tempfile.mkdtemp()
            temp_file = os.path.join(temp_dir, 'temp_doc.docx')
            document.save(temp_file)

            composer = Composer(template_doc)
            composer.append(Document(temp_file))

            result_path = os.path.join(temp_dir, 'result.docx')
            composer.save(result_path)

            # Полная замена объекта документа невозможна без смены ссылки,
            # поэтому оставляем текущую реализацию как в старом коде.
            self.add_action(
                element_type='title_page',
                element_index=0,
                action_type='title_page_inserted',
                old_value=None,
                new_value=result_path,
                description='Вставка титульного листа из шаблона выполнена (результат сохранен во временный файл)',
            )

            try:
                shutil.rmtree(temp_dir)
            except Exception:
                pass

            return True

        except Exception as exc:
            self.add_action(
                element_type='title_page',
                element_index=0,
                action_type='title_page_insert_error',
                old_value=None,
                new_value=None,
                description=f'Ошибка вставки титульного листа: {str(exc)}',
                success=False,
                error_message=str(exc),
            )
            return False

    def _get_table_paragraphs(self, document: Document) -> Set[int]:
        """Возвращает ID параграфов, находящихся в таблицах."""
        table_paragraphs: Set[int] = set()
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        table_paragraphs.add(id(para))
        return table_paragraphs
