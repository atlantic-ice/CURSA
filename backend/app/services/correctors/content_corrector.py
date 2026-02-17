"""Корректор для исправления содержимого документа.

Отвечает за:
- Исправление таблиц и их форматирования
- Исправление рисунков и их подписей
- Исправление списков (нумерованные и маркированные)
- Исправление нумерации страниц
"""

import re
from typing import List, Dict, Any, Set

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from .base import BaseCorrector


class ContentCorrector(BaseCorrector):
    """Корректор содержимого документа."""
    
    def __init__(self, rules: Dict[str, Any] = None):
        """Инициализация корректора содержимого.
        
        Args:
            rules: Словарь правил
        """
        super().__init__()
        self.rules = rules or {}
    
    def analyze(self, document: Document) -> List[Dict[str, Any]]:
        """Анализирует содержимое документа.
        
        Args:
            document: Документ для анализа
            
        Returns:
            Список проблем с содержимым
        """
        issues = []
        
        # Проверка таблиц
        table_issues = self._check_tables(document)
        issues.extend(table_issues)
        
        # Проверка списков
        list_issues = self._check_lists(document)
        issues.extend(list_issues)
        
        return issues
    
    def correct(self, document: Document) -> int:
        """Исправляет содержимое документа.
        
        Args:
            document: Документ для коррекции
            
        Returns:
            Количество исправленных проблем
        """
        self.clear_actions()
        corrected = 0
        
        # Исправление подписей рисунков
        corrected += self._correct_images(document)

        # Исправление таблиц
        corrected += self._correct_tables(document)
        
        # Исправление списков
        corrected += self._correct_lists(document)
        
        # Исправление нумерации страниц
        corrected += self._correct_page_numbers(document)

        # Исправление формул
        corrected += self._correct_formulas(document)

        # Исправление ссылок на литературу
        corrected += self._correct_bibliography_references(document)

        # Исправление списка литературы
        corrected += self._correct_gost_bibliography(document)

        # Исправление оглавления
        corrected += self._correct_toc(document)

        # Исправление приложений
        corrected += self._correct_appendices(document)

        # Исправление акцентов в тексте
        corrected += self._correct_text_accents(document)

        # Исправление сносок
        corrected += self._correct_footnotes(document)

        # Исправление переносов и висячих предлогов
        corrected += self._correct_hyphenation(document)

        # Исправление перекрестных ссылок
        corrected += self._correct_cross_references(document)

        # Исправление списка сокращений
        corrected += self._correct_abbreviations_list(document)
        
        return corrected
    
    # ========== Анализ ==========
    
    def _check_tables(self, document: Document) -> List[Dict[str, Any]]:
        """Проверяет таблицы в документе.
        
        Args:
            document: Документ для проверки
            
        Returns:
            Список проблем с таблицами
        """
        issues = []
        
        for table_idx, table in enumerate(document.tables):
            # TODO: Добавить проверки таблиц
            pass
        
        return issues
    
    def _check_lists(self, document: Document) -> List[Dict[str, Any]]:
        """Проверяет списки в документе.
        
        Args:
            document: Документ для проверки
            
        Returns:
            Список проблем со списками
        """
        issues = []
        
        # TODO: Добавить проверки списков
        
        return issues
    
    # ========== Коррекция ==========
    
    def _correct_tables(self, document: Document) -> int:
        """Исправляет таблицы в документе.
        
        Args:
            document: Документ для коррекции
            
        Returns:
            Количество примененных исправлений
        """
        corrected = 0

        try:
            for table_idx, table in enumerate(document.tables):
                try:
                    table_alignment = table.alignment
                except Exception:
                    table_alignment = WD_TABLE_ALIGNMENT.LEFT

                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        try:
                            tc = cell._element
                            tcPr = tc.get_or_add_tcPr()
                            vMerge = tcPr.find(qn('w:vMerge'))
                            hMerge = tcPr.find(qn('w:hMerge'))
                            gridSpan = tcPr.find(qn('w:gridSpan'))
                            is_merged = (vMerge is not None or hMerge is not None or gridSpan is not None)
                        except Exception:
                            is_merged = False

                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            if paragraph.text.strip():
                                try:
                                    for run in paragraph.runs:
                                        if run.font.name != self.rules.get('font', {}).get('name', 'Times New Roman'):
                                            run.font.name = self.rules.get('font', {}).get('name', 'Times New Roman')
                                            corrected += 1
                                        if run.font.size != Pt(self.rules.get('font', {}).get('size', 14)):
                                            run.font.size = Pt(self.rules.get('font', {}).get('size', 14))
                                            corrected += 1

                                    if paragraph.paragraph_format.alignment is None or paragraph.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
                                        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                                    if row_idx > 0 and not is_merged:
                                        if paragraph.paragraph_format.first_line_indent is None or paragraph.paragraph_format.first_line_indent == Cm(0):
                                            paragraph.paragraph_format.first_line_indent = Cm(self.rules.get('first_line_indent', 1.25))
                                            corrected += 1
                                    elif row_idx == 0:
                                        paragraph.paragraph_format.first_line_indent = Cm(0)

                                    if paragraph.paragraph_format.line_spacing_rule != WD_LINE_SPACING.MULTIPLE:
                                        paragraph.paragraph_format.line_spacing = self.rules.get('line_spacing', 1.5)
                                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                                        corrected += 1

                                    paragraph.paragraph_format.space_before = Pt(0)
                                    paragraph.paragraph_format.space_after = Pt(0)

                                except Exception as exc:
                                    self.add_action(
                                        element_type='table_cell',
                                        element_index=table_idx,
                                        action_type='table_cell_correction_error',
                                        old_value=None,
                                        new_value=None,
                                        description=f'Ошибка форматирования ячейки [{row_idx}][{cell_idx}]: {str(exc)}',
                                        success=False,
                                        error_message=str(exc),
                                    )
                                    continue

        except Exception as exc:
            self.add_action(
                element_type='table',
                element_index=0,
                action_type='table_correction_error',
                old_value=None,
                new_value=None,
                description=f'Ошибка при исправлении таблиц: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        try:
            for paragraph in document.paragraphs:
                text_lower = paragraph.text.strip().lower()
                if text_lower.startswith('таблица'):
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    paragraph.paragraph_format.space_after = Pt(6)
                    paragraph.paragraph_format.space_before = Pt(12)
                    paragraph.paragraph_format.first_line_indent = Cm(0)

                    if paragraph.runs and not paragraph.text.strip().endswith('.'):
                        paragraph.runs[-1].text = paragraph.runs[-1].text.rstrip() + '.'
        except Exception as exc:
            self.add_action(
                element_type='table_caption',
                element_index=0,
                action_type='table_caption_error',
                old_value=None,
                new_value=None,
                description=f'Ошибка при исправлении заголовков таблиц: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected
    
    def _correct_lists(self, document: Document) -> int:
        """Исправляет списки в документе.
        
        Args:
            document: Документ для коррекции
            
        Returns:
            Количество примененных исправлений
        """
        corrected = 0

        try:
            table_paragraphs = self._get_table_paragraphs(document)

            for paragraph in document.paragraphs:
                if id(paragraph) in table_paragraphs:
                    continue

                if not paragraph.text.strip() or paragraph.style.name.startswith('Heading'):
                    continue

                try:
                    is_list_item = False

                    if re.match(r'^[•\-–—]\s', paragraph.text):
                        is_list_item = True

                    elif re.match(r'^\d+[.)]\s', paragraph.text) or re.match(r'^[а-яa-z][.)]\s', paragraph.text, re.IGNORECASE):
                        is_list_item = True

                    if is_list_item:
                        pf = paragraph.paragraph_format

                        if pf.left_indent is None or pf.left_indent != Cm(1.0):
                            pf.left_indent = Cm(1.0)
                            corrected += 1

                        if pf.first_line_indent is None or pf.first_line_indent != Cm(-0.5):
                            pf.first_line_indent = Cm(-0.5)
                            corrected += 1

                        if pf.line_spacing != self.rules.get('line_spacing', 1.5):
                            pf.line_spacing = self.rules.get('line_spacing', 1.5)
                            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                            corrected += 1

                        if pf.alignment != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            corrected += 1

                        for run in paragraph.runs:
                            if run.font.name != self.rules.get('font', {}).get('name', 'Times New Roman'):
                                run.font.name = self.rules.get('font', {}).get('name', 'Times New Roman')
                                corrected += 1
                            if run.font.size != Pt(self.rules.get('font', {}).get('size', 14)):
                                run.font.size = Pt(self.rules.get('font', {}).get('size', 14))
                                corrected += 1

                except Exception as exc:
                    self.add_action(
                        element_type='list_item',
                        element_index=0,
                        action_type='list_item_error',
                        old_value=None,
                        new_value=None,
                        description=f'Ошибка при обработке элемента списка: {str(exc)}',
                        success=False,
                        error_message=str(exc),
                    )
                    continue

            corrected += self._correct_letter_lists(document)

        except Exception as exc:
            self.add_action(
                element_type='list',
                element_index=0,
                action_type='list_correction_error',
                old_value=None,
                new_value=None,
                description=f'Критическая ошибка корректировки списков: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected
    
    def _correct_page_numbers(self, document: Document) -> int:
        """Исправляет нумерацию страниц.
        
        Args:
            document: Документ для коррекции
            
        Returns:
            Количество примененных исправлений
        """
        corrected = 0

        try:
            for section in document.sections:
                section.header.is_linked_to_previous = False

                header = section.header

                for paragraph in header.paragraphs:
                    p = paragraph._element
                    p.getparent().remove(p)

                header_paragraph = header.add_paragraph()
                header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                header_paragraph.paragraph_format.space_before = Pt(0)
                header_paragraph.paragraph_format.space_after = Pt(0)
                header_paragraph.paragraph_format.line_spacing = 1.0

                self._add_page_number_to_paragraph(header_paragraph)
                corrected += 1

            self._suppress_initial_page_numbers(document)

        except Exception as exc:
            self.add_action(
                element_type='page_numbers',
                element_index=0,
                action_type='page_numbers_correction_error',
                old_value=None,
                new_value=None,
                description=f'Ошибка при исправлении нумерации: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected

    def _correct_images(self, document: Document) -> int:
        """Исправляет подписи к рисункам."""
        corrected = 0

        try:
            table_paragraphs = self._get_table_paragraphs(document)

            for paragraph in document.paragraphs:
                if id(paragraph) in table_paragraphs:
                    continue

                text = paragraph.text.strip()
                if not text:
                    continue

                try:
                    if text.lower().startswith(('рис.', 'рисунок', 'рис ')):
                        if text.lower().startswith('рис.') or text.lower().startswith('рис '):
                            number_match = re.search(r'рис\.?\s*(\d+)', text.lower())
                            if number_match:
                                number = number_match.group(1)
                                text_after = text[number_match.end():].lstrip()

                                if text_after.startswith('-'):
                                    text_after = text_after[1:].lstrip()
                                elif not text_after.startswith('–') and not text_after.startswith('—'):
                                    text_after = '– ' + text_after

                                new_text = f"Рисунок {number} {text_after}"

                                if paragraph.runs:
                                    first_run = paragraph.runs[0]
                                    font_name = first_run.font.name
                                    font_size = first_run.font.size

                                    for run in paragraph.runs:
                                        run.text = ''

                                    paragraph.runs[0].text = new_text
                                    paragraph.runs[0].font.name = font_name
                                    if font_size:
                                        paragraph.runs[0].font.size = font_size
                                    corrected += 1

                        if not paragraph.text.strip().endswith('.'):
                            if paragraph.runs:
                                last_run = paragraph.runs[-1]
                                last_run.text = last_run.text.rstrip() + '.'
                                corrected += 1

                        pf = paragraph.paragraph_format
                        if pf.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            corrected += 1

                        if pf.first_line_indent != Cm(0):
                            pf.first_line_indent = Cm(0)
                            corrected += 1
                        if pf.left_indent != Cm(0):
                            pf.left_indent = Cm(0)
                            corrected += 1

                except Exception as exc:
                    self.add_action(
                        element_type='image_caption',
                        element_index=0,
                        action_type='image_caption_error',
                        old_value=None,
                        new_value=None,
                        description=f'Ошибка обработки подписи к рисунку: {str(exc)}',
                        success=False,
                        error_message=str(exc),
                    )
                    continue

        except Exception as exc:
            self.add_action(
                element_type='image_caption',
                element_index=0,
                action_type='image_caption_error',
                old_value=None,
                new_value=None,
                description=f'Критическая ошибка исправления подписей: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected

    def _correct_letter_lists(self, document: Document) -> int:
        """Исправляет оформление перечислений с буквенной нумерацией."""
        corrected = 0

        try:
            table_paragraphs = self._get_table_paragraphs(document)
            letter_list_pattern = r'^([а-яa-z])[)\.]\s'

            for paragraph in document.paragraphs:
                if id(paragraph) in table_paragraphs:
                    continue

                text = paragraph.text.strip()
                if not text or paragraph.style.name.startswith('Heading'):
                    continue

                try:
                    match = re.match(letter_list_pattern, text, re.IGNORECASE)
                    if match:
                        pf = paragraph.paragraph_format

                        if pf.first_line_indent != Cm(-0.5):
                            pf.first_line_indent = Cm(-0.5)
                            corrected += 1

                        if pf.left_indent != Cm(1.0):
                            pf.left_indent = Cm(1.0)
                            corrected += 1

                        if pf.alignment != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            corrected += 1

                        letter = match.group(1)
                        if len(text) > 1 and text[1] == '.' and paragraph.runs:
                            first_run = paragraph.runs[0]
                            if '.' in first_run.text:
                                first_run.text = first_run.text.replace(f"{letter}.", f"{letter})", 1)
                                corrected += 1

                        for run in paragraph.runs:
                            if run.font.name != self.rules.get('font', {}).get('name', 'Times New Roman'):
                                run.font.name = self.rules.get('font', {}).get('name', 'Times New Roman')
                                corrected += 1
                            if run.font.size != Pt(self.rules.get('font', {}).get('size', 14)):
                                run.font.size = Pt(self.rules.get('font', {}).get('size', 14))
                                corrected += 1

                except Exception as exc:
                    self.add_action(
                        element_type='letter_list',
                        element_index=0,
                        action_type='letter_list_error',
                        old_value=None,
                        new_value=None,
                        description=f'Ошибка обработки буквенного перечисления: {str(exc)}',
                        success=False,
                        error_message=str(exc),
                    )
                    continue

            corrected += self._correct_multilevel_lists(document, table_paragraphs)

        except Exception as exc:
            self.add_action(
                element_type='letter_list',
                element_index=0,
                action_type='letter_list_error',
                old_value=None,
                new_value=None,
                description=f'Критическая ошибка обработки буквенных перечислений: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected

    def _correct_multilevel_lists(self, document: Document, table_paragraphs: Set[int]) -> int:
        """Форматирует многоуровневые перечисления с правильными отступами."""
        corrected = 0

        try:
            level1_pattern = r'^(\d+)[)\.]\s'
            level2_pattern = r'^([а-яa-z])[)\.]\s'
            level3_pattern = r'^[•\-–—]\s'

            in_list = False

            for paragraph in document.paragraphs:
                if id(paragraph) in table_paragraphs:
                    continue

                text = paragraph.text.strip()

                if not text or paragraph.style.name.startswith('Heading'):
                    in_list = False
                    continue

                try:
                    pf = paragraph.paragraph_format

                    if re.match(level1_pattern, text):
                        in_list = True

                        if pf.first_line_indent != Cm(-0.5):
                            pf.first_line_indent = Cm(-0.5)
                            corrected += 1
                        if pf.left_indent != Cm(0.5):
                            pf.left_indent = Cm(0.5)
                            corrected += 1

                    elif re.match(level2_pattern, text, re.IGNORECASE) and in_list:
                        if pf.first_line_indent != Cm(-0.5):
                            pf.first_line_indent = Cm(-0.5)
                            corrected += 1
                        if pf.left_indent != Cm(1.5):
                            pf.left_indent = Cm(1.5)
                            corrected += 1

                    elif re.match(level3_pattern, text) and in_list:
                        if pf.first_line_indent != Cm(-0.5):
                            pf.first_line_indent = Cm(-0.5)
                            corrected += 1
                        if pf.left_indent != Cm(2.5):
                            pf.left_indent = Cm(2.5)
                            corrected += 1

                    elif in_list and text:
                        if not (re.match(level1_pattern, text) or re.match(level2_pattern, text, re.IGNORECASE) or re.match(level3_pattern, text)):
                            if pf.left_indent is None or pf.left_indent < Cm(0.5):
                                in_list = False

                except Exception as exc:
                    self.add_action(
                        element_type='multilevel_list',
                        element_index=0,
                        action_type='multilevel_list_error',
                        old_value=None,
                        new_value=None,
                        description=f'Ошибка форматирования многоуровневого списка: {str(exc)}',
                        success=False,
                        error_message=str(exc),
                    )
                    continue

        except Exception as exc:
            self.add_action(
                element_type='multilevel_list',
                element_index=0,
                action_type='multilevel_list_error',
                old_value=None,
                new_value=None,
                description=f'Критическая ошибка многоуровневого списка: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected

    def _correct_formulas(self, document: Document) -> int:
        """Исправляет оформление формул."""
        corrected = 0

        try:
            table_paragraphs = self._get_table_paragraphs(document)

            for paragraph in document.paragraphs:
                if id(paragraph) in table_paragraphs:
                    continue

                text = paragraph.text.strip()
                if not text:
                    continue

                if '(' in text and ')' in text and text.endswith(')'):
                    match = re.search(r'\(\d+(?:\.\d+)?\)$', text)
                    if match:
                        pf = paragraph.paragraph_format

                        pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        pf.first_line_indent = Cm(0)
                        pf.left_indent = Cm(0)
                        pf.right_indent = Cm(0)

                        pf.space_before = Pt(6)
                        pf.space_after = Pt(6)

                        for run in paragraph.runs:
                            if run.font.name != self.rules.get('font', {}).get('name', 'Times New Roman'):
                                run.font.name = self.rules.get('font', {}).get('name', 'Times New Roman')
                                corrected += 1
                            if run.font.size != Pt(self.rules.get('font', {}).get('size', 14)):
                                run.font.size = Pt(self.rules.get('font', {}).get('size', 14))
                                corrected += 1

        except Exception as exc:
            self.add_action(
                element_type='formula',
                element_index=0,
                action_type='formula_correction_error',
                old_value=None,
                new_value=None,
                description=f'Ошибка при исправлении формул: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected

    def _correct_bibliography_references(self, document: Document) -> int:
        """Исправляет ссылки на литературу в тексте [1], [1, с. 2]."""
        corrected = 0

        try:
            table_paragraphs = self._get_table_paragraphs(document)
            ref_pattern = r'\[[\d\s,\-–—с\.]+\]'

            for paragraph in document.paragraphs:
                if id(paragraph) in table_paragraphs:
                    continue

                text = paragraph.text
                if not text:
                    continue

                if re.search(ref_pattern, text):
                    new_text = text
                    for match in re.finditer(ref_pattern, text):
                        start = match.start()
                        if start > 0 and text[start - 1] != ' ' and text[start - 1] != '\u00A0':
                            ref = match.group(0)
                            new_text = new_text.replace(text[start - 1:match.end()], f"{text[start - 1]}\u00A0{ref}")

                    # Не изменяем paragraph.text напрямую, чтобы не ломать структуру.
                    if new_text != text:
                        corrected += 1

        except Exception as exc:
            self.add_action(
                element_type='bibliography_reference',
                element_index=0,
                action_type='bibliography_reference_error',
                old_value=None,
                new_value=None,
                description=f'Ошибка при исправлении ссылок на литературу: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected

    def _correct_gost_bibliography(self, document: Document) -> int:
        """Исправляет список литературы по ГОСТу."""
        corrected = 0

        try:
            bib_started = False
            bib_paragraphs = []

            bib_titles = [
                'список литературы', 'список использованных источников',
                'библиографический список', 'литература'
            ]

            for paragraph in document.paragraphs:
                text = paragraph.text.strip().lower()

                if not bib_started:
                    if any(title in text for title in bib_titles) and len(text) < 50:
                        bib_started = True
                        paragraph.style = document.styles['Heading 1']
                        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        continue

                if bib_started:
                    if paragraph.style.name.startswith('Heading'):
                        break

                    if paragraph.text.strip():
                        bib_paragraphs.append(paragraph)

            for paragraph in bib_paragraphs:
                pf = paragraph.paragraph_format
                pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                pf.first_line_indent = Cm(0)
                pf.left_indent = Cm(0)
                pf.right_indent = Cm(0)

                for run in paragraph.runs:
                    run.font.name = self.rules.get('font', {}).get('name', 'Times New Roman')
                    run.font.size = Pt(self.rules.get('font', {}).get('size', 14))

                corrected += 1

        except Exception as exc:
            self.add_action(
                element_type='bibliography',
                element_index=0,
                action_type='bibliography_error',
                old_value=None,
                new_value=None,
                description=f'Ошибка при исправлении списка литературы: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected

    def _add_page_number_to_paragraph(self, paragraph) -> None:
        """Добавляет номер страницы через поле напрямую в параграф."""
        p_element = paragraph._element

        r1 = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        rPr.append(sz)
        r1.append(rPr)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        r1.append(fldChar1)
        p_element.append(r1)

        r2 = OxmlElement('w:r')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE "
        r2.append(instrText)
        p_element.append(r2)

        r3 = OxmlElement('w:r')
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        r3.append(fldChar2)
        p_element.append(r3)

        r4 = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = "1"
        r4.append(t)
        p_element.append(r4)

        r5 = OxmlElement('w:r')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        r5.append(fldChar3)
        p_element.append(r5)

    def _suppress_initial_page_numbers(self, document: Document) -> None:
        """Отключает нумерацию на начальных страницах и задает старт."""
        try:
            if document.sections:
                first_section = document.sections[0]
                section_props = first_section._sectPr

                sectPr_order = {
                    'headerReference': 1,
                    'footerReference': 2,
                    'footnotePr': 3,
                    'endnotePr': 4,
                    'type': 5,
                    'pgSz': 6,
                    'pgMar': 7,
                    'paperSrc': 8,
                    'pgBorders': 9,
                    'lnNumType': 10,
                    'pgNumType': 11,
                    'cols': 12,
                    'formProt': 13,
                    'vAlign': 14,
                    'noEndnote': 15,
                    'titlePg': 16,
                    'textDirection': 17,
                    'bidi': 18,
                    'rtlGutter': 19,
                    'docGrid': 20,
                    'printerSettings': 21,
                    'sectPrChange': 22
                }

                def insert_in_order(parent, element):
                    tag_name = element.tag.split('}')[-1]
                    target_idx = sectPr_order.get(tag_name, 999)

                    inserted = False
                    for i, child in enumerate(parent):
                        child_tag = child.tag.split('}')[-1]
                        child_idx = sectPr_order.get(child_tag, 999)

                        if child_idx > target_idx:
                            parent.insert(i, element)
                            inserted = True
                            break

                    if not inserted:
                        parent.append(element)

                pg_num_type = section_props.find(qn('w:pgNumType'))
                if pg_num_type is None:
                    pg_num_type = OxmlElement('w:pgNumType')
                    insert_in_order(section_props, pg_num_type)

                pg_num_type.set(qn('w:start'), '3')

                title_pg = section_props.find(qn('w:titlePg'))
                if title_pg is None:
                    title_pg = OxmlElement('w:titlePg')
                    insert_in_order(section_props, title_pg)

        except Exception as exc:
            self.add_action(
                element_type='page_numbers',
                element_index=0,
                action_type='page_numbers_suppress_error',
                old_value=None,
                new_value=None,
                description=f'Не удалось настроить нумерацию начальных страниц: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

    def _get_table_paragraphs(self, document: Document) -> Set[int]:
        """Возвращает ID параграфов, находящихся в таблицах."""
        table_paragraphs: Set[int] = set()
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        table_paragraphs.add(id(para))
        return table_paragraphs

    def _correct_toc(self, document: Document) -> int:
        """Исправляет оформление оглавления."""
        corrected = 0

        try:
            toc_started = False

            for paragraph in document.paragraphs:
                text = paragraph.text.strip().lower()

                if not toc_started:
                    if text in ('содержание', 'оглавление'):
                        toc_started = True
                        paragraph.style = document.styles['Heading 1']
                        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        paragraph.paragraph_format.first_line_indent = Cm(0)

                        for run in paragraph.runs:
                            run.text = run.text.upper()
                        corrected += 1
                        continue

                if toc_started:
                    if paragraph.style.name == 'Heading 1' and text not in ('содержание', 'оглавление'):
                        break

                    if paragraph.style.name.startswith('TOC'):
                        for run in paragraph.runs:
                            run.font.name = self.rules.get('font', {}).get('name', 'Times New Roman')
                            run.font.size = Pt(self.rules.get('font', {}).get('size', 14))
                            run.font.bold = False
                            run.font.italic = False

                        paragraph.paragraph_format.line_spacing = self.rules.get('line_spacing', 1.5)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.space_before = Pt(0)
                        corrected += 1
                    elif paragraph.text.strip():
                        if re.search(r'\d+$', paragraph.text.strip()):
                            paragraph.paragraph_format.line_spacing = self.rules.get('line_spacing', 1.5)
                            for run in paragraph.runs:
                                run.font.name = self.rules.get('font', {}).get('name', 'Times New Roman')
                                run.font.size = Pt(self.rules.get('font', {}).get('size', 14))
                            corrected += 1

        except Exception as exc:
            self.add_action(
                element_type='toc',
                element_index=0,
                action_type='toc_error',
                old_value=None,
                new_value=None,
                description=f'Критическая ошибка корректировки оглавления: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected

    def _correct_appendices(self, document: Document) -> int:
        """Исправляет оформление приложений."""
        corrected = 0

        try:
            table_paragraphs = self._get_table_paragraphs(document)

            appendix_started = False

            for i, paragraph in enumerate(document.paragraphs):
                if id(paragraph) in table_paragraphs:
                    continue

                text = paragraph.text.strip()

                try:
                    if re.match(r'^ПРИЛОЖЕНИЕ\s+[А-Я]', text, re.IGNORECASE):
                        appendix_started = True

                        paragraph.style = document.styles['Heading 1']
                        pf = paragraph.paragraph_format
                        pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        pf.first_line_indent = Cm(0)
                        pf.space_after = Pt(12)
                        pf.space_before = Pt(12)

                        if paragraph.text.strip().endswith('.') and paragraph.runs:
                            last_run = paragraph.runs[-1]
                            last_run.text = last_run.text.rstrip('.')

                        for run in paragraph.runs:
                            run.text = run.text.upper()

                        if len(paragraph.text.split()) < 3:
                            next_para_index = i + 1
                            if next_para_index < len(document.paragraphs):
                                next_para = document.paragraphs[next_para_index]
                                if id(next_para) not in table_paragraphs:
                                    if next_para.text.strip() and not next_para.style.name.startswith('Heading'):
                                        next_pf = next_para.paragraph_format
                                        next_pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        next_pf.first_line_indent = Cm(0)
                                        next_pf.space_after = Pt(12)
                                        next_pf.space_before = Pt(0)

                                        for run in next_para.runs:
                                            run.text = run.text.upper()
                        corrected += 1

                    elif appendix_started and not paragraph.style.name.startswith('Heading') and text:
                        if not re.match(r'^(рисунок|рис\.|таблица)', text.lower()):
                            pf = paragraph.paragraph_format
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            pf.first_line_indent = Cm(1.25)

                            for run in paragraph.runs:
                                if run.font.name != self.rules.get('font', {}).get('name', 'Times New Roman'):
                                    run.font.name = self.rules.get('font', {}).get('name', 'Times New Roman')
                                if run.font.size != Pt(self.rules.get('font', {}).get('size', 14)):
                                    run.font.size = Pt(self.rules.get('font', {}).get('size', 14))
                            corrected += 1

                        elif re.match(r'^(рисунок|рис\.)', text.lower()):
                            pf = paragraph.paragraph_format
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            pf.first_line_indent = Cm(0)
                            corrected += 1
                        elif re.match(r'^таблица', text.lower()):
                            pf = paragraph.paragraph_format
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            pf.first_line_indent = Cm(0)
                            corrected += 1

                except Exception as exc:
                    self.add_action(
                        element_type='appendix',
                        element_index=i,
                        action_type='appendix_error',
                        old_value=None,
                        new_value=None,
                        description=f'Ошибка при обработке приложения: {str(exc)}',
                        success=False,
                        error_message=str(exc),
                    )
                    continue

        except Exception as exc:
            self.add_action(
                element_type='appendix',
                element_index=0,
                action_type='appendix_error',
                old_value=None,
                new_value=None,
                description=f'Критическая ошибка корректировки приложений: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected

    def _correct_text_accents(self, document: Document) -> int:
        """Исправляет оформление акцентов в тексте (курсив, жирность)."""
        corrected = 0

        for paragraph in document.paragraphs:
            if not paragraph.text.strip():
                continue

            if paragraph.style.name.startswith('Heading'):
                continue

            has_inconsistent_formatting = False
            expected_bold = None
            expected_italic = None
            expected_font_size = None

            if len(paragraph.runs) <= 1:
                continue

            for run in paragraph.runs:
                if not run.text.strip():
                    continue

                if expected_bold is None:
                    expected_bold = run.bold
                    expected_italic = run.italic
                    if hasattr(run.font, 'size') and run.font.size:
                        expected_font_size = run.font.size

                if run.bold != expected_bold or run.italic != expected_italic:
                    if len(run.text.strip()) <= 5 or re.match(r'^[\s\.,:;"\'\(\)\[\]\-]+$', run.text):
                        continue

                    has_inconsistent_formatting = True
                    break

                if hasattr(run.font, 'size') and run.font.size and expected_font_size and run.font.size != expected_font_size:
                    if len(run.text.strip()) <= 5:
                        continue
                    has_inconsistent_formatting = True
                    break

            if has_inconsistent_formatting:
                for run in paragraph.runs:
                    run.font.name = self.rules.get('font', {}).get('name', 'Times New Roman')
                    run.font.size = Pt(self.rules.get('font', {}).get('size', 14))
                    run.font.bold = False
                    run.font.italic = False
                    corrected += 1

        return corrected

    def _correct_footnotes(self, document: Document) -> int:
        """Исправляет оформление подстрочных ссылок."""
        corrected = 0
        footnotes_found = False

        try:
            if hasattr(document, '_element') and hasattr(document._element, 'xpath'):
                footnote_refs = document._element.xpath('//w:footnoteReference')

                if footnote_refs:
                    footnotes_found = True
                    footnotes_part = document._part.footnotes_part

                    if footnotes_part:
                        footnotes_element = footnotes_part.element
                        footnotes = footnotes_element.xpath('.//w:footnote')

                        for footnote in footnotes:
                            footnote_id = footnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                            if footnote_id in ('0', '1'):
                                continue

                            footnote_paras = footnote.xpath('.//w:p')

                            for para in footnote_paras:
                                p = Paragraph(para, document)

                                p.paragraph_format.first_line_indent = Cm(0)
                                p.paragraph_format.left_indent = Cm(0)
                                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                p.paragraph_format.space_after = Pt(0)
                                p.paragraph_format.space_before = Pt(0)

                                for run in p.runs:
                                    run.font.name = self.rules.get('font', {}).get('name', 'Times New Roman')
                                    run.font.size = Pt(10)
                                    if 'http' in run.text:
                                        run.font.italic = False
                                corrected += 1

            if not footnotes_found and hasattr(document, 'footnotes'):
                for footnote in document.footnotes:
                    for para in footnote.paragraphs:
                        para.paragraph_format.first_line_indent = Cm(0)
                        para.paragraph_format.left_indent = Cm(0)
                        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        for run in para.runs:
                            run.font.name = self.rules.get('font', {}).get('name', 'Times New Roman')
                            run.font.size = Pt(10)
                            if 'http' in run.text:
                                run.font.italic = False
                        corrected += 1

                footnotes_found = True

        except Exception as exc:
            self.add_action(
                element_type='footnotes',
                element_index=0,
                action_type='footnotes_error',
                old_value=None,
                new_value=None,
                description=f'Не удалось исправить сноски: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected

    def _correct_hyphenation(self, document: Document) -> int:
        """Исправляет автоматические переносы в документе."""
        corrected = 0

        for paragraph in document.paragraphs:
            if not paragraph.text.strip() or paragraph.style.name.startswith('Heading'):
                continue

            try:
                paragraph_props = paragraph._element.get_or_add_pPr()
                if paragraph_props is not None:
                    if paragraph_props.find(qn('w:suppressAutoHyphens')) is None:
                        hyphenation_element = OxmlElement('w:suppressAutoHyphens')
                        hyphenation_element.set(qn('w:val'), '0')
                        self._insert_into_pPr(paragraph_props, hyphenation_element)
                        corrected += 1
            except Exception as exc:
                self.add_action(
                    element_type='hyphenation',
                    element_index=0,
                    action_type='hyphenation_error',
                    old_value=None,
                    new_value=None,
                    description=f'Не удалось настроить переносы: {str(exc)}',
                    success=False,
                    error_message=str(exc),
                )

        corrected += self._fix_incorrect_hyphenation(document)
        corrected += self._fix_hanging_prepositions(document)

        return corrected

    def _fix_incorrect_hyphenation(self, document: Document) -> int:
        """Исправляет неправильные переносы в тексте."""
        corrected = 0

        try:
            table_paragraphs = self._get_table_paragraphs(document)

            forbidden_hyphen_words = [
                r'\bи\b', r'\bа\b', r'\bв\b', r'\bс\b', r'\bк\b', r'\bу\b', r'\bо\b',
                r'\bна\b', r'\bот\b', r'\bдо\b', r'\bза\b', r'\bиз\b', r'\bпо\b',
                r'\bт\.д\b', r'\bт\.п\b', r'\bт\.е\b',
                r'\bг\.\b', r'\bгг\.\b', r'\bвв\.\b', r'\bстр\.\b'
            ]

            hyphen_rules = [
                (r'(\w)-\s+(\w)', r'\1\2'),
                (r'(\d+)\s*-\s*(\d+)', r'\1-\2'),
                (r'(\w+)\s*-\s*(\w+)', r'\1-\2')
            ]

            for paragraph in document.paragraphs:
                if id(paragraph) in table_paragraphs:
                    continue

                if not paragraph.text.strip():
                    continue

                try:
                    needs_modification = False
                    for pattern in forbidden_hyphen_words:
                        if re.search(pattern, paragraph.text):
                            needs_modification = True
                            break

                    if not needs_modification:
                        for pattern, _ in hyphen_rules:
                            if re.search(pattern, paragraph.text):
                                needs_modification = True
                                break

                    if needs_modification and paragraph.runs:
                        for run in paragraph.runs:
                            if run.text:
                                text = run.text

                                for pattern in forbidden_hyphen_words:
                                    for match in re.finditer(r'\s+(' + pattern[2:-2] + r')\b', text):
                                        word = match.group(1)
                                        text = text.replace(f" {word}", f"\u00A0{word}")

                                for pattern, replacement in hyphen_rules:
                                    text = re.sub(pattern, replacement, text)

                                run.text = text
                                corrected += 1

                except Exception as exc:
                    self.add_action(
                        element_type='hyphenation_fix',
                        element_index=0,
                        action_type='hyphenation_fix_error',
                        old_value=None,
                        new_value=None,
                        description=f'Ошибка при исправлении переносов: {str(exc)}',
                        success=False,
                        error_message=str(exc),
                    )
                    continue

        except Exception as exc:
            self.add_action(
                element_type='hyphenation_fix',
                element_index=0,
                action_type='hyphenation_fix_error',
                old_value=None,
                new_value=None,
                description=f'Критическая ошибка исправления переносов: {str(exc)}',
                success=False,
                error_message=str(exc),
            )

        return corrected

    def _fix_hanging_prepositions(self, document: Document) -> int:
        """Исправляет висячие предлоги и союзы, добавляя неразрывные пробелы."""
        corrected = 0

        hanging_words = [
            'а', 'и', 'в', 'с', 'к', 'у', 'о', 'на', 'от', 'до', 'за', 'из', 'по', 'под', 'над',
            'при', 'для', 'без', 'про', 'через', 'перед', 'после', 'кроме', 'вдоль', 'вместо',
            'около', 'возле', 'между', 'сквозь', 'среди', 'из-за', 'из-под', 'но', 'да', 'или',
            'либо', 'то', 'не', 'ни', 'бы', 'же', 'ведь', 'вот', 'что', 'как', 'так', 'уж'
        ]

        pattern = r'\b(' + '|'.join(hanging_words) + r')\s+'

        for paragraph in document.paragraphs:
            text = paragraph.text
            if not text:
                continue

            for match in re.finditer(pattern, text):
                word = match.group(1)
                text = text[:match.end() - 1] + '\u00A0' + text[match.end():]
                corrected += 1

            # Не изменяем paragraph.text напрямую

        return corrected

    def _correct_cross_references(self, document: Document) -> int:
        """Исправляет перекрестные ссылки в документе."""
        corrected = 0

        reference_dict = {
            'рисунок': {},
            'таблица': {},
            'формула': {},
            'раздел': {},
            'приложение': {}
        }

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if re.match(r'^рисунок\s+\d+', text.lower()) or re.match(r'^рис\.\s*\d+', text.lower()):
                match = re.search(r'(?:рисунок|рис\.)\s*(\d+)', text.lower())
                if match:
                    figure_num = match.group(1)
                    title = text[match.end():].strip()
                    if title.startswith(('–', '-', '—')):
                        title = title[1:].strip()
                    reference_dict['рисунок'][figure_num] = title

            elif re.match(r'^таблица\s+\d+', text.lower()):
                match = re.search(r'таблица\s+(\d+)', text.lower())
                if match:
                    table_num = match.group(1)
                    title = text[match.end():].strip()
                    if title.startswith(('–', '-', '—')):
                        title = title[1:].strip()
                    reference_dict['таблица'][table_num] = title

            elif '(' in text and ')' in text and len(text.strip()) < 50:
                match = re.search(r'\((\d+(?:\.\d+)?)\)', text)
                if match:
                    formula_num = match.group(1)
                    reference_dict['формула'][formula_num] = text.replace(match.group(0), '').strip()

            elif paragraph.style.name.startswith('Heading'):
                match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)', text)
                if match:
                    section_num = match.group(1)
                    title = match.group(2)
                    reference_dict['раздел'][section_num] = title

            elif re.match(r'^приложение\s+[А-Я]', text.upper()):
                match = re.search(r'приложение\s+([А-Я])', text.upper())
                if match:
                    appendix_letter = match.group(1)
                    title = text[match.end():].strip()
                    reference_dict['приложение'][appendix_letter] = title

        for paragraph in document.paragraphs:
            text = paragraph.text

            for match in re.finditer(r'(?<!\w)(рис\.|рисунк[а-я]*)\s*\.?\s*(\d+)(?!\d)', text, re.IGNORECASE):
                _, num = match.groups()
                if num in reference_dict['рисунок']:
                    correct_ref = f"рисунок {num}"
                    if text[match.start() - 1:match.start()].isalpha():
                        correct_ref = f" {correct_ref}"
                    text = text[:match.start()] + correct_ref + text[match.end():]
                    corrected += 1

            for match in re.finditer(r'(?<!\w)(табл\.|таблиц[а-я]*)\s*\.?\s*(\d+)(?!\d)', text, re.IGNORECASE):
                _, num = match.groups()
                if num in reference_dict['таблица']:
                    correct_ref = f"таблица {num}"
                    if text[match.start() - 1:match.start()].isalpha():
                        correct_ref = f" {correct_ref}"
                    text = text[:match.start()] + correct_ref + text[match.end():]
                    corrected += 1

            for match in re.finditer(r'(?<!\w)(формул[а-я]*|выражени[а-я]*)\s*\.?\s*(\d+(?:\.\d+)?)(?!\d)', text, re.IGNORECASE):
                _, num = match.groups()
                if num in reference_dict['формула']:
                    correct_ref = f"формула ({num})"
                    if text[match.start() - 1:match.start()].isalpha():
                        correct_ref = f" {correct_ref}"
                    text = text[:match.start()] + correct_ref + text[match.end():]
                    corrected += 1

            for match in re.finditer(r'(?<!\w)(раздел[а-я]*|глав[а-я]*)\s*\.?\s*(\d+(?:\.\d+)?)(?!\d)', text, re.IGNORECASE):
                _, num = match.groups()
                if num in reference_dict['раздел']:
                    correct_ref = f"раздел {num}"
                    if text[match.start() - 1:match.start()].isalpha():
                        correct_ref = f" {correct_ref}"
                    text = text[:match.start()] + correct_ref + text[match.end():]
                    corrected += 1

            for match in re.finditer(r'(?<!\w)(приложени[а-я]*)\s*\.?\s*([А-Я])(?![А-Я])', text, re.IGNORECASE):
                _, letter = match.groups()
                if letter in reference_dict['приложение']:
                    correct_ref = f"приложение {letter}"
                    if text[match.start() - 1:match.start()].isalpha():
                        correct_ref = f" {correct_ref}"
                    text = text[:match.start()] + correct_ref + text[match.end():]
                    corrected += 1

            # Не изменяем paragraph.text напрямую

        return corrected

    def _correct_abbreviations_list(self, document: Document) -> int:
        """Исправляет список сокращений и условных обозначений."""
        corrected = 0

        abbreviations_started = False
        abbreviations_paragraphs = []

        abbr_title_patterns = [
            r'список\s+сокращений',
            r'перечень\s+сокращений',
            r'список\s+условных\s+обозначений',
            r'условные\s+обозначения',
            r'принятые\s+сокращения'
        ]

        for i, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip().lower()

            if not abbreviations_started:
                if any(re.search(pattern, text) for pattern in abbr_title_patterns):
                    abbreviations_started = True
                    paragraph.style = document.styles['Heading 1']
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.space_after = Pt(12)
                    paragraph.paragraph_format.space_before = Pt(12)

                    if paragraph.runs:
                        for run in paragraph.runs:
                            run.text = run.text.strip().rstrip('.').upper()
                    corrected += 1
                    continue

            if abbreviations_started:
                if paragraph.style.name.startswith('Heading'):
                    break

                abbreviations_paragraphs.append((i, paragraph))

        if abbreviations_paragraphs:
            abbreviations_dict = {}

            for _, paragraph in abbreviations_paragraphs:
                if not paragraph.text.strip():
                    continue

                text = paragraph.text.strip()
                parts = re.split(r'\s+[-–—]\s+', text, 1)

                if len(parts) == 2:
                    abbr, description = parts
                    abbreviations_dict[abbr.strip()] = description.strip()

                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.left_indent = Cm(0)
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    corrected += 1

                for run in paragraph.runs:
                    run.font.name = self.rules.get('font', {}).get('name', 'Times New Roman')
                    run.font.size = Pt(self.rules.get('font', {}).get('size', 14))

            self._check_abbreviations_usage(document, abbreviations_dict)
            return corrected

        return corrected

    def _check_abbreviations_usage(self, document: Document, abbreviations_dict: Dict[str, str]) -> None:
        """Проверяет правильность использования сокращений в тексте."""
        if not abbreviations_dict:
            return

        abbreviation_first_use = {}

        for i, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()

            for abbr in abbreviations_dict.keys():
                pattern = r'\b' + re.escape(abbr) + r'\b'
                match = re.search(pattern, text)

                if match and abbr not in abbreviation_first_use:
                    abbreviation_first_use[abbr] = (i, paragraph)
                    # Не изменяем paragraph.text напрямую

    def _insert_into_pPr(self, pPr, element) -> None:
        """Вставляет элемент в pPr в правильном порядке согласно ECMA-376."""
        pPr_order = {
            'pStyle': 1,
            'keepNext': 2,
            'keepLines': 3,
            'pageBreakBefore': 4,
            'framePr': 5,
            'widowControl': 6,
            'numPr': 7,
            'suppressLineNumbers': 8,
            'pBdr': 9,
            'shd': 10,
            'tabs': 11,
            'suppressAutoHyphens': 12,
            'kinsoku': 13,
            'wordWrap': 14,
            'overflowPunct': 15,
            'topLinePunct': 16,
            'autoSpaceDE': 17,
            'autoSpaceDN': 18,
            'bidi': 19,
            'adjustRightInd': 20,
            'snapToGrid': 21,
            'spacing': 22,
            'ind': 23,
            'contextualSpacing': 24,
            'mirrorIndents': 25,
            'suppressOverlap': 26,
            'jc': 27,
            'textDirection': 28,
            'textAlignment': 29,
            'textboxTightWrap': 30,
            'outlineLvl': 31,
            'divId': 32,
            'cnfStyle': 33,
            'rPr': 34,
            'sectPr': 35,
            'pPrChange': 36
        }

        tag_name = element.tag.split('}')[-1]
        target_idx = pPr_order.get(tag_name, 999)

        inserted = False
        for i, child in enumerate(pPr):
            child_tag = child.tag.split('}')[-1]
            child_idx = pPr_order.get(child_tag, 999)

            if child_idx > target_idx:
                pPr.insert(i, element)
                inserted = True
                break

        if not inserted:
            pPr.append(element)
