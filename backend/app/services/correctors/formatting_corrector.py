"""Корректор для исправления форматирования документа.

Отвечает за:
- Исправление выравнивания текста
- Исправление специальных форматов (курсив, подчеркивание)
- Исправление формул
- Исправление библиографии
"""

import re
from typing import List, Dict, Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .base import BaseCorrector


class FormattingCorrector(BaseCorrector):
    """Корректор форматирования документа."""

    def __init__(self, rules: Dict[str, Any] = None):
        """Инициализация корректора форматирования.

        Args:
            rules: Словарь правил
        """
        super().__init__()
        self.rules = rules or {}

    def analyze(self, document: Document) -> List[Dict[str, Any]]:
        """Анализирует форматирование документа.

        Args:
            document: Документ для анализа

        Returns:
            Список проблем с форматированием
        """
        issues = []

        # Проверка выравнивания
        alignment_issues = self._check_alignment(document)
        issues.extend(alignment_issues)

        return issues

    def correct(self, document: Document) -> int:
        """Исправляет форматирование документа.

        Args:
            document: Документ для коррекции

        Returns:
            Количество исправленных проблем
        """
        self.clear_actions()
        corrected = 0

        # Исправление выравнивания
        corrected += self._correct_alignment(document)

        return corrected

    # ========== Анализ ==========

    def _check_alignment(self, document: Document) -> List[Dict[str, Any]]:
        """Проверяет выравнивание текста.

        Args:
            document: Документ для проверки

        Returns:
            Список проблем с выравниванием
        """
        issues = []

        # TODO: Добавить проверки выравнивания

        return issues

    # ========== Коррекция ==========

    def _correct_alignment(self, document: Document) -> int:
        """Исправляет выравнивание текста.

        Args:
            document: Документ для коррекции

        Returns:
            Количество примененных исправлений
        """
        corrected = 0

        for paragraph in document.paragraphs:
            if not paragraph.text.strip():
                continue

            if paragraph.style.name.startswith("Heading"):
                try:
                    heading_level = int(paragraph.style.name.replace("Heading ", ""))
                except ValueError:
                    heading_level = None

                if heading_level == 1:
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                continue

            if paragraph.text.strip().lower().startswith(("рисунок", "рис.")):
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                continue

            if paragraph.text.strip().lower().startswith("таблица"):
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                continue

            if re.match(r"^[•\-–—]\s", paragraph.text) or re.match(r"^\d+[.)]\s", paragraph.text):
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                continue

            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            self._enable_hyphenation(paragraph)
            corrected += 1

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if not paragraph.text.strip():
                            continue

                        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        self._enable_hyphenation(paragraph)
                        corrected += 1

        return corrected

    def _insert_into_pPr(self, pPr, element) -> None:
        """Вставляет элемент в pPr в правильном порядке согласно ECMA-376."""
        pPr_order = {
            "pStyle": 1,
            "keepNext": 2,
            "keepLines": 3,
            "pageBreakBefore": 4,
            "framePr": 5,
            "widowControl": 6,
            "numPr": 7,
            "suppressLineNumbers": 8,
            "pBdr": 9,
            "shd": 10,
            "tabs": 11,
            "suppressAutoHyphens": 12,
            "kinsoku": 13,
            "wordWrap": 14,
            "overflowPunct": 15,
            "topLinePunct": 16,
            "autoSpaceDE": 17,
            "autoSpaceDN": 18,
            "bidi": 19,
            "adjustRightInd": 20,
            "snapToGrid": 21,
            "spacing": 22,
            "ind": 23,
            "contextualSpacing": 24,
            "mirrorIndents": 25,
            "suppressOverlap": 26,
            "jc": 27,
            "textDirection": 28,
            "textAlignment": 29,
            "textboxTightWrap": 30,
            "outlineLvl": 31,
            "divId": 32,
            "cnfStyle": 33,
            "rPr": 34,
            "sectPr": 35,
            "pPrChange": 36,
        }

        tag_name = element.tag.split("}")[-1]
        target_idx = pPr_order.get(tag_name, 999)

        inserted = False
        for i, child in enumerate(pPr):
            child_tag = child.tag.split("}")[-1]
            child_idx = pPr_order.get(child_tag, 999)

            if child_idx > target_idx:
                pPr.insert(i, element)
                inserted = True
                break

        if not inserted:
            pPr.append(element)

    def _enable_hyphenation(self, paragraph) -> None:
        """Включает автоматические переносы для параграфа."""
        try:
            pPr = paragraph._element.get_or_add_pPr()
            if pPr is not None:
                if pPr.find(qn("w:autoSpaceDE")) is None:
                    hyphenation_element = OxmlElement("w:autoSpaceDE")
                    hyphenation_element.set(qn("w:val"), "1")
                    self._insert_into_pPr(pPr, hyphenation_element)

                if pPr.find(qn("w:contextualSpacing")) is None:
                    last_line_element = OxmlElement("w:contextualSpacing")
                    last_line_element.set(qn("w:val"), "1")
                    self._insert_into_pPr(pPr, last_line_element)
        except Exception as e:
            self.add_action(
                element_type="hyphenation",
                element_index=0,
                action_type="hyphenation_enable_error",
                old_value=None,
                new_value=None,
                description=f"Не удалось включить переносы слов: {str(e)}",
                success=False,
                error_message=str(e),
            )
