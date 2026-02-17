"""Корректор для исправления стилей документа.

Отвечает за:
- Исправление шрифтов (Times New Roman, размеры)
- Исправление межстрочного интервала (1.5)
- Исправление первой строки отступа (1.25 см)
- Исправление полей страницы (левое 3см, остальные 2см)
"""

from typing import List, Dict, Any, Set
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING

from .base import BaseCorrector


class StyleCorrector(BaseCorrector):
    """Корректор стилей документа.
    
    Attributes:
        rules: Словарь правил для исправления стилей
    """
    
    # Правила по умолчанию (ГОСТ 7.32-2017)
    DEFAULT_RULES = {
        'font': {
            'name': 'Times New Roman',
            'size': 14,
        },
        'headings': {
            'h1': {'font_size': 14, 'bold': True},
            'h2': {'font_size': 14, 'bold': True},
        },
        'line_spacing': 1.5,
        'first_line_indent': 1.25,  # cm
        'margins': {
            'left': 3.0,    # cm
            'right': 1.5,   # cm
            'top': 2.0,     # cm
            'bottom': 2.0,  # cm
        },
    }
    
    def __init__(self, rules: Dict[str, Any] = None):
        """Инициализация корректора стилей.
        
        Args:
            rules: Словарь правил (используются DEFAULT_RULES если не указано)
        """
        super().__init__()
        self.rules = rules or self.DEFAULT_RULES
    
    def analyze(self, document: Document) -> List[Dict[str, Any]]:
        """Анализирует стили в документе.
        
        Args:
            document: Документ для анализа
            
        Returns:
            Список проблем со стилями
        """
        issues = []
        
        # Проверка шрифтов
        font_issues = self._check_fonts(document)
        issues.extend(font_issues)
        
        # Проверка интервалов
        spacing_issues = self._check_spacing(document)
        issues.extend(spacing_issues)
        
        # Проверка полей
        margin_issues = self._check_margins(document)
        issues.extend(margin_issues)
        
        return issues
    
    def correct(self, document: Document) -> int:
        """Исправляет стили в документе.
        
        Args:
            document: Документ для коррекции
            
        Returns:
            Количество исправленных проблем
        """
        self.clear_actions()
        corrected = 0
        
        # Исправление шрифтов
        corrected += self._correct_font(document)
        
        # Исправление интервалов
        corrected += self._correct_line_spacing(document)
        
        # Исправление отступов
        corrected += self._correct_first_line_indent(document)
        
        # Исправление полей
        corrected += self._correct_margins(document)
        
        return corrected
    
    # ========== Анализ ==========
    
    def _check_fonts(self, document: Document) -> List[Dict[str, Any]]:
        """Проверяет шрифты в документе.
        
        Args:
            document: Документ для проверки
            
        Returns:
            Список проблем со шрифтами
        """
        issues = []
        target_font = self.rules['font']['name']
        target_size = self.rules['font']['size']
        
        for para_idx, paragraph in enumerate(document.paragraphs):
            if not paragraph.text.strip():
                continue
            
            for run in paragraph.runs:
                if run.font.name != target_font:
                    issues.append({
                        'type': 'font_name',
                        'paragraph_idx': para_idx,
                        'current': run.font.name,
                        'expected': target_font,
                    })
                
                if run.font.size and run.font.size != Pt(target_size):
                    issues.append({
                        'type': 'font_size',
                        'paragraph_idx': para_idx,
                        'current': run.font.size,
                        'expected': Pt(target_size),
                    })
        
        return issues
    
    def _check_spacing(self, document: Document) -> List[Dict[str, Any]]:
        """Проверяет межстрочный интервал.
        
        Args:
            document: Документ для проверки
            
        Returns:
            Список проблем с интервалом
        """
        issues = []
        target_spacing = self.rules['line_spacing']
        
        for para_idx, paragraph in enumerate(document.paragraphs):
            if not paragraph.text.strip():
                continue
            
            pf = paragraph.paragraph_format
            if pf.line_spacing != target_spacing:
                issues.append({
                    'type': 'line_spacing',
                    'paragraph_idx': para_idx,
                    'current': pf.line_spacing,
                    'expected': target_spacing,
                })
        
        return issues
    
    def _check_margins(self, document: Document) -> List[Dict[str, Any]]:
        """Проверяет поля страницы.
        
        Args:
            document: Документ для проверки
            
        Returns:
            Список проблем с полями
        """
        issues = []
        target_margins = self.rules['margins']
        
        for section_idx, section in enumerate(document.sections):
            for side in ['left', 'right', 'top', 'bottom']:
                margin_attr = f'{side}_margin'
                current = getattr(section, margin_attr)
                expected = Cm(target_margins[side])
                
                if current != expected:
                    issues.append({
                        'type': f'margin_{side}',
                        'section_idx': section_idx,
                        'current': current,
                        'expected': expected,
                    })
        
        return issues
    
    # ========== Коррекция ==========
    
    def _correct_font(self, document: Document) -> int:
        """Исправляет шрифты в документе.
        
        Args:
            document: Документ для коррекции
            
        Returns:
            Количество примененных исправлений
        """
        corrected = 0
        target_font = self.rules['font']['name']
        target_size = self.rules['font']['size']
        
        # Получаем параграфы в таблицах для специальной обработки
        table_paragraphs = self._get_table_paragraphs(document)
        
        for para_idx, paragraph in enumerate(document.paragraphs):
            if not paragraph.text.strip():
                continue
            
            is_heading = paragraph.style.name.startswith('Heading')
            
            try:
                for run_idx, run in enumerate(paragraph.runs):
                    # Исправление имени шрифта
                    if run.font.name != target_font:
                        old_name = run.font.name
                        run.font.name = target_font
                        
                        self.add_action(
                            element_type='font_name',
                            element_index=para_idx,
                            action_type='font_name_change',
                            old_value=old_name,
                            new_value=target_font,
                            description=f'Изменен шрифт на {target_font}',
                        )
                        corrected += 1
                    
                    # Исправление размера шрифта
                    expected_size = self._get_expected_font_size(paragraph)
                    if run.font.size != Pt(expected_size):
                        old_size = run.font.size
                        run.font.size = Pt(expected_size)
                        
                        self.add_action(
                            element_type='font_size',
                            element_index=para_idx,
                            action_type='font_size_change',
                            old_value=old_size,
                            new_value=Pt(expected_size),
                            description=f'Изменен размер шрифта на {expected_size}pt',
                        )
                        corrected += 1
            
            except Exception as e:
                self.add_action(
                    element_type='font',
                    element_index=para_idx,
                    action_type='font_correction_error',
                    old_value=None,
                    new_value=None,
                    description=f'Ошибка при исправлении шрифта: {str(e)}',
                    success=False,
                    error_message=str(e),
                )
        
        return corrected
    
    def _correct_line_spacing(self, document: Document) -> int:
        """Исправляет межстрочный интервал.
        
        Args:
            document: Документ для коррекции
            
        Returns:
            Количество примененных исправлений
        """
        corrected = 0
        target_spacing = self.rules['line_spacing']
        table_paragraphs = self._get_table_paragraphs(document)
        
        for para_idx, paragraph in enumerate(document.paragraphs):
            # Пропускаем параграфы в таблицах и пустые
            if id(paragraph) in table_paragraphs or not paragraph.text.strip():
                continue
            
            try:
                pf = paragraph.paragraph_format
                
                if pf.line_spacing != target_spacing:
                    old_spacing = pf.line_spacing
                    pf.line_spacing = target_spacing
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    
                    self.add_action(
                        element_type='line_spacing',
                        element_index=para_idx,
                        action_type='spacing_change',
                        old_value=old_spacing,
                        new_value=target_spacing,
                        description=f'Установлен интервал {target_spacing}',
                    )
                    corrected += 1
                    
                    # Сбрасываем интервалы до/после для обычного текста
                    if not paragraph.style.name.startswith('Heading'):
                        if pf.space_before != Pt(0):
                            pf.space_before = Pt(0)
                        if pf.space_after != Pt(0):
                            pf.space_after = Pt(0)
            
            except Exception as e:
                self.add_action(
                    element_type='line_spacing',
                    element_index=para_idx,
                    action_type='spacing_correction_error',
                    old_value=None,
                    new_value=None,
                    description=f'Ошибка при исправлении интервала: {str(e)}',
                    success=False,
                    error_message=str(e),
                )
        
        return corrected
    
    def _correct_first_line_indent(self, document: Document) -> int:
        """Исправляет отступ первой строки абзаца.
        
        Args:
            document: Документ для коррекции
            
        Returns:
            Количество примененных исправлений
        """
        corrected = 0
        target_indent = Cm(self.rules['first_line_indent'])
        
        for para_idx, paragraph in enumerate(document.paragraphs):
            # Пропускаем пустые параграфы и заголовки
            if not paragraph.text.strip() or paragraph.style.name.startswith('Heading'):
                continue
            
            # Пропускаем подписи к рисункам и таблицам
            para_text = paragraph.text.strip().lower()
            if para_text.startswith(('рисунок', 'рис.', 'таблица', 'табл.')):
                continue
            
            try:
                pf = paragraph.paragraph_format
                
                if pf.first_line_indent != target_indent:
                    old_indent = pf.first_line_indent
                    pf.first_line_indent = target_indent
                    
                    self.add_action(
                        element_type='first_line_indent',
                        element_index=para_idx,
                        action_type='indent_change',
                        old_value=old_indent,
                        new_value=target_indent,
                        description=f'Установлен отступ {self.rules["first_line_indent"]} см',
                    )
                    corrected += 1
            
            except Exception as e:
                self.add_action(
                    element_type='first_line_indent',
                    element_index=para_idx,
                    action_type='indent_correction_error',
                    old_value=None,
                    new_value=None,
                    description=f'Ошибка при исправлении отступа: {str(e)}',
                    success=False,
                    error_message=str(e),
                )
        
        return corrected
    
    def _correct_margins(self, document: Document) -> int:
        """Исправляет поля страницы.
        
        Args:
            document: Документ для коррекции
            
        Returns:
            Количество примененных исправлений
        """
        corrected = 0
        target_margins = self.rules['margins']
        
        for section_idx, section in enumerate(document.sections):
            try:
                # Исправление левого поля
                expected_left = Cm(target_margins['left'])
                if section.left_margin != expected_left:
                    section.left_margin = expected_left
                    corrected += 1
                
                # Исправление правого поля
                expected_right = Cm(target_margins['right'])
                if section.right_margin != expected_right:
                    section.right_margin = expected_right
                    corrected += 1
                
                # Исправление верхнего поля
                expected_top = Cm(target_margins['top'])
                if section.top_margin != expected_top:
                    section.top_margin = expected_top
                    corrected += 1
                
                # Исправление нижнего поля
                expected_bottom = Cm(target_margins['bottom'])
                if section.bottom_margin != expected_bottom:
                    section.bottom_margin = expected_bottom
                    corrected += 1
            
            except Exception as e:
                self.add_action(
                    element_type='margins',
                    element_index=section_idx,
                    action_type='margin_correction_error',
                    old_value=None,
                    new_value=None,
                    description=f'Ошибка при исправлении полей: {str(e)}',
                    success=False,
                    error_message=str(e),
                )
        
        return corrected
    
    # ========== Вспомогательные методы ==========
    
    def _get_table_paragraphs(self, document: Document) -> Set[int]:
        """Получает ID параграфов, находящихся в таблицах.
        
        Args:
            document: Документ
            
        Returns:
            Множество ID параграфов в таблицах
        """
        table_paragraphs: Set[int] = set()
        
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        table_paragraphs.add(id(para))
        
        return table_paragraphs
    
    def _get_expected_font_size(self, paragraph) -> int:
        """Определяет ожидаемый размер шрифта для параграфа.
        
        Args:
            paragraph: Параграф
            
        Returns:
            Ожидаемый размер в pt
        """
        is_heading = paragraph.style.name.startswith('Heading')
        
        if is_heading:
            try:
                heading_level = int(paragraph.style.name.replace('Heading ', ''))
                heading_key = f'h{heading_level}'
                
                if heading_key in self.rules['headings']:
                    return self.rules['headings'][heading_key]['font_size']
            except (ValueError, KeyError):
                pass
        
        return self.rules['font']['size']
