import docx
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table, _Row, _Cell
import re
from docx.shared import Length, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import json
import uuid
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from .norm_control_checker import NormControlChecker
from .document_corrector import DocumentCorrector
from datetime import datetime

class DocumentProcessor:
    """
    Класс для обработки DOCX документов и извлечения данных
    """
    def __init__(self, file_path):
        self.file_path = file_path
        
        # Проверяем существование файла перед открытием
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден по пути: {file_path}")
            
        # Проверяем, что это действительно файл
        if not os.path.isfile(file_path):
            raise ValueError(f"Указанный путь не является файлом: {file_path}")
            
        # Проверяем размер файла
        if os.path.getsize(file_path) == 0:
            raise ValueError(f"Файл по пути {file_path} пуст")
            
        # Проверяем расширение файла
        if not file_path.lower().endswith('.docx'):
            raise ValueError(f"Файл должен иметь расширение .docx: {file_path}")
            
        try:
            self.document = docx.Document(file_path)
        except Exception as e:
            print(f"Ошибка при открытии DOCX файла {file_path}: {str(e)}")
            raise
    
    def extract_data(self):
        """
        Извлекает все необходимые данные из документа для анализа
        """
        document_data = {}
        
        # Извлекаем разные типы данных, защищая каждый вызов от ошибок
        try:
            document_data['paragraphs'] = self._extract_paragraphs()
        except Exception as e:
            print(f"Ошибка при извлечении параграфов: {str(e)}")
            document_data['paragraphs'] = []
            
        try:
            document_data['tables'] = self._extract_tables()
        except Exception as e:
            print(f"Ошибка при извлечении таблиц: {str(e)}")
            document_data['tables'] = []
            
        try:
            document_data['headings'] = self._extract_headings()
        except Exception as e:
            print(f"Ошибка при извлечении заголовков: {str(e)}")
            document_data['headings'] = []
            
        try:
            document_data['bibliography'] = self._extract_bibliography()
        except Exception as e:
            print(f"Ошибка при извлечении библиографии: {str(e)}")
            document_data['bibliography'] = []
            
        try:
            document_data['styles'] = self._extract_styles()
        except Exception as e:
            print(f"Ошибка при извлечении стилей: {str(e)}")
            document_data['styles'] = {}
            
        try:
            document_data['page_setup'] = self._extract_page_setup()
        except Exception as e:
            print(f"Ошибка при извлечении настроек страницы: {str(e)}")
            document_data['page_setup'] = {}
            
        try:
            document_data['images'] = self._extract_images()
        except Exception as e:
            print(f"Ошибка при извлечении изображений: {str(e)}")
            document_data['images'] = []
            
        try:
            document_data['page_numbers'] = self._extract_page_numbers()
        except Exception as e:
            print(f"Ошибка при извлечении нумерации страниц: {str(e)}")
            document_data['page_numbers'] = {
                'has_page_numbers': False,
                'position': None,
                'first_numbered_page': None,
                'alignment': None
            }
            
        try:
            document_data['document_properties'] = self._extract_document_properties()
        except Exception as e:
            print(f"Ошибка при извлечении свойств документа: {str(e)}")
            document_data['document_properties'] = {}
            
        # Выделяем титульный лист
        document_data['title_page'] = self._extract_title_page(document_data.get('paragraphs', []))
        
        return document_data
    
    def _extract_paragraphs(self):
        """
        Извлекает все параграфы документа с их стилями
        """
        paragraphs = []
        for i, para in enumerate(self.document.paragraphs):
            if not para.text.strip():
                continue  # Пропускаем пустые параграфы
                
            paragraphs.append({
                'index': i,
                'text': para.text,
                'style': para.style.name if para.style else 'Normal',
                'alignment': self._get_paragraph_alignment(para),
                'font': self._get_paragraph_font(para),
                'line_spacing': self._get_paragraph_line_spacing(para),
                'paragraph_format': self._get_paragraph_format(para),
                'is_heading': para.style.name.startswith('Heading') if para.style else False,
                'list_info': self._get_list_info(para)
            })
        return paragraphs
    
    def _extract_tables(self):
        """
        Извлекает все таблицы документа с их содержимым и стилями
        """
        tables = []
        for i, table in enumerate(self.document.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
                
            # Пытаемся получить заголовок таблицы из предыдущего параграфа
            table_title = None
            table_xml_element = table._element
            prev_elem = table_xml_element.getprevious()
            
            if prev_elem is not None and isinstance(prev_elem, CT_P):
                prev_para = Paragraph(prev_elem, table._parent)
                if prev_para.text.strip().lower().startswith('таблица'):
                    table_title = prev_para.text.strip()
            
            tables.append({
                'index': i,
                'rows': rows,
                'style': table.style.name if hasattr(table, 'style') and table.style else 'TableNormal',
                'num_rows': len(table.rows),
                'num_cols': len(table.rows[0].cells) if table.rows else 0,
                'title': table_title,
                'has_header': self._table_has_header(table)
            })
        return tables
    
    def _table_has_header(self, table):
        """
        Проверяет, имеет ли таблица заголовочную строку
        """
        if not table.rows:
            return False
            
        # Проверяем, отличается ли форматирование первой строки от других
        if len(table.rows) > 1:
            # Если первая строка имеет другое форматирование (например, жирный шрифт)
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.bold:
                            return True
        return False
    
    def _extract_headings(self):
        """
        Извлекает заголовки документа по уровням
        """
        headings = []
        for i, para in enumerate(self.document.paragraphs):
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                except ValueError:
                    level = 0
                    
                font_info = self._get_paragraph_font(para)
                
                headings.append({
                    'index': i,
                    'text': para.text,
                    'level': level,
                    'style': para.style.name,
                    'font': font_info,
                    'alignment': self._get_paragraph_alignment(para),
                    'has_number': bool(re.match(r'^\d+(\.\d+)*\.?\s', para.text)),
                    'has_ending_dot': para.text.strip().endswith('.'),
                    'all_caps': all(c.isupper() for c in para.text if c.isalpha()),
                    'para_format': self._get_paragraph_format(para)
                })
        return headings
    
    def _extract_bibliography(self):
        """
        Пытается найти и извлечь список литературы
        """
        bibliography_items = []
        bibliography_started = False
        bibliography_section_titles = [
            'список литературы', 'список используемых источников', 
            'список использованных источников', 'список источников',
            'библиографический список', 'библиография',
            'список использованной литературы', 'литература',
            'использованные источники', 'источники', 'использованная литература'
        ]
        
        # Параграфы, которые могут означать окончание списка литературы
        end_section_identifiers = [
            'приложение', 'глоссарий', 'алфавитный указатель', 
            'предметный указатель', 'указатель имен'
        ]
        
        paragraphs = self.document.paragraphs
        
        for i, para in enumerate(paragraphs):
            para_text = para.text.lower().strip()
            
            # Поиск начала списка литературы
            if not bibliography_started:
                if any(para_text == title for title in bibliography_section_titles) or \
                   any(para_text.startswith(title) for title in bibliography_section_titles):
                    bibliography_started = True
                    bibliography_heading_index = i
                    continue
            else:
                # Проверка на окончание списка литературы
                if para.style.name.startswith('Heading') or \
                   any(para_text.startswith(end) for end in end_section_identifiers):
                    break
                
                # Пропускаем пустые параграфы
                if not para_text:
                    continue
                
                # Проверка на наличие нумерации (например, "1. ", "1) ", "[1]", и т.д.)
                is_numbered = bool(re.match(r'^\d+[\.\)\]]', para_text)) or \
                              bool(re.match(r'^\[\d+\]', para_text))
                
                # Если параграф выглядит как библиографическая запись
                if is_numbered or self._looks_like_bibliography_item(para_text):
                    # Обрабатываем нумерацию, убирая её из текста
                    clean_text = re.sub(r'^\d+[\.\)\]]?\s*', '', para_text)
                    clean_text = re.sub(r'^\[\d+\]\s*', '', clean_text)
                    
                    # Добавляем в список, если это не просто номер
                    if len(clean_text) > 3:  # проверка, что это не просто номер
                        bibliography_items.append({
                            'text': para.text,
                            'index': i,
                            'font': self._get_paragraph_font(para),
                            'is_numbered': is_numbered,
                            'alignment': self._get_paragraph_alignment(para)
                        })
                    
                # Если параграф похож на продолжение предыдущей записи
                elif bibliography_items and len(para_text) > 3 and not para.style.name.startswith('Heading'):
                    # Проверяем, не начинается ли параграф с заглавной буквы 
                    # (что может указывать на новую запись)
                    if not (para_text[0].isupper() and bibliography_items[-1]['text'].endswith('.')):
                        # Добавляем к предыдущей записи
                        bibliography_items[-1]['text'] = f"{bibliography_items[-1]['text']} {para.text}"
                    else:
                        bibliography_items.append({
                            'text': para.text,
                            'index': i,
                            'font': self._get_paragraph_font(para),
                            'is_numbered': False,
                            'alignment': self._get_paragraph_alignment(para)
                        })
                
        return bibliography_items
    
    def _looks_like_bibliography_item(self, text):
        """
        Проверяет, похож ли текст на библиографическую запись
        """
        # Типичные признаки библиографической записи
        patterns = [
            r'^[А-Я][а-я]+,\s[А-Я]\.',  # Фамилия, И.О.
            r'\d{4}\.',  # Год издания с точкой
            r'[–—-]\s\d+\sс\.',  # Указание на количество страниц
            r'\[Электронный\sресурс\]',  # Электронный ресурс
            r'URL:',  # URL
            r'дата\sобращения',  # Дата обращения
            r'ГОСТ\s',  # ГОСТ
            r'№\s?\d+',  # Номер (для законов и постановлений)
            r'от\s\d{2}\.\d{2}\.\d{4}',  # Дата (для законов и постановлений)
        ]
        
        # Если текст соответствует хотя бы одному из паттернов
        return any(bool(re.search(pattern, text, re.IGNORECASE)) for pattern in patterns)
    
    def _extract_styles(self):
        """
        Извлекает информацию о стилях документа
        """
        styles = {}
        for style in self.document.styles:
            style_info = {}
            if hasattr(style, 'font') and style.font:
                font = style.font
                style_info['font'] = {
                    'name': font.name if hasattr(font, 'name') else None,
                    'size': font.size.pt if hasattr(font, 'size') and font.size else None,
                    'bold': font.bold if hasattr(font, 'bold') else None,
                    'italic': font.italic if hasattr(font, 'italic') else None,
                }
                
            if hasattr(style, 'paragraph_format') and style.paragraph_format:
                pf = style.paragraph_format
                style_info['paragraph_format'] = {
                    'alignment': pf.alignment if hasattr(pf, 'alignment') else None,
                    'line_spacing': pf.line_spacing if hasattr(pf, 'line_spacing') else None,
                    'space_before': pf.space_before.pt if hasattr(pf, 'space_before') and pf.space_before else None,
                    'space_after': pf.space_after.pt if hasattr(pf, 'space_after') and pf.space_after else None,
                    'first_line_indent': pf.first_line_indent.cm if hasattr(pf, 'first_line_indent') and pf.first_line_indent else None,
                }
                
            styles[style.name] = style_info
            
        return styles
    
    def _extract_page_setup(self):
        """
        Извлекает настройки страницы
        """
        page_setup = {}
        
        # Извлекаем информацию о разделах документа
        for i, section in enumerate(self.document.sections):
            section_data = {}
            
            # Размеры полей
            section_data['left_margin'] = section.left_margin.cm if section.left_margin else None
            section_data['right_margin'] = section.right_margin.cm if section.right_margin else None
            section_data['top_margin'] = section.top_margin.cm if section.top_margin else None
            section_data['bottom_margin'] = section.bottom_margin.cm if section.bottom_margin else None
            
            # Размер страницы
            section_data['page_width'] = section.page_width.cm if section.page_width else None
            section_data['page_height'] = section.page_height.cm if section.page_height else None
            
            # Ориентация
            section_data['orientation'] = 'portrait' if section.orientation == 0 else 'landscape'
            
            # Колонтитулы
            section_data['header_distance'] = section.header_distance.cm if section.header_distance else None
            section_data['footer_distance'] = section.footer_distance.cm if section.footer_distance else None
            
            page_setup[f'section_{i+1}'] = section_data
        
        return page_setup
    
    def _extract_images(self):
        """
        Извлекает изображения из документа и их описания
        """
        images = []
        
        # Проход по всем параграфам для поиска рисунков
        for i, paragraph in enumerate(self.document.paragraphs):
            if paragraph.text.strip().lower().startswith(('рис.', 'рисунок')):
                # Это похоже на подпись к рисунку
                image_caption = paragraph.text.strip()
                
                # Ищем предыдущий параграф, который может содержать изображение
                j = i - 1
                while j >= 0:
                    prev_para = self.document.paragraphs[j]
                    if hasattr(prev_para, 'runs'):
                        for run in prev_para.runs:
                            if hasattr(run, '_element') and run._element.findall('.//pic:pic', {'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}):
                                images.append({
                                    'caption': image_caption,
                                    'caption_index': i,
                                    'image_para_index': j,
                                    'has_number': bool(re.search(r'рис\w*\s+\d+', image_caption.lower())),
                                    'ends_with_dot': image_caption.endswith('.'),
                                    'alignment': self._get_paragraph_alignment(paragraph)
                                })
                                break
                    j -= 1
                    # Ограничиваем поиск до 3 параграфов назад
                    if j < i - 3:
                        break
        
        return images

    def _extract_page_numbers(self):
        """
        Извлекает информацию о нумерации страниц
        """
        page_numbers = {
            'has_page_numbers': False,
            'position': None,
            'first_numbered_page': None,
            'alignment': None
        }
        
        try:
            # Проверяем колонтитулы на наличие нумерации
            for section in self.document.sections:
                # Проверка нижнего колонтитула
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        # Проверка на None и пустой текст
                        if paragraph.text is None:
                            continue
                            
                        if '{PAGE}' in paragraph.text or '{PAGE \\*' in paragraph.text or '{PAGE *' in paragraph.text:
                            page_numbers['has_page_numbers'] = True
                            page_numbers['position'] = 'footer'
                            
                            # Определяем позицию (центр, справа, слева)
                            alignment = self._get_paragraph_alignment(paragraph)
                            if alignment:
                                if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                                    page_numbers['alignment'] = 'center'
                                elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                                    page_numbers['alignment'] = 'right'
                                elif alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
                                    page_numbers['alignment'] = 'left'
                            break
                            
                # Проверка верхнего колонтитула
                if not page_numbers['has_page_numbers'] and section.header:
                    for paragraph in section.header.paragraphs:
                        # Проверка на None и пустой текст
                        if paragraph.text is None:
                            continue
                            
                        if '{PAGE}' in paragraph.text or '{PAGE \\*' in paragraph.text or '{PAGE *' in paragraph.text:
                            page_numbers['has_page_numbers'] = True
                            page_numbers['position'] = 'header'
                            
                            # Определяем позицию (центр, справа, слева)
                            alignment = self._get_paragraph_alignment(paragraph)
                            if alignment:
                                if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                                    page_numbers['alignment'] = 'center'
                                elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                                    page_numbers['alignment'] = 'right'
                                elif alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
                                    page_numbers['alignment'] = 'left'
                            break
        except Exception as e:
            # В случае ошибки при извлечении номеров страниц, логируем её и возвращаем базовую информацию
            print(f"Ошибка при извлечении нумерации страниц: {str(e)}")
            # Возвращаем базовую информацию о нумерации
            return {
                'has_page_numbers': False,
                'position': None,
                'first_numbered_page': None,
                'alignment': None
            }
        
        return page_numbers

    def _extract_document_properties(self):
        """
        Извлекает метаданные документа
        """
        properties = {}
        
        # Извлекаем свойства из Core Properties
        if hasattr(self.document, 'core_properties'):
            cp = self.document.core_properties
            properties['title'] = cp.title
            properties['author'] = cp.author
            properties['created'] = cp.created.isoformat() if cp.created else None
            properties['modified'] = cp.modified.isoformat() if cp.modified else None
            properties['last_modified_by'] = cp.last_modified_by
            properties['revision'] = cp.revision
            
        # Статистика документа
        statistics = {
            'paragraph_count': len(self.document.paragraphs),
            'table_count': len(self.document.tables),
            'section_count': len(self.document.sections),
            'heading_count': sum(1 for para in self.document.paragraphs if para.style and para.style.name.startswith('Heading'))
        }
        properties['statistics'] = statistics
        
        return properties
    
    def _get_paragraph_alignment(self, paragraph):
        """
        Определяет выравнивание параграфа
        """
        if paragraph.paragraph_format and paragraph.paragraph_format.alignment:
            return paragraph.paragraph_format.alignment
        return None
    
    def _get_paragraph_font(self, paragraph):
        """
        Извлекает информацию о шрифте и форматировании текста в параграфе
        """
        font_info = {}
        
        try:
            # Если в параграфе нет текста, возвращаем пустой словарь
            if not paragraph or not hasattr(paragraph, 'text') or not paragraph.text or not paragraph.text.strip():
                return font_info
                
            # Проверяем runs (части текста с одинаковым форматированием)
            if not hasattr(paragraph, 'runs'):
                return font_info
                
            runs = paragraph.runs
            if runs and len(runs) > 0:
                # Берем первый run как основной для получения шрифта
                main_run = runs[0]
                if main_run and hasattr(main_run, 'font') and main_run.font:
                    font = main_run.font
                    font_info['name'] = font.name if hasattr(font, 'name') else None
                    font_info['size'] = font.size.pt if hasattr(font, 'size') and font.size else None
                    font_info['bold'] = font.bold if hasattr(font, 'bold') else None
                    font_info['italic'] = font.italic if hasattr(font, 'italic') else None
                    font_info['underline'] = font.underline if hasattr(font, 'underline') else None
                    
                    # Проверка атрибутов цвета
                    if hasattr(font, 'color') and font.color:
                        font_info['color'] = font.color.rgb if hasattr(font.color, 'rgb') else None
                    else:
                        font_info['color'] = None
                
                # Проверяем, одинаково ли форматирование во всех runs
                if len(runs) > 1:
                    try:
                        font_info['consistent_formatting'] = all(
                            ((run.font.name == main_run.font.name) if run.font and hasattr(run.font, 'name') and main_run.font and hasattr(main_run.font, 'name') else True) and
                            ((run.font.size == main_run.font.size) if run.font and hasattr(run.font, 'size') and main_run.font and hasattr(main_run.font, 'size') else True) and
                            ((run.font.bold == main_run.font.bold) if run.font and hasattr(run.font, 'bold') and main_run.font and hasattr(main_run.font, 'bold') else True)
                            for run in runs[1:]
                        )
                    except Exception as inner_e:
                        print(f"Ошибка при проверке согласованности форматирования: {str(inner_e)}")
                        font_info['consistent_formatting'] = False
                
        except Exception as e:
            print(f"Ошибка при извлечении информации о шрифте: {str(e)}")
            # В случае ошибки возвращаем базовую информацию
            return {
                'name': None,
                'size': None,
                'bold': None,
                'italic': None,
                'underline': None,
                'color': None,
                'consistent_formatting': False
            }
            
        return font_info
    
    def _get_paragraph_line_spacing(self, paragraph):
        """
        Определяет межстрочный интервал параграфа
        """
        if paragraph.paragraph_format:
            pf = paragraph.paragraph_format
            if pf.line_spacing:
                return pf.line_spacing
        return None

    def _get_paragraph_format(self, paragraph):
        """
        Извлекает данные о форматировании параграфа
        """
        para_format = {}
        
        if paragraph.paragraph_format:
            pf = paragraph.paragraph_format
            para_format['alignment'] = pf.alignment
            para_format['line_spacing'] = pf.line_spacing
            para_format['line_spacing_rule'] = pf.line_spacing_rule
            para_format['space_before'] = pf.space_before.pt if pf.space_before else None
            para_format['space_after'] = pf.space_after.pt if pf.space_after else None
            para_format['first_line_indent'] = pf.first_line_indent.cm if pf.first_line_indent else None
            para_format['left_indent'] = pf.left_indent.cm if pf.left_indent else None
            para_format['right_indent'] = pf.right_indent.cm if pf.right_indent else None
            
        return para_format

    def _get_list_info(self, paragraph):
        """
        Определяет, является ли параграф элементом списка и возвращает информацию о нем
        """
        list_info = {
            'is_list_item': False,
            'list_type': None,  # 'bullet' или 'numbered'
            'list_level': 0
        }
        
        # Проверка на нумерованный список по тексту
        if re.match(r'^\d+[.)]\s', paragraph.text) or re.match(r'^[a-z][.)]\s', paragraph.text):
            list_info['is_list_item'] = True
            list_info['list_type'] = 'numbered'
            
            # Проверка уровня вложенности по отступу
            if paragraph.paragraph_format and paragraph.paragraph_format.left_indent:
                indent = paragraph.paragraph_format.left_indent.cm
                list_info['list_level'] = int(indent / 0.5) if indent > 0 else 0
        
        # Проверка на маркированный список по тексту
        elif re.match(r'^[•\-–—]\s', paragraph.text):
            list_info['is_list_item'] = True
            list_info['list_type'] = 'bullet'
            
            # Проверка уровня вложенности по отступу
            if paragraph.paragraph_format and paragraph.paragraph_format.left_indent:
                indent = paragraph.paragraph_format.left_indent.cm
                list_info['list_level'] = int(indent / 0.5) if indent > 0 else 0
        
        return list_info 

    def process_document(self, file_path):
        """
        Обрабатывает документ и возвращает результаты проверки
        
        Args:
            file_path: Путь к документу
            
        Returns:
            dict: Результаты проверки
        """
        try:
            # Извлечение структурированных данных из документа
            document_data = self._extract_document_data(file_path)
            
            # Проверка документа на соответствие требованиям нормоконтроля
            check_results = self.checker.check_document(document_data)
            
            # Подготовка результатов
            results = {
                'success': True,
                'check_results': check_results,
                'file_name': os.path.basename(file_path),
                'timestamp': datetime.now().isoformat(),
                'document_data': {
                    'title': document_data.get('title', 'Неизвестное название'),
                    'author': document_data.get('author', 'Неизвестный автор'),
                    'paragraphs_count': len(document_data.get('paragraphs', [])),
                    'headings_count': len(document_data.get('headings', [])),
                    'images_count': len(document_data.get('images', [])),
                    'tables_count': len(document_data.get('tables', [])),
                }
            }
            
            return results
            
        except Exception as e:
            # В случае ошибки возвращаем сообщение об ошибке
            return {
                'success': False,
                'error_message': str(e)
            }
    
    def correct_document(self, file_path, check_results):
        """
        Исправляет найденные несоответствия в документе
        
        Args:
            file_path: Путь к документу
            check_results: Результаты проверки
            
        Returns:
            str: Путь к исправленному документу
        """
        try:
            # Исправление документа
            corrected_file_path = self.corrector.correct_document(file_path, check_results)
            
            return corrected_file_path
            
        except Exception as e:
            # В случае ошибки возвращаем None
            print(f"Ошибка при исправлении документа: {str(e)}")
            return None
    
    def generate_report_document(self, check_results, file_name):
        """
        Генерирует документ Word с подробным отчетом о проверке
        
        Args:
            check_results: Результаты проверки
            file_name: Имя исходного файла
            
        Returns:
            str: Путь к сгенерированному отчету
        """
        try:
            # Создаем новый документ Word для отчета
            doc = Document()
            
            # Настраиваем стили документа
            self._setup_document_styles(doc)
            
            # Формируем заголовок отчета
            title_para = doc.add_paragraph()
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title_para.add_run("ОТЧЕТ О ПРОВЕРКЕ ДОКУМЕНТА НА СООТВЕТСТВИЕ ТРЕБОВАНИЯМ НОРМОКОНТРОЛЯ")
            title_run.bold = True
            title_run.font.size = Pt(16)
            
            # Добавляем информацию о файле
            doc.add_paragraph()
            file_para = doc.add_paragraph()
            file_run = file_para.add_run(f"Документ: {file_name}")
            file_run.bold = True
            
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Дата проверки: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            
            # Добавляем раздел с итогами проверки
            doc.add_paragraph()
            doc.add_heading("Итоги проверки", level=1)
            
            summary_para = doc.add_paragraph()
            total_issues = check_results.get('total_issues_count', 0)
            summary_para.add_run(f"Всего найдено несоответствий: ").bold = True
            summary_para.add_run(f"{total_issues}")
            
            if total_issues > 0:
                # Добавляем статистику по категориям
                doc.add_heading("Статистика по категориям", level=2)
                
                statistics = check_results.get('statistics', {})
                categories = statistics.get('categories', {})
                severity = statistics.get('severity', {})
                
                # Статистика по серьезности
                severity_para = doc.add_paragraph()
                severity_para.add_run("Распределение по серьезности:").bold = True
                severity_para.add_run(f"\n• Критических: {severity.get('high', 0)}")
                severity_para.add_run(f"\n• Средних: {severity.get('medium', 0)}")
                severity_para.add_run(f"\n• Незначительных: {severity.get('low', 0)}")
                
                # Статистика по категориям
                categories_para = doc.add_paragraph()
                categories_para.add_run("Распределение по категориям:").bold = True
                
                for category, count in categories.items():
                    category_name = self._get_category_name(category)
                    categories_para.add_run(f"\n• {category_name}: {count}")
                
                # Добавляем перечень всех найденных проблем
                doc.add_heading("Детальный отчет о несоответствиях", level=1)
                
                # Группируем проблемы по категориям
                issues_by_category = {}
                for issue in check_results.get('issues', []):
                    category = issue.get('type', '').split('_')[0]
                    if category not in issues_by_category:
                        issues_by_category[category] = []
                    issues_by_category[category].append(issue)
                
                # Добавляем проблемы по категориям
                for category, issues in issues_by_category.items():
                    category_name = self._get_category_name(category)
                    doc.add_heading(f"{category_name}", level=2)
                    
                    for i, issue in enumerate(issues, 1):
                        issue_para = doc.add_paragraph(style='List Bullet')
                        description = issue.get('description', 'Нет описания')
                        location = issue.get('location', 'Неизвестно')
                        severity = issue.get('severity', 'low')
                        
                        severity_text = {
                            'high': 'Критическая',
                            'medium': 'Средняя',
                            'low': 'Незначительная'
                        }.get(severity, 'Неизвестная')
                        
                        issue_para.add_run(f"{description}").bold = True
                        issue_para.add_run(f"\nРасположение: {location}")
                        issue_para.add_run(f"\nСерьезность: {severity_text}")
                        
                        if issue.get('auto_fixable'):
                            issue_para.add_run("\nМожет быть исправлено автоматически").italic = True
                
                # Добавляем рекомендации
                doc.add_heading("Рекомендации по исправлению", level=1)
                recommendations_para = doc.add_paragraph()
                recommendations_para.add_run("Для исправления обнаруженных несоответствий рекомендуется:").bold = True
                recommendations_para.add_run("\n\n1. Проверить и исправить все критические несоответствия в первую очередь.")
                recommendations_para.add_run("\n2. Обратить внимание на форматирование текста и структуру документа.")
                recommendations_para.add_run("\n3. Убедиться, что список литературы оформлен в соответствии с ГОСТ Р 7.0.100-2018.")
                recommendations_para.add_run("\n4. Проверить корректность оформления всех таблиц и изображений.")
                
                # Статистика автоматически исправляемых ошибок
                auto_fixable = statistics.get('auto_fixable_count', 0)
                if auto_fixable > 0:
                    auto_fix_para = doc.add_paragraph()
                    auto_fix_ratio = round((auto_fixable / total_issues) * 100, 1)
                    auto_fix_para.add_run(f"\nВажно: {auto_fixable} из {total_issues} несоответствий ({auto_fix_ratio}%) могут быть исправлены автоматически с помощью нашего сервиса.").bold = True
            else:
                # Если проблем нет
                success_para = doc.add_paragraph()
                success_para.add_run("Поздравляем! Документ полностью соответствует требованиям нормоконтроля.").bold = True
            
            # Добавляем заключение
            doc.add_heading("Заключение", level=1)
            conclusion_para = doc.add_paragraph()
            if total_issues > 0:
                if check_results.get('statistics', {}).get('severity', {}).get('high', 0) > 0:
                    conclusion_para.add_run("Документ требует доработки. Необходимо исправить все критические несоответствия перед сдачей.")
                else:
                    conclusion_para.add_run("Документ требует незначительных корректировок. Рекомендуется устранить найденные несоответствия.")
            else:
                conclusion_para.add_run("Документ соответствует всем требованиям нормоконтроля и готов к сдаче.")
            
            # Сохраняем отчет
            report_filename = f"report_{uuid.uuid4().hex[:8]}.docx"
            report_path = os.path.join(self.app.static_folder, 'reports', report_filename)
            doc.save(report_path)
            
            return f"/static/reports/{report_filename}"
            
        except Exception as e:
            print(f"Ошибка при генерации отчета: {str(e)}")
            return None
    
    def _setup_document_styles(self, doc):
        """
        Настраивает стили документа
        """
        # Настройка стилей заголовков
        for i in range(1, 10):
            style_name = f'Heading {i}'
            if style_name in doc.styles:
                style = doc.styles[style_name]
                style.font.name = 'Times New Roman'
                style.font.size = Pt(14)
                style.font.bold = True
                
                # Особые настройки для разных уровней
                if i == 1:
                    style.font.size = Pt(16)
                    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
        # Настройка стиля обычного текста
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.5
        
        # Настройка стиля маркированного списка
        if 'List Bullet' in doc.styles:
            style = doc.styles['List Bullet']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(14)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.5
    
    def _get_category_name(self, category):
        """
        Возвращает русское название категории проблем
        """
        categories = {
            'font': 'Шрифты и форматирование',
            'margins': 'Поля страницы',
            'line': 'Межстрочный интервал',
            'paragraphs': 'Параграфы и отступы',
            'headings': 'Заголовки',
            'bibliography': 'Список литературы',
            'images': 'Изображения',
            'tables': 'Таблицы',
            'page': 'Нумерация страниц',
            'lists': 'Списки',
            'references': 'Ссылки',
            'structure': 'Структура документа'
        }
        return categories.get(category, f'Категория "{category}"')
    
    def _extract_document_data(self, file_path):
        """
        Извлекает структурированные данные из документа формата DOCX
        
        Args:
            file_path: Путь к документу
            
        Returns:
            dict: Структурированные данные документа
        """
        # Открываем документ
        doc = Document(file_path)
        
        # Инициализируем структуру данных
        document_data = {
            'paragraphs': [],
            'headings': [],
            'tables': [],
            'images': [],
            'page_setup': {},
            'bibliography': [],
            'title_page': []
        }
        
        # Получаем свойства документа
        prop_names = ['title', 'author', 'company', 'comments', 'category', 'subject']
        for name in prop_names:
            if hasattr(doc.core_properties, name):
                document_data[name] = getattr(doc.core_properties, name)
        
        # Получаем настройки страницы
        if doc.sections:
            section = doc.sections[0]
            document_data['page_setup'] = {
                'section_0': {
                    'left_margin': section.left_margin.cm,
                    'right_margin': section.right_margin.cm,
                    'top_margin': section.top_margin.cm,
                    'bottom_margin': section.bottom_margin.cm,
                    'page_height': section.page_height.cm,
                    'page_width': section.page_width.cm,
                    'orientation': 'portrait' if section.page_width < section.page_height else 'landscape',
                    'header_distance': section.header_distance.cm,
                    'footer_distance': section.footer_distance.cm,
                }
            }
        
        # Обрабатываем параграфы
        title_page_paragraphs = []
        title_page_found = False
        for i, para in enumerate(doc.paragraphs):
            # Пропускаем пустые параграфы
            if not para.text.strip():
                continue
                
            # Собираем данные о параграфе
            para_data = {
                'index': i,
                'text': para.text,
                'style': para.style.name if para.style else 'Normal',
                'is_heading': para.style.name.startswith('Heading') if para.style else False,
                'heading_level': int(para.style.name.split(' ')[1]) if para.style and para.style.name.startswith('Heading') else 0,
                'alignment': para.alignment,
                'line_spacing': para.paragraph_format.line_spacing if para.paragraph_format.line_spacing else None,
                'paragraph_format': {
                    'first_line_indent': para.paragraph_format.first_line_indent.cm if para.paragraph_format.first_line_indent else None,
                    'left_indent': para.paragraph_format.left_indent.cm if para.paragraph_format.left_indent else None,
                    'right_indent': para.paragraph_format.right_indent.cm if para.paragraph_format.right_indent else None,
                    'space_before': para.paragraph_format.space_before.pt if para.paragraph_format.space_before else None,
                    'space_after': para.paragraph_format.space_after.pt if para.paragraph_format.space_after else None,
                    'alignment': para.alignment
                }
            }
            
            # Анализ списков
            para_data['list_info'] = {
                'is_list_item': bool(para._p.pPr and para._p.pPr.numPr)
            }
            
            # Анализ шрифта (берем из первого run или из всех, если они отличаются)
            font_info = {}
            if para.runs:
                # Анализируем первый run
                first_run = para.runs[0]
                font_info = {
                    'name': first_run.font.name,
                    'size': first_run.font.size.pt if first_run.font.size else None,
                    'bold': first_run.font.bold,
                    'italic': first_run.font.italic,
                    'underline': first_run.font.underline,
                    'color': first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None,
                    'consistent_formatting': True
                }
                
                # Проверяем, отличается ли форматирование в разных участках текста
                for run in para.runs[1:]:
                    if (run.font.name != font_info['name'] or
                        (run.font.size and run.font.size.pt != font_info['size']) or
                        run.font.bold != font_info['bold'] or
                        run.font.italic != font_info['italic']):
                        font_info['consistent_formatting'] = False
                        break
            
            para_data['font'] = font_info
            
            # Добавляем параграф в список
            document_data['paragraphs'].append(para_data)
            
            # Логика выделения титульного листа
            if not title_page_found:
                text_lower = para.text.strip().lower()
                if para_data['is_heading'] and para_data['heading_level'] == 1:
                    title_page_found = True
                elif any(word in text_lower for word in ['содержание', 'введение']):
                    title_page_found = True
                else:
                    title_page_paragraphs.append({'index': i, 'text': para.text})
            
            # Если это заголовок, добавляем его в отдельный список
            if para_data['is_heading']:
                heading_data = {
                    'index': i,
                    'text': para.text,
                    'level': para_data['heading_level'],
                    'alignment': para.alignment,
                    'font': font_info,
                    'has_ending_dot': para.text.strip().endswith('.')
                }
                document_data['headings'].append(heading_data)
                
            # Проверяем, является ли этот параграф частью списка литературы
            # (для простоты ищем заголовок "Список литературы" или "Библиография" и берем все параграфы после него)
            lower_text = para.text.lower()
            if lower_text.startswith('список литературы') or lower_text.startswith('библиография'):
                # Помечаем, что следующие параграфы входят в библиографию
                bibliography_section = True
                
            # Для упрощения, будем считать любой параграф с цифрой в начале как элемент списка литературы
            if bibliography_section and re.match(r'^\d+\.\s', para.text):
                bibliography_item = {
                    'index': len(document_data['bibliography']) + 1,
                    'text': para.text,
                    'font': font_info,
                    'alignment': para.alignment
                }
                document_data['bibliography'].append(bibliography_item)
        
        # Получаем информацию о таблицах
        for i, table in enumerate(doc.tables):
            table_data = {
                'index': i,
                'rows': len(table.rows),
                'columns': len(table.columns) if table.rows else 0,
            }
            
            # Пытаемся найти заголовок таблицы (параграф перед таблицей, начинающийся с "Таблица")
            if i > 0 and document_data['paragraphs']:
                for para in reversed(document_data['paragraphs']):
                    if para['text'].strip().startswith('Таблица'):
                        table_data['title'] = para['text'].strip()
                        break
            
            document_data['tables'].append(table_data)
        
        # Получаем информацию о колонтитулах и номерах страниц
        document_data['page_numbers'] = {
            'has_page_numbers': False,
            'position': None,
            'alignment': None
        }
        
        for section in doc.sections:
            # Проверяем наличие номеров страниц в нижнем колонтитуле
            if section.footer.paragraphs and any('PAGE' in p._element.xml for p in section.footer.paragraphs):
                document_data['page_numbers'] = {
                    'has_page_numbers': True,
                    'position': 'footer',
                    'alignment': 'center'  # Упрощенно считаем, что номера страниц по центру
                }
                break
            
            # Проверяем наличие номеров страниц в верхнем колонтитуле
            if section.header.paragraphs and any('PAGE' in p._element.xml for p in section.header.paragraphs):
                document_data['page_numbers'] = {
                    'has_page_numbers': True,
                    'position': 'header',
                    'alignment': 'center'  # Упрощенно считаем, что номера страниц по центру
                }
                break
        
        document_data['title_page'] = title_page_paragraphs
        
        return document_data
    
    def _extract_title_page(self, paragraphs):
        """
        Выделяет параграфы титульного листа (до первого Heading 1 или до 'СОДЕРЖАНИЕ'/'ВВЕДЕНИЕ')
        """
        title_page = []
        for para in paragraphs:
            text_lower = para.get('text', '').strip().lower()
            if para.get('style', '').startswith('Heading') and para.get('style', '') == 'Heading 1':
                break
            if any(word in text_lower for word in ['содержание', 'введение']):
                break
            title_page.append({'index': para.get('index'), 'text': para.get('text')})
        return title_page 