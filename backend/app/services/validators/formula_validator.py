"""
Валидатор для проверки оформления формул в документе.
"""

from typing import Dict, Any, List
import time
import re
from . import BaseValidator, ValidationResult, ValidationIssue, Severity


class FormulaValidator(BaseValidator):
    """
    Валидатор для проверки оформления формул.

    Проверяет:
    - Наличие нумерации формул
    - Правильность нумерации формул (последовательность)
    - Наличие ссылок на формулы в тексте
    - Форматирование подписи формулы
    - Интервалы до и после формул
    """

    @property
    def name(self) -> str:
        return "FormulaValidator"

    def validate(self, document: Any, document_data: Dict[str, Any]) -> ValidationResult:
        """
        Проверяет оформление формул в документе.

        Args:
            document: python-docx Document объект
            document_data: Извлечённые данные документа

        Returns:
            ValidationResult с найденными проблемами
        """
        start_time = time.time()
        issues: List[ValidationIssue] = []

        # Получаем требования из профиля
        formula_config = self._get_rule_config("formulas", {})

        require_numbering = formula_config.get("require_numbering", True)
        numbering_format = formula_config.get("numbering_format", "({number})")
        require_references = formula_config.get("require_references", True)

        # Ищем формулы в документе
        # В Python-docx формулы обычно хранятся как поля (нет встроенной поддержки MathML)
        # Поэтому ищем шаблоны: (1), (1.1), [1], и т.д.
        formula_patterns = self._find_formulas(document)

        if not formula_patterns:
            # Нет формул - проверка не требуется
            execution_time = time.time() - start_time
            return ValidationResult(
                validator_name=self.name, passed=True, issues=[], execution_time=execution_time
            )

        # Проверяем каждую формулу
        formula_numbers = []
        for formula_info in formula_patterns:
            # Проверяем нумерацию
            if require_numbering and not formula_info.get("has_number"):
                issues.append(
                    self._create_issue(
                        rule_id=17,
                        rule_name="Нумерация формул",
                        description="Формула не пронумерована",
                        severity=Severity.WARNING,
                        location=self._format_location(
                            paragraph_index=formula_info.get("paragraph_index"),
                            text_preview=formula_info.get("text")[:50],
                        ),
                        expected="Формула должна быть пронумерована",
                        actual="Номер не найден",
                        suggestion=f"Добавьте номер формулы в формате {numbering_format}",
                        can_autocorrect=False,
                    )
                )
            elif formula_info.get("number"):
                formula_numbers.append((formula_info.get("number"), formula_info))

        # Проверяем последовательность нумерации
        if formula_numbers:
            formula_numbers.sort(key=lambda x: x[0])
            for idx, (number, formula_info) in enumerate(formula_numbers, start=1):
                if number != idx:
                    issues.append(
                        self._create_issue(
                            rule_id=17,
                            rule_name="Последовательность нумерации формул",
                            description=(
                                "Нарушена последовательность нумерации "
                                "формул"),
                            severity=Severity.ERROR,
                            location=self._format_location(
                                paragraph_index=formula_info.get("paragraph_index")
                            ),
                            expected=f"Формула {idx}",
                            actual=f"Формула {number}",
                            suggestion=f"Перенумеруйте формулу на номер {idx}",
                            can_autocorrect=True,
                        )
                    )

        # Проверяем ссылки на формулы в тексте (если требуется)
        if require_references and formula_numbers:
            issues.extend(self._check_formula_references(document, len(formula_numbers)))

        # Проверяем интервалы и выравнивание
        issues.extend(self._check_formula_spacing(document, formula_patterns))

        execution_time = time.time() - start_time
        passed = len([i for i in issues if i.severity in [Severity.CRITICAL, Severity.ERROR]]) == 0

        return ValidationResult(
            validator_name=self.name, passed=passed, issues=issues, execution_time=execution_time
        )

    def _find_formulas(self, document: Any) -> List[Dict[str, Any]]:
        """
        Находит формулы в документе по стандартным шаблонам нумерации.

        Args:
            document: Document объект

        Returns:
            Список найденных формул
        """
        formulas = []

        # Паттерны нумерации формул
        patterns = [
            r"\((\d+(?:\.\d+)?)\)",  # (1), (1.1), (1.1.1)
            r"\[(\d+(?:\.\d+)?)\]",  # [1], [1.1]
            r"Формула\s+(\d+(?:\.\d+)?)",  # Формула 1, Формула 1.1
            r"формула\s+(\d+(?:\.\d+)?)",
        ]

        for idx, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()

            if not text:
                continue

            # Проверяем, содержит ли параграф математическое выражение или формулу
            # Ищем типичные символы формул
            if self._looks_like_formula(text):
                # Ищем номер формулы
                number = None
                for pattern in patterns:
                    match = re.search(pattern, text)
                    if match:
                        try:
                            number = int(
                                match.group(1).split(".")[0]
                            )  # Берем только первую часть числа
                        except (ValueError, IndexError):
                            pass
                        break

                formulas.append(
                    {
                        "text": text,
                        "paragraph_index": idx,
                        "number": number,
                        "has_number": number is not None,
                        "raw_text": text,
                    }
                )

        return formulas

    def _looks_like_formula(self, text: str) -> bool:
        """
        Проверяет, похожа ли строка на математическую формулу.

        Args:
            text: Текст для проверки

        Returns:
            True если это похоже на формулу
        """
        # Признаки формулы:
        # - Содержит математические символы
        # - Содержит греческие буквы
        # - Содержит математические операции
        # - Выглядит как уравнение

        formula_indicators = [
            r"[=≠<>≤≥]",  # Операторы сравнения
            r"[+−×÷/]",  # Арифметические операторы
            r"[αβγδεζηθικλμνξπρστυφχψω]",  # Греческие буквы в нижнем регистре
            r"[ΑΒΓΔΕΖΗΘΙΚΛΜΝΞΠΡΣΤΥΦΧΨΩ]",  # Греческие буквы в верхнем регистре
            r"∑|∏|√|∫|∂|∇",  # Математические символы
            r"\b(sin|cos|tan|log|ln|exp|dx|dy)\b",  # Популярные функции
            r"\b(где|при|если)|d[xyz]|Δ[xyz]",  # Условные выражения
            r"\((\d+(?:\.\d+)?)\)",  # Нумерованная формула
        ]

        for indicator in formula_indicators:
            if re.search(indicator, text, re.IGNORECASE):
                return True

        # Проверяем, является ли это простой переменной или выражением в скобках
        if re.match(r"^\s*\(.*\)\s*$", text) or re.match(r"^\s*\[.*\]\s*$", text):
            # Проверяем, что это не просто ссылка
            if not re.match(r"^\s*\([а-яА-ЯёЁ0-9\s,\.\-и]\)", text):
                return True

        return False

    def _check_formula_references(self, document: Any, formula_count: int) -> List[ValidationIssue]:
        """
        Проверяет наличие ссылок на формулы в тексте.

        Args:
            document: Document объект
            formula_count: Количество формул

        Returns:
            Список проблем
        """
        issues = []

        # Собираем весь текст документа
        full_text = " ".join([p.text for p in document.paragraphs])

        # Ищем ссылки на формулы
        reference_patterns = [
            r"формул[аеуы]\s+(?:\()?(\d+)(?:\))?",
            r"ф-л[а-яеёу]?\s+(?:\()?(\d+)(?:\))?",
            r"(?:\((?:см\.|из)\s+)?(?:формул[аеуы]\s+)?(?:\()?(\d+)(?:\))?(?:\s*\))?",
        ]

        referenced_formulas = set()

        for pattern in reference_patterns:
            matches = re.finditer(pattern, full_text, re.IGNORECASE)
            for match in matches:
                try:
                    number = int(match.group(1))
                    referenced_formulas.add(number)
                except (ValueError, IndexError):
                    pass
        for formula_num in range(1, formula_count + 1):
            if formula_num not in referenced_formulas:
                issues.append(
                    self._create_issue(
                        rule_id=18,
                        rule_name="Ссылки на формулы",
                        description=f"Не найдена ссылка на формулу {formula_num} в тексте",
                        severity=Severity.INFO,
                        location="Документ в целом",
                        expected=f"Ссылка на формулу {formula_num}",
                        actual="Ссылка не найдена",
                        suggestion=f"Добавьте ссылку на формулу ({formula_num}) в текст документа",
                        can_autocorrect=False,
                    )
                )

        return issues

    def _check_formula_spacing(
        self, document: Any, formulas: List[Dict[str, Any]]
    ) -> List[ValidationIssue]:
        """
        Проверяет интервалы до и после формул.

        Args:
            document: Document объект
            formulas: Список найденных формул

        Returns:
            Список проблем
        """
        issues = []

        # Требуемые интервалы из профиля
        space_before = self._get_rule_config(
            'formulas.space_before_pt', 6)
        expected_alignment = self._get_rule_config(
            'formulas.alignment', 'center')

        for formula_info in formulas:
            para_idx = formula_info["paragraph_index"]
            paragraph = document.paragraphs[para_idx]

            # Проверяем выравнивание (формулы обычно по центру)
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            if expected_alignment == "center" and paragraph.alignment not in [
                None,
                WD_PARAGRAPH_ALIGNMENT.CENTER,
            ]:
                issues.append(
                    self._create_issue(
                        rule_id=18,
                        rule_name="Выравнивание формулы",
                        description="Формула должна быть выравнена по центру",
                        severity=Severity.INFO,
                        location=self._format_location(
                            paragraph_index=para_idx, text_preview=formula_info["text"][:30]
                        ),
                        expected="По центру",
                        actual=self._get_alignment_name(paragraph.alignment),
                        suggestion="Выровняйте формулу по центру",
                        can_autocorrect=True,
                    )
                )

            # Проверяем интервалы
            if para_idx > 0:
                prev_paragraph = document.paragraphs[para_idx - 1]
                if prev_paragraph.paragraph_format.space_after:
                    actual_space = prev_paragraph.paragraph_format.space_after.pt
                    if abs(actual_space - space_before) > 2:
                        issues.append(
                            self._create_issue(
                                rule_id=18,
                                rule_name="Интервал перед формулой",
                                description="Неправильный интервал перед формулой",
                                severity=Severity.INFO,
                                location=self._format_location(paragraph_index=para_idx),
                                expected=f"{space_before}pt",
                                actual=f"{actual_space}pt",
                                can_autocorrect=True,
                            )
                        )

        return issues

    def _get_alignment_name(self, alignment) -> str:
        """Возвращает название выравнивания"""
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        if alignment is None:
            return "По умолчанию"

        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: "По левому краю",
            WD_PARAGRAPH_ALIGNMENT.CENTER: "По центру",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "По правому краю",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "По ширине",
        }

        return alignment_map.get(alignment, f"Неизвестно ({alignment})")
