"""
Валидатор для проверки оформления заголовков документа.
"""

from typing import Dict, Any, List
import time
import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from . import BaseValidator, ValidationResult, ValidationIssue, Severity


class HeadingValidator(BaseValidator):
    """
    Валидатор для проверки оформления заголовков в соответствии с требованиями.

    Проверяет:
    - Форматирование структурных элементов (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ и т.д.)
    - Форматирование разделов (ГЛАВА 1, 1. НАЗВАНИЕ)
    - Форматирование подразделов (1.1, 1.1.1)
    - Нумерацию заголовков
    - Отсутствие переносов, подчеркиваний
    - Интервалы до и после заголовков
    - Запрет точки в конце заголовка
    """

    # Структурные элементы (должны быть ЗАГЛАВНЫМИ)
    STRUCTURAL_ELEMENTS = [
        "введение",
        "заключение",
        "список литературы",
        "список использованных источников",
        "содержание",
        "оглавление",
        "реферат",
        "приложение",
        "приложения",
        "библиография",
        "глоссарий",
        "определения",
        "обозначения и сокращения",
    ]

    @property
    def name(self) -> str:
        return "HeadingValidator"

    def validate(self, document: Any, document_data: Dict[str, Any]) -> ValidationResult:
        """
        Проверяет оформление заголовков в документе.

        Args:
            document: python-docx Document объект
            document_data: Извлечённые данные документа

        Returns:
            ValidationResult с найденными проблемами
        """
        start_time = time.time()
        issues: List[ValidationIssue] = []

        # Получаем требования из профиля
        h1_config = self._get_rule_config("headings.h1", {})
        h2_config = self._get_rule_config("headings.h2", {})
        h3_config = self._get_rule_config("headings.h3", {})

        # Проверяем каждый параграф
        for idx, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()

            if not text:
                continue

            # Определяем тип заголовка
            heading_info = self._identify_heading(paragraph, text)

            if heading_info["type"] is None:
                continue  # Не заголовок

            # Проверяем форматирование в зависимости от типа
            if heading_info["type"] == "structural":
                issues.extend(self._check_structural_element(paragraph, text, idx, h1_config))
            elif heading_info["type"] == "chapter":
                issues.extend(self._check_chapter_heading(paragraph, text, idx, h2_config))
            elif heading_info["type"] == "section":
                issues.extend(self._check_section_heading(paragraph, text, idx, h3_config))

            # Общие проверки для всех заголовков
            issues.extend(self._check_common_heading_issues(paragraph, text, idx))

        # Проверка нумерации заголовков
        issues.extend(self._check_heading_numbering(document))

        # Проверка интервалов между заголовками и текстом
        issues.extend(self._check_heading_spacing(document))

        execution_time = time.time() - start_time
        passed = len([i for i in issues if i.severity in [Severity.CRITICAL, Severity.ERROR]]) == 0

        return ValidationResult(
            validator_name=self.name, passed=passed, issues=issues, execution_time=execution_time
        )

    def _identify_heading(self, paragraph: Any, text: str) -> Dict[str, Any]:
        """
        Определяет тип заголовка.

        Args:
            paragraph: Параграф документа
            text: Текст параграфа

        Returns:
            Словарь с информацией о типе заголовка
        """
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        text_lower = text.lower().rstrip(".")  # Убираем точку для сравнения

        # Проверка стиля
        if "heading" in style_name or "заголовок" in style_name:
            if "heading 1" in style_name or "заголовок 1" in style_name:
                return {"type": "structural", "level": 1}
            elif "heading 2" in style_name or "заголовок 2" in style_name:
                return {"type": "chapter", "level": 2}
            elif "heading 3" in style_name or "заголовок 3" in style_name:
                return {"type": "section", "level": 3}

        # Проверка по контенту: структурные элементы
        for element in self.STRUCTURAL_ELEMENTS:
            if text_lower == element or text_lower == element.upper():
                return {"type": "structural", "level": 1}

        # Проверка по нумерации: ГЛАВА 1 или 1. НАЗВАНИЕ
        if re.match(r"^(глава\s+\d+|раздел\s+\d+|\d+\.?\s+[А-ЯЁ])", text, re.IGNORECASE):
            return {"type": "chapter", "level": 2}

        # Проверка подразделов: 1.1, 1.1.1
        if re.match(r"^\d+\.\d+\.?\s+", text):
            return {"type": "section", "level": 3}

        return {"type": None, "level": 0}

    def _check_structural_element(
        self, paragraph: Any, text: str, idx: int, config: Dict[str, Any]
    ) -> List[ValidationIssue]:
        """
        Проверяет оформление структурного элемента (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ и т.д.).

        Args:
            paragraph: Параграф
            text: Текст
            idx: Индекс параграфа
            config: Конфигурация из профиля

        Returns:
            Список проблем
        """
        issues = []

        # Проверка: должно быть ЗАГЛАВНЫМИ
        expected_all_caps = config.get("all_caps", True)
        if expected_all_caps and not text.isupper():
            issues.append(
                self._create_issue(
                    rule_id=13,
                    rule_name="Заголовки структурных элементов",
                    description=f"Структурный элемент должен быть ЗАГЛАВНЫМИ буквами: '{text}'",
                    severity=Severity.ERROR,
                    location=self._format_location(paragraph_index=idx, text_preview=text),
                    expected="ЗАГЛАВНЫМИ БУКВАМИ",
                    actual=text,
                    suggestion="Преобразуйте текст в верхний регистр",
                    can_autocorrect=True,
                )
            )

        # Проверка выравнивания (обычно по центру)
        expected_alignment = config.get("alignment", WD_PARAGRAPH_ALIGNMENT.CENTER)
        if paragraph.alignment != expected_alignment:
            alignment_name = self._get_alignment_name(paragraph.alignment)
            expected_alignment_name = self._get_alignment_name(expected_alignment)

            issues.append(
                self._create_issue(
                    rule_id=13,
                    rule_name="Выравнивание структурных элементов",
                    description=f"Неправильное выравнивание структурного элемента",
                    severity=Severity.ERROR,
                    location=self._format_location(paragraph_index=idx, text_preview=text),
                    expected=expected_alignment_name,
                    actual=alignment_name,
                    suggestion=f"Выровняйте по центру",
                    can_autocorrect=True,
                )
            )

        # Проверка: не должно быть точки в конце
        if text.endswith("."):
            issues.append(
                self._create_issue(
                    rule_id=13,
                    rule_name="Точка в конце заголовка",
                    description=f"Структурный элемент не должен заканчиваться точкой",
                    severity=Severity.ERROR,
                    location=self._format_location(paragraph_index=idx, text_preview=text),
                    expected="Без точки в конце",
                    actual=text,
                    suggestion="Удалите точку в конце",
                    can_autocorrect=True,
                )
            )

        return issues

    def _check_chapter_heading(
        self, paragraph: Any, text: str, idx: int, config: Dict[str, Any]
    ) -> List[ValidationIssue]:
        """
        Проверяет оформление заголовка раздела (ГЛАВА 1, 1. НАЗВАНИЕ).

        Args:
            paragraph: Параграф
            text: Текст
            idx: Индекс параграфа
            config: Конфигурация

        Returns:
            Список проблем
        """
        issues = []

        # Разделы должны быть ЗАГЛАВНЫМИ
        expected_all_caps = config.get("all_caps", False)

        # Извлекаем номер и название
        match = re.match(r"^(глава\s+\d+|раздел\s+\d+|\d+\.?)\s+(.+)$", text, re.IGNORECASE)

        if match:
            number_part = match.group(1)
            title_part = match.group(2)

            if expected_all_caps and not title_part.isupper():
                issues.append(
                    self._create_issue(
                        rule_id=15,
                        rule_name="Оформление заголовков разделов",
                        description=f"Название раздела должно быть ЗАГЛАВНЫМИ: '{title_part}'",
                        severity=Severity.ERROR,
                        location=self._format_location(paragraph_index=idx, text_preview=text),
                        expected="ЗАГЛАВНЫМИ БУКВАМИ",
                        actual=title_part,
                        can_autocorrect=True,
                    )
                )

        # Проверка выравнивания (по ширине с отступом)
        expected_alignment = config.get("alignment", WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        if paragraph.alignment and paragraph.alignment != expected_alignment:
            issues.append(
                self._create_issue(
                    rule_id=15,
                    rule_name="Выравнивание разделов",
                    description="Неправильное выравнивание раздела",
                    severity=Severity.WARNING,
                    location=self._format_location(paragraph_index=idx, text_preview=text[:50]),
                    expected=self._get_alignment_name(expected_alignment),
                    actual=self._get_alignment_name(paragraph.alignment),
                    can_autocorrect=True,
                )
            )

        # Проверка: не должно быть точки в конце
        if text.endswith("."):
            issues.append(
                self._create_issue(
                    rule_id=15,
                    rule_name="Точка в конце заголовка",
                    description="Заголовок раздела не должен заканчиваться точкой",
                    severity=Severity.ERROR,
                    location=self._format_location(paragraph_index=idx),
                    can_autocorrect=True,
                )
            )

        return issues

    def _check_section_heading(
        self, paragraph: Any, text: str, idx: int, config: Dict[str, Any]
    ) -> List[ValidationIssue]:
        """
        Проверяет оформление подраздела (1.1, 1.1.1).

        Args:
            paragraph: Параграф
            text: Текст
            idx: Индекс
            config: Конфигурация

        Returns:
            Список проблем
        """
        issues = []

        # Подразделы: первая буква заглавная, остальные строчные
        match = re.match(r"^(\d+\.\d+(?:\.\d+)?\.?)\s+(.+)$", text)

        if match:
            number_part = match.group(1)
            title_part = match.group(2)

            # Проверка: первая буква заглавная
            if title_part and not title_part[0].isupper():
                issues.append(
                    self._create_issue(
                        rule_id=15,
                        rule_name="Оформление подразделов",
                        description="Название подраздела должно начинаться с заглавной буквы",
                        severity=Severity.ERROR,
                        location=self._format_location(paragraph_index=idx, text_preview=text),
                        expected="С заглавной буквы",
                        actual=title_part,
                        can_autocorrect=True,
                    )
                )

            # Не должно быть полностью в верхнем регистре
            if title_part.isupper() and len(title_part) > 1:
                issues.append(
                    self._create_issue(
                        rule_id=15,
                        rule_name="Оформление подразделов",
                        description="Название подраздела не должно быть полностью ЗАГЛАВНЫМИ",
                        severity=Severity.WARNING,
                        location=self._format_location(paragraph_index=idx),
                        actual=title_part,
                        can_autocorrect=True,
                    )
                )

        # Проверка: не должно быть точки в конце (кроме точки после номера)
        if text.rstrip().endswith(".") and not re.match(r".*\d+\.$", text.rstrip()):
            issues.append(
                self._create_issue(
                    rule_id=15,
                    rule_name="Точка в конце заголовка",
                    description="Заголовок подраздела не должен заканчиваться точкой",
                    severity=Severity.ERROR,
                    location=self._format_location(paragraph_index=idx),
                    can_autocorrect=True,
                )
            )

        return issues

    def _check_common_heading_issues(
        self, paragraph: Any, text: str, idx: int
    ) -> List[ValidationIssue]:
        """
        Проверяет общие проблемы заголовков.

        Args:
            paragraph: Параграф
            text: Текст
            idx: Индекс

        Returns:
            Список проблем
        """
        issues = []

        # Проверка на переносы слов (дефисы в конце строк)
        if "-\n" in paragraph.text or "- " in text:
            issues.append(
                self._create_issue(
                    rule_id=16,
                    rule_name="Переносы в заголовках",
                    description="В заголовках не допускаются переносы слов",
                    severity=Severity.ERROR,
                    location=self._format_location(paragraph_index=idx, text_preview=text[:50]),
                    suggestion="Удалите переносы слов в заголовке",
                    can_autocorrect=True,
                )
            )

        # Проверка на подчеркивание
        for run in paragraph.runs:
            if run.underline:
                issues.append(
                    self._create_issue(
                        rule_id=16,
                        rule_name="Подчеркивание в заголовках",
                        description="Заголовки не должны быть подчеркнуты",
                        severity=Severity.WARNING,
                        location=self._format_location(
                            paragraph_index=idx, text_preview=run.text[:30]
                        ),
                        suggestion="Удалите подчеркивание",
                        can_autocorrect=True,
                    )
                )
                break

        # Проверка на разрядку (избыточные пробелы между буквами)
        if re.search(r"[А-ЯЁа-яё]\s{2,}[А-ЯЁа-яё]", text):
            issues.append(
                self._create_issue(
                    rule_id=16,
                    rule_name="Разрядка в заголовках",
                    description="В заголовках не допускается разрядка (лишние пробелы между буквами)",
                    severity=Severity.WARNING,
                    location=self._format_location(paragraph_index=idx),
                    suggestion="Удалите лишние пробелы",
                    can_autocorrect=True,
                )
            )

        return issues

    def _check_heading_numbering(self, document: Any) -> List[ValidationIssue]:
        """
        Проверяет правильность нумерации заголовков.

        Args:
            document: Document объект

        Returns:
            Список проблем
        """
        issues = []

        # Собираем все пронумерованные заголовки
        chapter_numbers = []
        section_numbers = {}

        for idx, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()

            # Главы/разделы: ГЛАВА 1, 1. НАЗВАНИЕ
            chapter_match = re.match(r"^(?:глава|раздел)\s+(\d+)", text, re.IGNORECASE)
            if not chapter_match:
                chapter_match = re.match(r"^(\d+)\.?\s+[А-ЯЁ]", text)

            if chapter_match:
                num = int(chapter_match.group(1))
                chapter_numbers.append((num, idx, text))

            # Подразделы: 1.1, 1.1.1
            section_match = re.match(r"^(\d+)\.(\d+)(?:\.(\d+))?", text)
            if section_match:
                chapter_num = int(section_match.group(1))
                section_num = int(section_match.group(2))
                subsection_num = int(section_match.group(3)) if section_match.group(3) else None

                if chapter_num not in section_numbers:
                    section_numbers[chapter_num] = []
                section_numbers[chapter_num].append((section_num, subsection_num, idx, text))

        # Проверка последовательности глав
        for i, (num, idx, text) in enumerate(chapter_numbers):
            expected = i + 1
            if num != expected:
                issues.append(
                    self._create_issue(
                        rule_id=15,
                        rule_name="Нумерация разделов",
                        description=f"Неправильная нумерация раздела: {num} вместо {expected}",
                        severity=Severity.ERROR,
                        location=self._format_location(paragraph_index=idx, text_preview=text[:50]),
                        expected=str(expected),
                        actual=str(num),
                        can_autocorrect=True,
                    )
                )

        # Проверка последовательности подразделов
        for chapter_num, sections in section_numbers.items():
            sections.sort(key=lambda x: (x[0], x[1] if x[1] else 0))

            expected_section = 1
            for section_num, subsection_num, idx, text in sections:
                if subsection_num is None:  # Это раздел типа 1.1
                    if section_num != expected_section:
                        issues.append(
                            self._create_issue(
                                rule_id=15,
                                rule_name="Нумерация подразделов",
                                description=f"Неправильная нумерация подраздела: {chapter_num}.{section_num}",
                                severity=Severity.ERROR,
                                location=self._format_location(paragraph_index=idx),
                                expected=f"{chapter_num}.{expected_section}",
                                actual=f"{chapter_num}.{section_num}",
                                can_autocorrect=True,
                            )
                        )
                    expected_section += 1

        return issues

    def _check_heading_spacing(self, document: Any) -> List[ValidationIssue]:
        """
        Проверяет интервалы до и после заголовков.

        Args:
            document: Document объект

        Returns:
            Список проблем
        """
        issues = []

        # Требуемые интервалы из профиля
        space_before_h1 = self._get_rule_config("headings.h1.space_before", 0)
        space_after_h1 = self._get_rule_config("headings.h1.space_after", 12)
        space_before_h2 = self._get_rule_config("headings.h2.space_before", 12)
        space_after_h2 = self._get_rule_config("headings.h2.space_after", 12)

        tolerance = 2  # Допуск в пунктах

        for idx, paragraph in enumerate(document.paragraphs):
            heading_info = self._identify_heading(paragraph, paragraph.text.strip())

            if heading_info["type"] is not None:
                # Проверяем интервалы
                actual_space_before = paragraph.paragraph_format.space_before
                actual_space_after = paragraph.paragraph_format.space_after

                if heading_info["level"] == 1:
                    expected_before = space_before_h1
                    expected_after = space_after_h1
                elif heading_info["level"] == 2:
                    expected_before = space_before_h2
                    expected_after = space_after_h2
                else:
                    continue

                # Проверка интервала до
                if (
                    actual_space_before
                    and abs(actual_space_before.pt - expected_before) > tolerance
                ):
                    issues.append(
                        self._create_issue(
                            rule_id=11,
                            rule_name="Интервалы до заголовка",
                            description=f"Неправильный интервал до заголовка",
                            severity=Severity.WARNING,
                            location=self._format_location(
                                paragraph_index=idx, text_preview=paragraph.text.strip()[:30]
                            ),
                            expected=f"{expected_before}pt",
                            actual=f"{actual_space_before.pt}pt" if actual_space_before else "0pt",
                            can_autocorrect=True,
                        )
                    )

                # Проверка интервала после
                if actual_space_after and abs(actual_space_after.pt - expected_after) > tolerance:
                    issues.append(
                        self._create_issue(
                            rule_id=11,
                            rule_name="Интервалы после заголовка",
                            description=f"Неправильный интервал после заголовка",
                            severity=Severity.WARNING,
                            location=self._format_location(
                                paragraph_index=idx, text_preview=paragraph.text.strip()[:30]
                            ),
                            expected=f"{expected_after}pt",
                            actual=f"{actual_space_after.pt}pt" if actual_space_after else "0pt",
                            can_autocorrect=True,
                        )
                    )

        return issues

    def _get_alignment_name(self, alignment) -> str:
        """Возвращает название выравнивания"""
        if alignment is None:
            return "Не задано"

        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: "По левому краю",
            WD_PARAGRAPH_ALIGNMENT.CENTER: "По центру",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "По правому краю",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "По ширине",
        }

        return alignment_map.get(alignment, f"Неизвестно ({alignment})")
