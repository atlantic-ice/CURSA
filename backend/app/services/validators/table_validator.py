"""
Валидатор для проверки оформления таблиц в документе.
"""

from typing import Dict, Any, List
import time
from . import BaseValidator, ValidationResult, ValidationIssue, Severity


class TableValidator(BaseValidator):
    """
    Валидатор для проверки оформления таблиц.

    Проверяет:
    - Наличие подписи таблицы
    - Форматирование подписи (Таблица X - Название)
    - Нумерацию таблиц
    - Выравнивание таблиц
    - Оформление ячеек таблицы
    - Размер шрифта в таблицах
    - Последовательность ссылок на таблицы в тексте
    """

    @property
    def name(self) -> str:
        return "TableValidator"

    def validate(self, document: Any, document_data: Dict[str, Any]) -> ValidationResult:
        """
        Проверяет оформление таблиц в документе.

        Args:
            document: python-docx Document объект
            document_data: Извлечённые данные документа

        Returns:
            ValidationResult с найденными проблемами
        """
        start_time = time.time()
        issues: List[ValidationIssue] = []

        # Получаем требования из профиля
        table_config = self._get_rule_config("tables", {})

        require_caption = table_config.get("require_caption", True)
        caption_format = table_config.get("caption_format", "Таблица {number} - {title}")
        min_font_size = table_config.get("min_font_size_pt", 10)
        max_font_size = table_config.get("max_font_size_pt", 14)

        # Получаем все таблицы
        tables = list(document.tables)

        if not tables:
            # Нет таблиц - проверка не требуется
            execution_time = time.time() - start_time
            return ValidationResult(
                validator_name=self.name, passed=True, issues=[], execution_time=execution_time
            )

        # Проверяем каждую таблицу
        for table_idx, table in enumerate(tables, start=1):
            # Ищем подпись таблицы (обычно перед или после таблицы)
            caption_info = self._find_table_caption(document, table, table_idx)

            if require_caption and not caption_info["found"]:
                issues.append(
                    self._create_issue(
                        rule_id=21,
                        rule_name="Подпись таблицы",
                        description=f"Не найдена подпись для таблицы {table_idx}",
                        severity=Severity.ERROR,
                        location=f"Таблица {table_idx}",
                        expected="Наличие подписи таблицы",
                        actual="Подпись не найдена",
                        suggestion=f"Добавьте подпись в формате: {caption_format.format(number=table_idx, title='Название')}",
                        can_autocorrect=False,
                    )
                )
            elif caption_info["found"]:
                # Проверяем формат подписи
                issues.extend(self._check_caption_format(caption_info, table_idx, caption_format))

            # Проверяем оформление ячеек таблицы
            issues.extend(self._check_table_cells(table, table_idx, min_font_size, max_font_size))

            # Проверяем выравнивание таблицы
            issues.extend(self._check_table_alignment(table, table_idx, table_config))

        # Проверяем последовательность нумерации таблиц
        issues.extend(self._check_table_numbering(document, tables))

        # Проверяем ссылки на таблицы в тексте
        issues.extend(self._check_table_references(document, len(tables)))

        execution_time = time.time() - start_time
        passed = len([i for i in issues if i.severity in [Severity.CRITICAL, Severity.ERROR]]) == 0

        return ValidationResult(
            validator_name=self.name, passed=passed, issues=issues, execution_time=execution_time
        )

    def _find_table_caption(self, document: Any, table: Any, table_number: int) -> Dict[str, Any]:
        """
        Ищет подпись таблицы (обычно параграф перед таблицей).

        Args:
            document: Document объект
            table: Таблица
            table_number: Номер таблицы

        Returns:
            Информация о подписи
        """
        # Таблицы и параграфы перемешаны в document
        # Нужно найти позицию таблицы и проверить соседние параграфы

        # Получаем все элементы документа
        body_elements = document.element.body
        table_element = table._element

        # Ищем индекс таблицы
        table_index = None
        for idx, element in enumerate(body_elements):
            if element == table_element:
                table_index = idx
                break

        if table_index is None:
            return {"found": False}

        # Проверяем параграф перед таблицей
        if table_index > 0:
            prev_element = body_elements[table_index - 1]
            if prev_element.tag.endswith("p"):  # Это параграф
                # Находим соответствующий параграф в документе
                for paragraph in document.paragraphs:
                    if paragraph._element == prev_element:
                        text = paragraph.text.strip()
                        text_lower = text.lower()

                        # Проверяем, является ли это подписью таблицы
                        if "таблица" in text_lower:
                            return {
                                "found": True,
                                "text": text,
                                "paragraph": paragraph,
                                "position": "before",
                            }

        # Проверяем параграф после таблицы
        if table_index < len(body_elements) - 1:
            next_element = body_elements[table_index + 1]
            if next_element.tag.endswith("p"):
                for paragraph in document.paragraphs:
                    if paragraph._element == next_element:
                        text = paragraph.text.strip()
                        text_lower = text.lower()

                        if "таблица" in text_lower:
                            return {
                                "found": True,
                                "text": text,
                                "paragraph": paragraph,
                                "position": "after",
                            }

        return {"found": False}

    def _check_caption_format(
        self, caption_info: Dict[str, Any], table_number: int, expected_format: str
    ) -> List[ValidationIssue]:
        """
        Проверяет формат подписи таблицы.

        Args:
            caption_info: Информация о подписи
            table_number: Номер таблицы
            expected_format: Ожидаемый формат

        Returns:
            Список проблем
        """
        issues = []

        text = caption_info["text"]
        paragraph = caption_info["paragraph"]

        # Проверяем формат: Таблица X - Название или Таблица X. Название
        import re

        # Паттерны для различных форматов
        patterns = [
            r"таблица\s+(\d+)\s*[-–—]\s*(.+)",  # Таблица 1 - Название
            r"таблица\s+(\d+)\.\s*(.+)",  # Таблица 1. Название
            r"таблица\s+(\d+)\s+(.+)",  # Таблица 1 Название
        ]

        match = None
        for pattern in patterns:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                break

        if not match:
            issues.append(
                self._create_issue(
                    rule_id=21,
                    rule_name="Формат подписи таблицы",
                    description="Неправильный формат подписи",
                    severity=Severity.WARNING,
                    location=f"Подпись таблицы {table_number}",
                    expected=expected_format.format(number=table_number, title="Название"),
                    actual=text,
                    suggestion="Используйте формат: Таблица X",
                    can_autocorrect=True,
                )
            )
            # Проверяем номер таблицы
            found_number = int(match.group(1))
            if found_number != table_number:
                issues.append(
                    self._create_issue(
                        rule_id=21,
                        rule_name="Нумерация таблицы",
                        description="Неправильный номер таблицы",
                        severity=Severity.ERROR,
                        location=f"Подпись таблицы {table_number}",
                        expected=f"Таблица {table_number}",
                        actual=f"Таблица {found_number}",
                        suggestion=f"Исправьте номер на {table_number}",
                        can_autocorrect=True,
                    )
                )

            # Проверяем, что название не пустое
            title = match.group(2).strip()
            if not title:
                issues.append(
                    self._create_issue(
                        rule_id=21,
                        rule_name="Название таблицы",
                        description="Отсутствует название таблицы",
                        severity=Severity.WARNING,
                        location=f"Подпись таблицы {table_number}",
                        expected="Таблица с названием",
                        actual=text,
                        suggestion="Добавьте описательное название",
                        can_autocorrect=False,
                    )
                )

        # Проверяем выравнивание подписи (обычно по левому краю или по центру)
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        if paragraph.alignment not in [
            None,
            WD_PARAGRAPH_ALIGNMENT.LEFT,
            WD_PARAGRAPH_ALIGNMENT.CENTER,
        ]:
            issues.append(
                self._create_issue(
                    rule_id=21,
                    rule_name="Выравнивание подписи таблицы",
                    description="Неправильное выравнивание подписи таблицы",
                    severity=Severity.INFO,
                    location=f"Подпись таблицы {table_number}",
                    expected="По левому краю или по центру",
                    actual=self._get_alignment_name(paragraph.alignment),
                    can_autocorrect=True,
                )
            )

        return issues

    def _check_table_cells(
        self, table: Any, table_number: int, min_font_size: int, max_font_size: int
    ) -> List[ValidationIssue]:
        """
        Проверяет оформление ячеек таблицы.

        Args:
            table: Таблица
            table_number: Номер таблицы
            min_font_size: Минимальный размер шрифта
            max_font_size: Максимальный размер шрифта

        Returns:
            Список проблем
        """
        issues = []

        # Проверяем каждую ячейку
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.size:
                            font_size_pt = run.font.size.pt

                            if font_size_pt < min_font_size:
                                issues.append(
                                    self._create_issue(
                                        rule_id=19,
                                        rule_name="Размер шрифта в таблице",
                                        description="Поривающий размер шрифта",
                                        severity=Severity.WARNING,
                                        location=(
                                            f"Таблица {table_number}, "
                                            f"ячейка ({row_idx+1}, {col_idx+1})"
                                        ),
                                        expected=f"≥{min_font_size}pt",
                                        actual=f"{font_size_pt}pt",
                                        suggestion=(f"Увеличьте но " f"{min_font_size}pt"),
                                        can_autocorrect=True,
                                    )
                                )
                            elif font_size_pt > max_font_size:
                                issues.append(
                                    self._create_issue(
                                        rule_id=19,
                                        rule_name="Размер шрифта в таблице",
                                        description="Превышающий размер шрифта",
                                        severity=Severity.INFO,
                                        location=(
                                            f"Таблица {table_number}, "
                                            f"ячейка ({row_idx+1}, {col_idx+1})"
                                        ),
                                        expected=f"≤{max_font_size}pt",
                                        actual=f"{font_size_pt}pt",
                                        suggestion=(f"Уменьшите но " f"{max_font_size}pt"),
                                        can_autocorrect=True,
                                    )
                                )

        return issues

    def _check_table_alignment(
        self, table: Any, table_number: int, config: Dict[str, Any]
    ) -> List[ValidationIssue]:
        """
        Проверяет выравнивание таблицы.

        Args:
            table: Таблица
            table_number: Номер таблицы
            config: Конфигурация

        Returns:
            Список проблем
        """
        issues = []

        # В python-docx сложно получить выравнивание таблицы
        # Это более продвинутая проверка, которую можно добавить позже

        return issues

    def _check_table_numbering(self, document: Any, tables: List[Any]) -> List[ValidationIssue]:
        """
        Проверяет последовательность нумерации таблиц.

        Args:
            document: Document объект
            tables: Список таблиц

        Returns:
            Список проблем
        """
        issues = []

        # Собираем номера таблиц из подписей
        import re

        table_numbers = []

        for table_idx, table in enumerate(tables, start=1):
            caption_info = self._find_table_caption(document, table, table_idx)

            if caption_info["found"]:
                text = caption_info["text"]
                match = re.search(r"таблица\s+(\d+)", text, re.IGNORECASE)

                if match:
                    number = int(match.group(1))
                    table_numbers.append((number, table_idx))

        # Проверяем последовательность
        for idx, (number, table_idx) in enumerate(table_numbers):
            expected = idx + 1

            if number != expected:
                issues.append(
                    self._create_issue(
                        rule_id=21,
                        rule_name="Последовательность нумерации",
                        description="Нарушена последовательность",
                        severity=Severity.ERROR,
                        location=f"Таблица {table_idx}",
                        expected=f"Таблица {expected}",
                        actual=f"Таблица {number}",
                        suggestion=f"Перенумеруйте таблицу на номер {expected}",
                        can_autocorrect=True,
                    )
                )

        return issues

    def _check_table_references(self, document: Any, table_count: int) -> List[ValidationIssue]:
        """
        Проверяет наличие ссылок на таблицы в тексте.

        Args:
            document: Document объект
            table_count: Количество таблиц

        Returns:
            Список проблем
        """
        issues = []

        import re

        # Собираем весь текст документа (кроме таблиц)
        full_text = " ".join([p.text for p in document.paragraphs])

        # Ищем ссылки на таблицы
        reference_patterns = [
            r"таблиц[аеуы]\s+(\d+)",
            r"табл\.\s*(\d+)",
            r"см\.\s*таблиц[уа]\s+(\d+)",
        ]

        referenced_tables = set()

        for pattern in reference_patterns:
            matches = re.finditer(pattern, full_text, re.IGNORECASE)
            for match in matches:
                number = int(match.group(1))
                referenced_tables.add(number)

        # Проверяем, что на все таблицы есть ссылки
        for table_num in range(1, table_count + 1):
            if table_num not in referenced_tables:
                issues.append(
                    self._create_issue(
                        rule_id=21,
                        rule_name="Ссылки на таблицы",
                        description=f"Не найдена ссылка на таблицу {table_num} в тексте",
                        severity=Severity.INFO,
                        location=f"Таблица {table_num}",
                        expected=f"Ссылка на таблицу в тексте",
                        actual="Ссылка не найдена",
                        suggestion=f"Добавьте ссылку на таблицу {table_num} в текст документа",
                        can_autocorrect=False,
                    )
                )

        return issues

    def _get_alignment_name(self, alignment) -> str:
        """Возвращает название выравнивания"""
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        if alignment is None:
            return "По умолчанию"

        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: "По левому краю",
            WD_PARAGRAPH_ALIGNMENT.CENTER: "По центру",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "По правому краю",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "По ширине",
        }

        return alignment_map.get(alignment, f"Неизвестно ({alignment})")
