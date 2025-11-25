import os
import re
import datetime
import tempfile
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import shutil
from docxtpl import DocxTemplate
from docxcompose.composer import Composer

class DocumentCorrector:
    """
    Класс для исправления ошибок в документе
    """
    def __init__(self, profile_data=None):
        # Если профиль не передан, пытаемся загрузить стандартный
        if profile_data is None:
            try:
                base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                profile_path = os.path.join(base_dir, 'profiles', 'default_gost.json')
                if os.path.exists(profile_path):
                    with open(profile_path, 'r', encoding='utf-8') as f:
                        profile_data = json.load(f)
            except Exception as e:
                print(f"Ошибка при загрузке стандартного профиля: {e}")

        # Если все еще нет профиля, используем жестко заданный минимум (fallback)
        if profile_data is None:
            self.profile = {
                'rules': {
                    'font': {'name': 'Times New Roman', 'size': 14.0},
                    'margins': {'left': 3.0, 'right': 1.0, 'top': 2.0, 'bottom': 2.0},
                    'line_spacing': 1.5,
                    'first_line_indent': 1.25,
                    'headings': {
                        'h1': {'font_size': 14.0, 'bold': True, 'alignment': 'CENTER', 'all_caps': True},
                        'h2': {'font_size': 14.0, 'bold': True, 'alignment': 'LEFT'}
                    }
                }
            }
        else:
            self.profile = profile_data

        # Для удобства доступа к правилам
        self.rules = self.profile.get('rules', {})
        
        self.errors = []
        self.temp_files = []
    
    def __del__(self):
        """
        Деструктор для очистки временных файлов
        """
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
                    print(f"Удален временный файл: {temp_file}")
                    
                # Проверяем, пуста ли директория, и если да, удаляем её
                temp_dir = os.path.dirname(temp_file)
                if os.path.exists(temp_dir) and not os.listdir(temp_dir):
                    os.rmdir(temp_dir)
                    print(f"Удалена пустая временная директория: {temp_dir}")
            except Exception as e:
                print(f"Ошибка при удалении временного файла {temp_file}: {str(e)}")
    
    def correct_document(self, file_path, errors=None, out_path=None):
        """
        Исправляет ошибки в документе
        
        Args:
            file_path: Путь к файлу для исправления
            errors: Список ошибок для исправления (если None, исправляем все возможные)
            out_path: Путь для сохранения исправленного файла (если None, генерируется автоматически)
            
        Returns:
            str: Путь к исправленному файлу
        """
        self.errors = errors
        
        # Проверяем существование файла
        if not os.path.exists(file_path):
            # Пробуем добавить расширение .docx если его нет
            if not file_path.lower().endswith('.docx'):
                corrected_path = file_path + '.docx'
                if os.path.exists(corrected_path):
                    file_path = corrected_path
                else:
                    raise FileNotFoundError(f"Файл не найден: {file_path}")
            else:
                raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        try:
            # Загружаем документ
            document = Document(file_path)
            
            # Если указан путь для сохранения
            if out_path:
                # Проверяем, что путь имеет правильное расширение
                if not out_path.lower().endswith('.docx'):
                    out_path = out_path + '.docx'
                
                # Создаем директорию, если ее нет
                out_dir = os.path.dirname(out_path)
                if out_dir and not os.path.exists(out_dir):
                    os.makedirs(out_dir, exist_ok=True)
            else:
                # Если путь не указан, создаем временный файл
                temp_dir = tempfile.mkdtemp()
                file_name = os.path.basename(file_path)
                
                # Убедимся, что имя имеет расширение .docx
                if not file_name.lower().endswith('.docx'):
                    file_name = os.path.splitext(file_name)[0] + '.docx'
                
                out_path = os.path.join(temp_dir, f"corrected_{file_name}")
                self.temp_files.append(out_path)
            
            # Если список ошибок не предоставлен, исправляем все, что можем
            if errors is None:
                # Применяем базовые стили перед точечными корректировками, чтобы документ выглядел системно
                self._apply_core_styles(document)
                self._correct_all(document)
            else:
                # Исправляем только указанные ошибки
                self._correct_specific_errors(document, errors)
            
            # Сохраняем исправленный документ
            document.save(out_path)
            return out_path
            
        except Exception as e:
            print(f"Ошибка при исправлении документа: {str(e)}")
            raise
    
    def _correct_all(self, document):
        """
        Исправляет все типичные ошибки в документе
        """
        # Сначала исправляем поля страницы и базовые настройки документа
        self._correct_margins(document)
        
        # Исправляем шрифт для всего документа
        self._correct_font(document)
        
        # Исправляем межстрочный интервал
        self._correct_line_spacing(document)

        # Продвигаем псевдозаголовки (обычный текст, похожий на заголовок) в корректные стили Heading
        self._promote_pseudo_headings_to_styles(document)
        
        # Исправляем заголовки разделов (улучшенная версия)
        self._correct_section_headings(document)
        
        # Исправляем оформление таблиц
        self._correct_tables(document)
        
        # Исправляем подписи к рисункам (точки в конце)
        self._correct_images(document)
        
        # Исправляем оформление формул
        self._correct_formulas(document)
        
        # Исправляем оформление списков
        self._correct_lists(document)
        
        # Исправляем библиографические ссылки
        self._correct_bibliography_references(document)
        
        # Исправляем список литературы по ГОСТу
        self._correct_gost_bibliography(document)
        
        # Исправляем оглавление
        self._correct_toc(document)
        
        # Исправляем список сокращений и условных обозначений
        self._correct_abbreviations_list(document)
        
        # Исправляем перекрестные ссылки
        self._correct_cross_references(document)
        
        # Исправляем оформление приложений
        self._correct_appendices(document)
        
        # Исправляем оформление акцентов в тексте
        self._correct_text_accents(document)
        
        # Исправляем подстрочные ссылки
        self._correct_footnotes(document)
        
        # Исправляем нумерацию страниц
        self._correct_page_numbers(document)
        
        # ОТКЛЮЧЕНО: Исправление титульного листа (удаляло весь контент)
        # self._correct_title_page(document)
        
        # Исправляем переносы в тексте
        self._correct_hyphenation(document)
        
        # В конце применяем форматирование абзацев и выравнивание
        # для гарантии правильного форматирования всего текста
        self._correct_first_line_indent(document)
        self._correct_paragraph_alignment(document)
        self._clean_extra_blank_lines(document)
    
    def _apply_core_styles(self, document):
        """Подстраивает ключевые стили Word под нормоконтрольный стандарт."""
        try:
            normal_style = document.styles['Normal']
            self._set_style_font_defaults(normal_style, self.rules['font']['size'])
            normal_style.paragraph_format.first_line_indent = Cm(1.25)
            normal_style.paragraph_format.left_indent = Cm(0)
            normal_style.paragraph_format.right_indent = Cm(0)
            normal_style.paragraph_format.space_before = Pt(0)
            normal_style.paragraph_format.space_after = Pt(0)
            normal_style.paragraph_format.line_spacing = self.rules['line_spacing']
            normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            normal_style.paragraph_format.keep_together = False
            normal_style.paragraph_format.keep_with_next = False
        except KeyError:
            pass

        # Нормоконтрольные интервалы для заголовков:
        # H1: перед 36 пт (≈3 одинарных), после 24 пт (≈2 одинарных)
        # H2: перед 24 пт (≈2 одинарных), после 24 пт (≈2 одинарных)
        heading_defaults = [
            ('Heading 1', self.rules['headings'].get('h1', {}), Pt(36), Pt(24)),
            ('Heading 2', self.rules['headings'].get('h2', {}), Pt(24), Pt(24)),
        ]

        for style_name, rule, space_before, space_after in heading_defaults:
            try:
                heading_style = document.styles[style_name]
            except KeyError:
                continue

            font_size = rule.get('font_size', self.rules['font']['size'])
            bold = rule.get('bold', True)
            all_caps = rule.get('all_caps', False)
            alignment = rule.get('alignment', WD_PARAGRAPH_ALIGNMENT.LEFT)

            self._set_style_font_defaults(heading_style, font_size, bold=bold, all_caps=all_caps)

            heading_style.paragraph_format.alignment = alignment
            heading_style.paragraph_format.first_line_indent = Cm(0)
            heading_style.paragraph_format.left_indent = Cm(0)
            heading_style.paragraph_format.right_indent = Cm(0)
            heading_style.paragraph_format.space_before = space_before
            heading_style.paragraph_format.space_after = space_after
            heading_style.paragraph_format.keep_with_next = True
            heading_style.paragraph_format.keep_together = True
            # Полуторный интервал и для заголовков (требование общего межстрочного интервала)
            heading_style.paragraph_format.line_spacing = self.rules['line_spacing']
            heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    def _set_style_font_defaults(self, style, font_size, *, bold=False, all_caps=False):
        """Применяет единый шрифт Times New Roman к стилю."""
        font_name = self.rules['font']['name']
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.bold = bold
        style.font.all_caps = all_caps

        rPr = style._element.get_or_add_rPr()
        rPr.rFonts_ascii = font_name
        rPr.rFonts_hAnsi = font_name
        rPr.rFonts_eastAsia = font_name
        rPr.rFonts_cs = font_name

    def _clean_extra_blank_lines(self, document):
        """Удаляет лишние подряд идущие пустые абзацы, сохраняя визуальную чистоту.
        ВАЖНО: НЕ трогает таблицы - они требуют особой осторожности!
        """
        def trim_paragraph_list(paragraphs, in_table=False):
            consecutive_blank = 0
            for paragraph in list(paragraphs):
                # КРИТИЧЕСКИ ВАЖНО: в таблицах НЕ удаляем пустые параграфы!
                # Они могут быть частью структуры ячеек
                if in_table:
                    continue
                    
                if paragraph.text.strip():
                    consecutive_blank = 0
                    continue

                consecutive_blank += 1
                # Удаляем только если больше 2 подряд пустых параграфов
                if consecutive_blank > 2:
                    try:
                        element = paragraph._element
                        parent = element.getparent()
                        if parent is not None:
                            parent.remove(element)
                    except Exception as e:
                        print(f"Предупреждение: не удалось удалить пустой параграф: {str(e)}")
                        continue

        # Обрабатываем только основные параграфы документа
        trim_paragraph_list(document.paragraphs, in_table=False)

        # Таблицы НЕ трогаем - они хрупкие!
        # Комментируем опасный код:
        # for table in document.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             trim_paragraph_list(cell.paragraphs, in_table=True)

    def _promote_pseudo_headings_to_styles(self, document):
        """Находит параграфы, оформленные как заголовки вручную, и присваивает им Heading 1/2.

        Эвристики:
        - Строки вида: "ГЛАВА 1 ...", "РАЗДЕЛ 2 ..." -> Heading 1
        - Нумерация вида: "1. ..." (верхний уровень) -> Heading 1
        - Нумерация вида: "1.1 ..." (второй уровень и глубже) -> Heading 2
        - ALL-CAPS короткие слова (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ, СПИСОК ЛИТЕРАТУРЫ) -> Heading 1
        
        ВАЖНО: НЕ использует paragraph.text = ... для сохранения форматирования
        """
        try:
            h1_style = None
            h2_style = None
            try:
                h1_style = document.styles['Heading 1']
            except KeyError:
                pass
            try:
                h2_style = document.styles['Heading 2']
            except KeyError:
                pass

            if not (h1_style or h2_style):
                return

            h1_keywords = {
                'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ', 'СПИСОК ЛИТЕРАТУРЫ', 'ПРИЛОЖЕНИЯ', 'ПРИЛОЖЕНИЕ'
            }

            # Получаем список всех параграфов внутри таблиц для исключения
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))

            for paragraph in document.paragraphs:
                # КРИТИЧЕСКАЯ ПРОВЕРКА: пропускаем параграфы внутри таблиц
                if id(paragraph) in table_paragraphs:
                    continue
                
                text = paragraph.text.strip()
                if not text:
                    continue
                # Уже нормальный заголовок — пропускаем
                if paragraph.style and paragraph.style.name.startswith('Heading'):
                    continue

                try:
                    lower = text.lower()

                    # ГЛАВА N / РАЗДЕЛ N (только верхний уровень)
                    if re.match(r'^(глава|раздел)\s+\d+\.?\s*', lower):
                        if h1_style:
                            paragraph.style = h1_style
                        continue

                    # Многоуровневая нумерация: 1., 1.1, 1.1.1 и т.д.
                    # Определяем уровень по количеству чисел в начале
                    m = re.match(r'^(\d+(?:\.\d+)*)(?:[\s\.:\-]+)(.*)', text)
                    if m:
                        numbering = m.group(1)
                        rest = m.group(2).strip()
                        level = numbering.count('.') + 1

                        # Выбираем стиль по уровню (1->Heading 1, 2->Heading 2, >=3->Heading 3)
                        try:
                            if level == 1 and h1_style:
                                paragraph.style = h1_style
                            elif level == 2 and h2_style:
                                paragraph.style = h2_style
                            else:
                                # Для третьего и более уровней используем Heading 3, если он есть
                                if 'Heading 3' in document.styles:
                                    paragraph.style = document.styles['Heading 3']
                        except Exception as e:
                            print(f"ОШИБКА при установке стиля заголовка: {str(e)}")

                        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: НЕ используем paragraph.text = ...
                        # Вместо этого изменяем runs
                        new_text = rest
                        if new_text.endswith('.'):
                            new_text = new_text.rstrip('.')
                        
                        if new_text and len(paragraph.runs) > 0:
                            # Находим первый run с нумерацией и заменяем его текст
                            for run in paragraph.runs:
                                if numbering in run.text:
                                    run.text = run.text.replace(numbering, '', 1).lstrip('. :\-–—')
                                    break
                            
                            # Удаляем завершающую точку из последнего run
                            if paragraph.runs:
                                last_run = paragraph.runs[-1]
                                if last_run.text.endswith('.'):
                                    last_run.text = last_run.text.rstrip('.')
                        continue

                    # Нумерация "1. ..." без вложений -> Heading 1
                    if re.match(r'^\d+\.(\s+|$)', text):
                        if h1_style:
                            paragraph.style = h1_style
                        continue

                    # ALL-CAPS короткие заголовки (часто структурные части)
                    if text == text.upper() and 2 <= len(text) <= 80 and not text.endswith('.'):
                        # Неформатированные ключевые слова — точно H1
                        if text in h1_keywords and h1_style:
                            paragraph.style = h1_style
                        continue
                
                except Exception as e:
                    print(f"ОШИБКА при обработке заголовка '{text[:50]}...': {str(e)}")
                    continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _promote_pseudo_headings_to_styles: {str(e)}")
            import traceback
            traceback.print_exc()

    def _correct_specific_errors(self, document, errors):
        """
        Исправляет только указанные ошибки
        """
        error_types = set(error.get('type') for error in errors if 'type' in error)
        
        # Группируем ошибки по типу
        font_errors = any(et.startswith('font_') for et in error_types)
        margin_errors = any(et.endswith('_margin') for et in error_types)
        spacing_errors = 'line_spacing' in error_types
        indent_errors = 'first_line_indent' in error_types
        heading_errors = any(et.startswith('heading_') for et in error_types)
        image_errors = any(et.startswith('image_') for et in error_types)
        table_errors = any(et.startswith('table_') for et in error_types)
        page_numbers_errors = any(et.startswith('page_numbers_') for et in error_types)
        paragraph_alignment_errors = 'paragraph_alignment' in error_types
        list_errors = any(et.startswith('list_') for et in error_types)
        title_page_errors = any(et.startswith('title_page_') for et in error_types)
        
        # Исправляем соответствующие группы ошибок
        if font_errors:
            self._correct_font(document)
        if margin_errors:
            self._correct_margins(document)
        if spacing_errors:
            self._correct_line_spacing(document)
        if indent_errors:
            self._correct_first_line_indent(document)
        if heading_errors:
            self._correct_section_headings(document)
        if image_errors:
            self._correct_images(document)
        if paragraph_alignment_errors:
            self._correct_paragraph_alignment(document)
        if table_errors:
            self._correct_tables(document)
        if page_numbers_errors:
            self._correct_page_numbers(document)
        if list_errors:
            self._correct_lists(document)
        # ОТКЛЮЧЕНО: Исправление титульного листа удаляет весь контент
        # if title_page_errors:
        #     self._correct_title_page(document)
    
    def _correct_font(self, document):
        """
        Исправляет шрифт для всего документа
        Добавлена защита и обработка ошибок
        """
        try:
            # Получаем список всех параграфов внутри таблиц для особой обработки
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            for paragraph in document.paragraphs:
                # Пропускаем пустые параграфы
                if not paragraph.text.strip():
                    continue
                
                try:
                    # Определяем, является ли параграф заголовком
                    is_heading = paragraph.style.name.startswith('Heading')
                    heading_level = None
                    
                    if is_heading:
                        try:
                            heading_level = int(paragraph.style.name.replace('Heading ', ''))
                        except ValueError:
                            heading_level = None
                    
                    # Применяем соответствующий стиль шрифта
                    for run in paragraph.runs:
                        try:
                            # Устанавливаем базовый шрифт для всех элементов
                            if run.font.name != self.rules['font']['name']:
                                run.font.name = self.rules['font']['name']
                            
                            if is_heading and heading_level == 1:
                                # Для заголовков 1 уровня
                                if run.font.size != Pt(self.rules['headings']['h1']['font_size']):
                                    run.font.size = Pt(self.rules['headings']['h1']['font_size'])
                                if run.font.bold != self.rules['headings']['h1']['bold']:
                                    run.font.bold = self.rules['headings']['h1']['bold']
                            elif is_heading and heading_level == 2:
                                # Для заголовков 2 уровня
                                if run.font.size != Pt(self.rules['headings']['h2']['font_size']):
                                    run.font.size = Pt(self.rules['headings']['h2']['font_size'])
                                if run.font.bold != self.rules['headings']['h2']['bold']:
                                    run.font.bold = self.rules['headings']['h2']['bold']
                            else:
                                # Для обычного текста
                                if run.font.size != Pt(self.rules['font']['size']):
                                    run.font.size = Pt(self.rules['font']['size'])
                        
                        except Exception as e:
                            print(f"ОШИБКА при установке шрифта для run: {str(e)}")
                            continue
                
                except Exception as e:
                    print(f"ОШИБКА при обработке параграфа '{paragraph.text[:50]}...': {str(e)}")
                    continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_font: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _correct_margins(self, document):
        """
        Исправляет поля страницы
        """
        # Устанавливаем правильные поля для всех секций документа
        for section in document.sections:
            # Устанавливаем значения полей в сантиметрах
            section.top_margin = Cm(self.rules['margins']['top'])
            section.bottom_margin = Cm(self.rules['margins']['bottom'])
            section.left_margin = Cm(self.rules['margins']['left'])
            section.right_margin = Cm(self.rules['margins']['right'])
            
            # Устанавливаем стандартную ориентацию страницы
            section.orientation = 0  # 0 - портретная ориентация
            
            # Устанавливаем размер страницы A4
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
    
    def _correct_line_spacing(self, document):
        """
        Исправляет межстрочный интервал
        Добавлена защита от таблиц и обработка ошибок
        """
        try:
            # Получаем список всех параграфов внутри таблиц
            table_paragraphs = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            table_paragraphs.add(id(para))
            
            for paragraph in document.paragraphs:
                # Пропускаем параграфы внутри таблиц - у них свои правила
                if id(paragraph) in table_paragraphs:
                    continue
                
                # Пропускаем пустые параграфы
                if not paragraph.text.strip():
                    continue

                try:
                    pf = paragraph.paragraph_format
                    
                    # Устанавливаем полуторный интервал (1.5) для всех абзацев, включая заголовки
                    if pf.line_spacing != self.rules['line_spacing']:
                        pf.line_spacing = self.rules['line_spacing']
                        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

                    # Для обычного текста сбрасываем интервалы до/после; для заголовков их задают стили
                    if not paragraph.style.name.startswith('Heading'):
                        if pf.space_before != Pt(0):
                            pf.space_before = Pt(0)
                        if pf.space_after != Pt(0):
                            pf.space_after = Pt(0)
                
                except Exception as e:
                    print(f"ОШИБКА при установке интервала для параграфа '{paragraph.text[:50]}...': {str(e)}")
                    continue
        
        except Exception as e:
            print(f"КРИТИЧЕСКАЯ ОШИБКА в _correct_line_spacing: {str(e)}")
            import traceback
            traceback.print_exc()
