try:
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    print(f"WD_PARAGRAPH_ALIGNMENT.LEFT: {WD_PARAGRAPH_ALIGNMENT.LEFT}")
except ImportError:
    print("WD_PARAGRAPH_ALIGNMENT not found")

try:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    print(f"WD_ALIGN_PARAGRAPH.LEFT: {WD_ALIGN_PARAGRAPH.LEFT}")
except ImportError:
    print("WD_ALIGN_PARAGRAPH not found")

from docx import Document
doc = Document()
p = doc.add_paragraph("test")
try:
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    print("Assigned WD_PARAGRAPH_ALIGNMENT.LEFT")
except Exception as e:
    print(f"Failed to assign WD_PARAGRAPH_ALIGNMENT.LEFT: {e}")

try:
    from docx.enum.text import WD_LINE_SPACING
    print(f"WD_LINE_SPACING.MULTIPLE: {WD_LINE_SPACING.MULTIPLE}")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    print("Assigned WD_LINE_SPACING.MULTIPLE")
except Exception as e:
    print(f"Failed WD_LINE_SPACING: {e}")

try:
    from docx.enum.table import WD_TABLE_ALIGNMENT
    print(f"WD_TABLE_ALIGNMENT.LEFT: {WD_TABLE_ALIGNMENT.LEFT}")
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    print("Assigned WD_TABLE_ALIGNMENT.LEFT")
except Exception as e:
    print(f"Failed WD_TABLE_ALIGNMENT: {e}")

try:
    from docx.enum.section import WD_SECTION
    print(f"WD_SECTION.NEW_PAGE: {WD_SECTION.NEW_PAGE}")
    # section.start_type = WD_SECTION.NEW_PAGE
    print("Checked WD_SECTION")
except Exception as e:
    print(f"Failed WD_SECTION: {e}")
