#!/usr/bin/env python3
"""
Быстрый тестовый скрипт для проверки системы валидации документов.
Создает тестовый документ и запускает полную валидацию.
"""

import sys
import os
from pathlib import Path
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Добавляем путь к приложению
sys.path.insert(0, str(Path(__file__).parent))

from app.services.validation_engine import ValidationEngine


def create_test_document(filepath: str) -> str:
    """
    Создает тестовый документ для проверки валидаторов.

    Args:
        filepath: Путь для сохранения документа

    Returns:
        Путь к созданному документу
    """
    print("📝 Создание тестового документа...")

    doc = Document()

    # Настройка полей
    for section in doc.sections:
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Добавляем структуру документа

    # 1. Содержание
    para = doc.add_paragraph("СОДЕРЖАНИЕ")
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.name = "Times New Roman"
    para.runs[0].bold = True
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("Введение")
    doc.add_paragraph("Глава 1. Основная часть")
    doc.add_paragraph("Глава 2. Дополнительная информация")
    doc.add_paragraph("Заключение")
    doc.add_paragraph("Список литературы")

    doc.add_paragraph("")  # Пустая строка

    # 2. Введение
    para = doc.add_paragraph("ВВЕДЕНИЕ")
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.name = "Times New Roman"
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Текст введения с правильным форматированием
    intro_text = (
        "Введение определяет актуальность и значимость исследования. "
        "В этом разделе обосновывается выбор темы, изложены основные задачи и цели работы. "
        "Введение должно содержать четкую формулировку проблемы, которая решается в работе."
    )

    para = doc.add_paragraph(intro_text)
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.name = "Times New Roman"
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.line_spacing = 1.5
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("")

    # 3. Глава 1
    para = doc.add_paragraph("ГЛАВА 1 ОСНОВНАЯ ЧАСТЬ")
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.name = "Times New Roman"
    para.runs[0].bold = True
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 1.1
    para = doc.add_paragraph("1.1 Первый раздел")
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.name = "Times New Roman"
    para.runs[0].bold = True

    content = (
        "Содержание первого раздела представляет собой теоретическую базу исследования. "
        "Рассматриваются основные подходы и методологии, которые будут использованы в дальнейшем. "
        "Следует отметить, что данный раздел должен содержать ссылки на список литературы."
    )

    para = doc.add_paragraph(content)
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.name = "Times New Roman"
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.line_spacing = 1.5
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("")

    # 1.2
    para = doc.add_paragraph("1.2 Второй раздел")
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.name = "Times New Roman"
    para.runs[0].bold = True

    content = (
        "Второй раздел продолжает развитие темы. "
        "Здесь могут быть представлены эмпирические данные и результаты анализа."
    )

    para = doc.add_paragraph(content)
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.name = "Times New Roman"
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.line_spacing = 1.5
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Добавляем таблицу
    doc.add_paragraph("Таблица 1 - Результаты исследования")

    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"

    # Заголовок таблицы
    cell = table.rows[0].cells[0]
    cell.text = "Параметр"
    cell = table.rows[0].cells[1]
    cell.text = "Значение"
    cell = table.rows[0].cells[2]
    cell.text = "Единица"

    # Строки данных
    cell = table.rows[1].cells[0]
    cell.text = "Показатель 1"
    cell = table.rows[1].cells[1]
    cell.text = "10"
    cell = table.rows[1].cells[2]
    cell.text = "шт"

    cell = table.rows[2].cells[0]
    cell.text = "Показатель 2"
    cell = table.rows[2].cells[1]
    cell.text = "20"
    cell = table.rows[2].cells[2]
    cell.text = "%"

    doc.add_paragraph("")

    # 4. Заключение
    para = doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.name = "Times New Roman"
    para.runs[0].bold = True
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    conclusion = (
        "В заключении необходимо кратко изложить основные выводы исследования. "
        "Следует подчеркнуть достигнутые результаты и их практическую значимость."
    )

    para = doc.add_paragraph(conclusion)
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.name = "Times New Roman"
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.line_spacing = 1.5
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph("")

    # 5. Список литературы
    para = doc.add_paragraph("СПИСОК ЛИТЕРАТУРЫ")
    para.runs[0].font.size = Pt(14)
    para.runs[0].font.name = "Times New Roman"
    para.runs[0].bold = True
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Примеры источников
    sources = [
        "1. Иванов И.И. Основы методологии научных исследований. М.: Наука, 2020. 250 с.",
        "2. Петров П.П. Современные подходы в аналитике. СПб.: Питер, 2021. 180 с.",
        "3. Сидоров С.С., Сорокин С.С. Теория и практика. М.: МГУ, 2019. 320 с.",
        "4. Кузнецов К.К. Электронные материалы и методы. М.: Техника, 2022. 150 с.",
        "5. Орлов О.О. Интегральные системы. Л.: ЛГУ, 2018. 210 с.",
        "6. Волков В.В. Практическое применение методов. М.: МГТУ, 2023. 280 с.",
        "7. Морозов М.М. Анализ данных в реальном времени. СПб.: СПбГУ, 2021. 200 с.",
        "8. Федоров Ф.Ф. Инновационные подходы. М.: РАН, 2022. 160 с.",
        "9. Лебедев Л.Л. Компьютерные методы обработки. М.: ВГТУ, 2020. 240 с.",
        "10. Соколов С.С., Глушков Г.Г. Современная редакция стандартов. М.: Издательство, 2023. 300 с.",
    ]

    for source in sources:
        para = doc.add_paragraph(source)
        para.runs[0].font.size = Pt(14)
        para.runs[0].font.name = "Times New Roman"
        para.paragraph_format.left_indent = Cm(1.25)
        para.paragraph_format.first_line_indent = Cm(-1.25)  # Висячий отступ
        para.paragraph_format.line_spacing = 1.5

    # Сохраняем документ
    doc.save(filepath)
    print(f"✅ Документ создан: {filepath}")

    return filepath


def run_validation_test(doc_path: str) -> None:
    """
    Запускает полную валидацию документа и выводит результаты.

    Args:
        doc_path: Путь к документу для проверки
    """
    print("\n🚀 Запуск полной валидации документа...\n")

    # Загружаем профиль
    profile = {
        "name": "Test GOST Profile",
        "version": "7.32-2017",
        "rules": {
            "font": {
                "name": "Times New Roman",
                "size": 14.0,
                "allowed_fonts": ["Times New Roman", "Arial"],
            },
            "margins": {"left": 3.0, "right": 1.5, "top": 2.0, "bottom": 2.0},
            "paragraph": {
                "first_line_indent_cm": 1.25,
                "alignment": "justify",
                "line_spacing": 1.5,
                "space_before_pt": 0,
                "space_after_pt": 0,
            },
            "headings": {
                "h1": {"all_caps": True, "alignment": "center"},
                "h2": {"all_caps": True, "alignment": "center"},
            },
            "structure": {
                "required_sections": ["содержание", "введение", "заключение", "список литературы"],
                "min_pages": 5,
                "max_pages": 100,
            },
            "tables": {"require_caption": True, "min_font_size_pt": 10, "max_font_size_pt": 14},
            "bibliography": {"min_sources": 5},
        },
    }

    # Создаем движок валидации
    engine = ValidationEngine(profile=profile)

    # Запускаем валидацию
    report = engine.validate_document(doc_path, document_data={})

    # Выводим результаты
    print("=" * 80)
    print(f"📊 ОТЧЕТ О ВАЛИДАЦИИ")
    print("=" * 80)

    print(f"\n📄 Документ: {report['document']['filename']}")
    print(f"📋 Профиль: {report['profile']['name']}")
    print(f"⏱️  Время проверки: {report['execution']['total_time']} сек")
    print(f"🔍 Валидаторов запущено: {report['execution']['validators_run']}")

    summary = report["summary"]
    print(f"\n📈 СТАТИСТИКА:")
    print(f"  • Всего проблем: {summary['total_issues']}")
    print(f"  • 🔴 Критические: {summary['critical']}")
    print(f"  • 🟠 Ошибки: {summary['errors']}")
    print(f"  • 🟡 Предупреждения: {summary['warnings']}")
    print(f"  • ℹ️  Информация: {summary['info']}")
    print(f"  • ✨ Автоисправимых: {summary['autocorrectable']}")
    print(f"  • 📊 Соответствие: {summary['completion_percentage']:.1f}%")

    print(f"\n📋 СТАТУС: {report['status'].upper()}")
    print(
        f"{'✅ ДОКУМЕНТ СООТВЕТСТВУЕТ' if report['passed'] else '❌ ДОКУМЕНТ НЕ СООТВЕТСТВУЕТ'} ТРЕБОВАНИЯМ"
    )

    # Выводим проблемы по валидаторам
    if report["validators"]:
        print(f"\n🔍 РЕЗУЛЬТАТЫ ПО ВАЛИДАТОРАМ:")
        print("=" * 80)

        for validator in report["validators"]:
            name = validator["validator_name"]
            passed = "✅" if validator["passed"] else "❌"
            stats = validator["statistics"]
            time_sec = validator["execution_time"]

            print(f"\n{passed} {name} ({time_sec:.3f}с)")
            print(f"   Проблем: {stats['total_issues']}", end="")
            if stats["critical"] > 0:
                print(f" | 🔴 {stats['critical']}", end="")
            if stats["errors"] > 0:
                print(f" | 🟠 {stats['errors']}", end="")
            if stats["warnings"] > 0:
                print(f" | 🟡 {stats['warnings']}", end="")
            print()

    # Выводим рекомендации
    if report["recommendations"]:
        print(f"\n💡 РЕКОМЕНДАЦИИ:")
        print("=" * 80)
        for rec in report["recommendations"]:
            print(f"  • {rec}")

    # Выводим критические и серьезные ошибки
    critical_and_errors = (
        report["issues_by_severity"]["critical"] + report["issues_by_severity"]["error"]
    )

    if critical_and_errors:
        print(f"\n🔴 КРИТИЧЕСКИЕ ОШИБКИ И ОШИБКИ ({len(critical_and_errors)}):")
        print("=" * 80)

        for issue in critical_and_errors[:10]:  # Показываем первые 10
            severity_icon = "🔴" if issue["severity"] == "critical" else "🟠"
            print(f"\n{severity_icon} {issue['rule_name']} (Правило {issue['rule_id']})")
            print(f"   {issue['description']}")
            if issue["expected"]:
                print(f"   Ожидается: {issue['expected']}")
            if issue["actual"]:
                print(f"   Найдено: {issue['actual']}")
            if issue["suggestion"]:
                print(f"   Решение: {issue['suggestion']}")

        if len(critical_and_errors) > 10:
            print(f"\n   ... и ещё {len(critical_and_errors) - 10} ошибок")

    print("\n" + "=" * 80)
    print()


def main():
    """Главная функция"""
    print("\n" + "=" * 80)
    print("🧪 БЫСТРЫЙ ТЕСТ СИСТЕМЫ ВАЛИДАЦИИ ДОКУМЕНТОВ")
    print("=" * 80 + "\n")

    # Создаем временный файл
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = os.path.join(tmpdir, "test_document.docx")

        try:
            # Создаем тестовый документ
            create_test_document(doc_path)

            # Запускаем валидацию
            run_validation_test(doc_path)

            print("✨ ТЕСТ ЗАВЕРШЕН УСПЕШНО!")

        except Exception as e:
            print(f"❌ ОШИБКА: {str(e)}")
            import traceback

            traceback.print_exc()
            sys.exit(1)


if __name__ == "__main__":
    main()
