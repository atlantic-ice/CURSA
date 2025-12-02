"""
Тест многопроходной коррекции документов.
Создаёт тестовый документ с ошибками и проверяет качество исправления.
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Добавляем путь для импорта модулей
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app.services.document_corrector import DocumentCorrector, CorrectionReport


def create_test_document_with_errors(path: str):
    """Создаёт тестовый документ с типичными ошибками форматирования"""
    doc = Document()
    
    # Заголовок 1 уровня с ошибками
    h1 = doc.add_paragraph("ВВЕДЕНИЕ")
    h1.style = 'Heading 1'
    for run in h1.runs:
        run.font.name = 'Arial'  # ОШИБКА: должен быть Times New Roman
        run.font.size = Pt(16)   # ОШИБКА: должен быть 14pt
    
    # Обычный текст с ошибками
    p1 = doc.add_paragraph(
        "Это тестовый документ для проверки многопроходной коррекции. "
        "В нём специально допущены различные ошибки форматирования, "
        "которые должны быть исправлены системой автоматической коррекции."
    )
    for run in p1.runs:
        run.font.name = 'Calibri'  # ОШИБКА: должен быть Times New Roman
        run.font.size = Pt(12)      # ОШИБКА: должен быть 14pt
    p1.paragraph_format.line_spacing = 1.0  # ОШИБКА: должен быть 1.5
    p1.paragraph_format.first_line_indent = Cm(0)  # ОШИБКА: должен быть 1.25 см
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # ОШИБКА: должен быть JUSTIFY
    
    # Ещё один абзац с ошибками
    p2 = doc.add_paragraph(
        "Второй абзац документа также содержит ошибки. "
        "Шрифт неправильный, интервал одинарный вместо полуторного, "
        "выравнивание по левому краю вместо выравнивания по ширине."
    )
    for run in p2.runs:
        run.font.name = 'Georgia'  # ОШИБКА
        run.font.size = Pt(11)      # ОШИБКА
    p2.paragraph_format.line_spacing = 1.0  # ОШИБКА
    p2.paragraph_format.first_line_indent = Cm(0.5)  # ОШИБКА: не 1.25 см
    
    # Заголовок 2 уровня
    h2 = doc.add_paragraph("1.1 Актуальность темы")
    h2.style = 'Heading 2'
    for run in h2.runs:
        run.font.name = 'Verdana'  # ОШИБКА
        run.font.bold = False       # ОШИБКА: должен быть жирный
    
    # Текст после заголовка
    p3 = doc.add_paragraph(
        "Текст после заголовка второго уровня. "
        "Также содержит ошибки форматирования для тестирования."
    )
    for run in p3.runs:
        run.font.name = 'Tahoma'
        run.font.size = Pt(13)
    p3.paragraph_format.line_spacing = 2.0  # ОШИБКА: слишком большой интервал
    
    # Сохраняем
    doc.save(path)
    print(f"Тестовый документ создан: {path}")
    return path


def count_errors(doc: Document) -> dict:
    """Подсчитывает количество ошибок в документе"""
    errors = {
        'font_name': 0,
        'font_size': 0,
        'line_spacing': 0,
        'alignment': 0,
        'indent': 0
    }
    
    for para in doc.paragraphs:
        # Проверка шрифтов
        for run in para.runs:
            if run.font.name and run.font.name != 'Times New Roman':
                errors['font_name'] += 1
            if run.font.size and run.font.size != Pt(14) and not para.style.name.startswith('Heading'):
                errors['font_size'] += 1
        
        # Проверка интервалов
        if para.paragraph_format.line_spacing and para.paragraph_format.line_spacing != 1.5:
            errors['line_spacing'] += 1
        
        # Проверка выравнивания (для обычного текста)
        if not para.style.name.startswith('Heading'):
            if para.alignment and para.alignment != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                errors['alignment'] += 1
    
    return errors


def test_multipass_correction():
    """Основной тест многопроходной коррекции"""
    print("=" * 60)
    print("ТЕСТ МНОГОПРОХОДНОЙ КОРРЕКЦИИ ДОКУМЕНТОВ")
    print("=" * 60)
    
    # Создаём тестовый документ
    test_dir = os.path.join(os.path.dirname(__file__), 'test_data')
    os.makedirs(test_dir, exist_ok=True)
    
    test_file = os.path.join(test_dir, 'test_multipass_input.docx')
    output_file = os.path.join(test_dir, 'test_multipass_output.docx')
    
    create_test_document_with_errors(test_file)
    
    # Подсчитываем ошибки до коррекции
    doc_before = Document(test_file)
    errors_before = count_errors(doc_before)
    total_before = sum(errors_before.values())
    
    print(f"\nОшибки ДО коррекции:")
    for error_type, count in errors_before.items():
        print(f"  {error_type}: {count}")
    print(f"  ВСЕГО: {total_before}")
    
    # Выполняем многопроходную коррекцию
    print("\n" + "-" * 60)
    print("Запуск многопроходной коррекции...")
    print("-" * 60)
    
    corrector = DocumentCorrector()
    corrector.verbose_logging = True
    corrector.max_passes = 3
    
    corrected_path, report = corrector.correct_document_multipass(
        test_file,
        out_path=output_file,
        max_passes=3
    )
    
    # Подсчитываем ошибки после коррекции
    doc_after = Document(corrected_path)
    errors_after = count_errors(doc_after)
    total_after = sum(errors_after.values())
    
    print(f"\nОшибки ПОСЛЕ коррекции:")
    for error_type, count in errors_after.items():
        print(f"  {error_type}: {count}")
    print(f"  ВСЕГО: {total_after}")
    
    # Выводим отчёт
    print("\n" + "=" * 60)
    print("ОТЧЁТ О КОРРЕКЦИИ")
    print("=" * 60)
    
    summary = report.get_summary()
    print(f"Проходов выполнено: {summary['passes_completed']}")
    print(f"Найдено проблем: {summary['total_issues_found']}")
    print(f"Исправлено: {summary['total_issues_fixed']}")
    print(f"Осталось: {summary['remaining_issues']}")
    print(f"Успешность: {summary['success_rate']}%")
    print(f"Время выполнения: {summary['duration_seconds']:.2f} сек")
    
    print("\nДействия по фазам:")
    for phase, count in summary['actions_by_phase'].items():
        print(f"  {phase}: {count}")
    
    print("\nРезультаты верификации:")
    for check, result in report.verification_results.items():
        status = "✓" if result.get('passed', False) else "✗"
        print(f"  {status} {check}: {result.get('message', '')}")
    
    # Итоговый результат
    print("\n" + "=" * 60)
    improvement = total_before - total_after
    improvement_percent = (improvement / max(total_before, 1)) * 100
    
    print(f"ИТОГО:")
    print(f"  Ошибок до коррекции: {total_before}")
    print(f"  Ошибок после коррекции: {total_after}")
    print(f"  Улучшение: {improvement} ({improvement_percent:.1f}%)")
    
    if total_after == 0:
        print("\n✓ ОТЛИЧНО! Все ошибки исправлены!")
    elif improvement_percent >= 80:
        print("\n✓ ХОРОШО! Большинство ошибок исправлено.")
    else:
        print("\n⚠ Требуется доработка алгоритмов коррекции.")
    
    print("=" * 60)
    
    return total_after == 0


if __name__ == '__main__':
    success = test_multipass_correction()
    sys.exit(0 if success else 1)
