"""
Тест гибридного подхода: python-docx + XML.
Сравнивает качество исправлений стандартного и гибридного методов.
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Добавляем путь для импорта модулей
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app.services.document_corrector import DocumentCorrector, CorrectionReport
from app.services.xml_document_editor import XMLDocumentEditor, apply_xml_corrections


def create_complex_test_document(path: str):
    """Создаёт документ со сложными ошибками форматирования"""
    doc = Document()
    
    # Заголовок с множественными ошибками
    h1 = doc.add_paragraph("ВВЕДЕНИЕ")
    h1.style = 'Heading 1'
    for run in h1.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = False
    
    # Текст с вложенными run-элементами и разными шрифтами
    p1 = doc.add_paragraph()
    run1 = p1.add_run("Начало абзаца с одним шрифтом. ")
    run1.font.name = 'Calibri'
    run1.font.size = Pt(11)
    
    run2 = p1.add_run("Продолжение с другим шрифтом. ")
    run2.font.name = 'Arial'
    run2.font.size = Pt(12)
    
    run3 = p1.add_run("И ещё один шрифт в конце.")
    run3.font.name = 'Verdana'
    run3.font.size = Pt(10)
    
    p1.paragraph_format.line_spacing = 1.0
    p1.paragraph_format.first_line_indent = Cm(0)
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Абзац с экстремальными отступами
    p2 = doc.add_paragraph(
        "Этот абзац имеет неправильные отступы и интервалы. "
        "Левый отступ слишком большой, интервал двойной."
    )
    for run in p2.runs:
        run.font.name = 'Georgia'
        run.font.size = Pt(13)
    p2.paragraph_format.line_spacing = 2.0
    p2.paragraph_format.left_indent = Cm(2)
    p2.paragraph_format.first_line_indent = Cm(0.5)
    
    # Заголовок 2 уровня
    h2 = doc.add_paragraph("1.1 Актуальность исследования")
    h2.style = 'Heading 2'
    for run in h2.runs:
        run.font.name = 'Tahoma'
        run.font.size = Pt(12)
        run.font.bold = False
    
    # Ещё несколько абзацев с ошибками
    for i in range(3):
        p = doc.add_paragraph(
            f"Абзац {i+3} с ошибками форматирования. "
            "Шрифт неправильный, интервал одинарный."
        )
        fonts = ['Comic Sans MS', 'Impact', 'Courier New']
        sizes = [9, 15, 11]
        for run in p.runs:
            run.font.name = fonts[i % len(fonts)]
            run.font.size = Pt(sizes[i % len(sizes)])
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Cm(0)
    
    doc.save(path)
    print(f"Тестовый документ создан: {path}")
    return path


def count_detailed_errors(doc: Document) -> dict:
    """Подсчитывает ошибки с детализацией"""
    errors = {
        'font_name': 0,
        'font_size': 0,
        'line_spacing': 0,
        'alignment': 0,
        'indent': 0,
        'runs_with_wrong_font': 0,
        'total_runs': 0
    }
    
    # Допуск для сравнения отступов (разница в округлении)
    INDENT_TOLERANCE = Cm(0.05)  # 0.5 мм допуска
    expected_indent = Cm(1.25)
    
    for para in doc.paragraphs:
        # Проверка run-элементов
        for run in para.runs:
            errors['total_runs'] += 1
            if run.font.name and run.font.name != 'Times New Roman':
                errors['font_name'] += 1
                errors['runs_with_wrong_font'] += 1
            if run.font.size and run.font.size != Pt(14) and not para.style.name.startswith('Heading'):
                errors['font_size'] += 1
        
        # Проверка интервалов
        if para.paragraph_format.line_spacing and para.paragraph_format.line_spacing != 1.5:
            errors['line_spacing'] += 1
        
        # Проверка выравнивания
        if not para.style.name.startswith('Heading'):
            if para.alignment and para.alignment != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                errors['alignment'] += 1
        
        # Проверка отступов с допуском
        if not para.style.name.startswith('Heading'):
            indent = para.paragraph_format.first_line_indent
            if indent:
                # Проверяем с допуском на погрешность округления
                diff = abs(indent - expected_indent)
                if diff > INDENT_TOLERANCE:
                    errors['indent'] += 1
    
    return errors


def test_xml_editor_standalone():
    """Тест XML-редактора отдельно"""
    print("\n" + "=" * 60)
    print("ТЕСТ XML-РЕДАКТОРА (STANDALONE)")
    print("=" * 60)
    
    test_dir = os.path.join(os.path.dirname(__file__), 'test_data')
    os.makedirs(test_dir, exist_ok=True)
    
    input_file = os.path.join(test_dir, 'xml_test_input.docx')
    output_file = os.path.join(test_dir, 'xml_test_output.docx')
    
    # Создаём тестовый документ
    create_complex_test_document(input_file)
    
    # Считаем ошибки до
    doc_before = Document(input_file)
    errors_before = count_detailed_errors(doc_before)
    total_before = sum(v for k, v in errors_before.items() if k != 'total_runs')
    
    print(f"\nОшибки ДО XML-коррекции:")
    for key, value in errors_before.items():
        print(f"  {key}: {value}")
    
    # Применяем XML-коррекцию
    print("\nПрименяем XML-коррекцию...")
    saved_path, report = apply_xml_corrections(input_file, output_file)
    
    print(f"\nОтчёт XML-редактора:")
    print(f"  Всего изменений: {report.total_edits}")
    print(f"  Успешных: {report.successful_edits}")
    print(f"  Ошибок: {report.failed_edits}")
    
    # Считаем ошибки после
    doc_after = Document(output_file)
    errors_after = count_detailed_errors(doc_after)
    total_after = sum(v for k, v in errors_after.items() if k != 'total_runs')
    
    print(f"\nОшибки ПОСЛЕ XML-коррекции:")
    for key, value in errors_after.items():
        print(f"  {key}: {value}")
    
    improvement = total_before - total_after
    improvement_pct = (improvement / max(total_before, 1)) * 100
    
    print(f"\nУлучшение: {improvement} ({improvement_pct:.1f}%)")
    
    return improvement_pct


def test_hybrid_approach():
    """Тест гибридного подхода python-docx + XML"""
    print("\n" + "=" * 60)
    print("ТЕСТ ГИБРИДНОГО ПОДХОДА (python-docx + XML)")
    print("=" * 60)
    
    test_dir = os.path.join(os.path.dirname(__file__), 'test_data')
    os.makedirs(test_dir, exist_ok=True)
    
    input_file = os.path.join(test_dir, 'hybrid_test_input.docx')
    output_standard = os.path.join(test_dir, 'hybrid_test_standard.docx')
    output_hybrid = os.path.join(test_dir, 'hybrid_test_hybrid.docx')
    
    # Создаём тестовый документ
    create_complex_test_document(input_file)
    
    # Считаем ошибки до
    doc_before = Document(input_file)
    errors_before = count_detailed_errors(doc_before)
    total_before = sum(v for k, v in errors_before.items() if k != 'total_runs')
    
    print(f"\nОшибки в исходном документе: {total_before}")
    
    # === ТЕСТ 1: Стандартный метод (только python-docx) ===
    print("\n--- Стандартный метод (python-docx) ---")
    corrector_standard = DocumentCorrector()
    corrector_standard.verbose_logging = False
    corrector_standard.enable_xml_correction = False  # Отключаем XML
    
    path_standard, report_standard = corrector_standard.correct_document_multipass(
        input_file, out_path=output_standard
    )
    
    doc_standard = Document(path_standard)
    errors_standard = count_detailed_errors(doc_standard)
    total_standard = sum(v for k, v in errors_standard.items() if k != 'total_runs')
    
    print(f"Ошибок после стандартной коррекции: {total_standard}")
    print(f"Успешность: {report_standard.get_summary()['success_rate']}%")
    
    # === ТЕСТ 2: Гибридный метод (python-docx + XML) ===
    print("\n--- Гибридный метод (python-docx + XML) ---")
    
    # Пересоздаём документ для честного сравнения
    create_complex_test_document(input_file)
    
    corrector_hybrid = DocumentCorrector()
    corrector_hybrid.verbose_logging = True
    corrector_hybrid.enable_xml_correction = True  # Включаем XML
    
    path_hybrid, report_hybrid = corrector_hybrid.correct_document_multipass(
        input_file, out_path=output_hybrid
    )
    
    doc_hybrid = Document(path_hybrid)
    errors_hybrid = count_detailed_errors(doc_hybrid)
    total_hybrid = sum(v for k, v in errors_hybrid.items() if k != 'total_runs')
    
    print(f"\nОшибок после гибридной коррекции: {total_hybrid}")
    print(f"Успешность: {report_hybrid.get_summary()['success_rate']}%")
    
    # === СРАВНЕНИЕ ===
    print("\n" + "=" * 60)
    print("СРАВНЕНИЕ РЕЗУЛЬТАТОВ")
    print("=" * 60)
    
    print(f"\nИсходный документ:     {total_before} ошибок")
    print(f"Стандартный метод:     {total_standard} ошибок ({100 - total_standard/max(total_before,1)*100:.1f}% исправлено)")
    print(f"Гибридный метод:       {total_hybrid} ошибок ({100 - total_hybrid/max(total_before,1)*100:.1f}% исправлено)")
    
    if total_hybrid < total_standard:
        improvement = total_standard - total_hybrid
        print(f"\n✓ Гибридный метод лучше на {improvement} ошибок!")
    elif total_hybrid == total_standard:
        print("\n= Методы показали одинаковый результат")
    else:
        print("\n⚠ Стандартный метод показал лучший результат")
    
    # Детальное сравнение
    print("\nДетальное сравнение:")
    print(f"{'Тип ошибки':<25} {'До':>8} {'Станд.':>8} {'Гибрид':>8}")
    print("-" * 50)
    for key in errors_before:
        if key != 'total_runs':
            print(f"{key:<25} {errors_before[key]:>8} {errors_standard[key]:>8} {errors_hybrid[key]:>8}")
    
    return total_hybrid == 0


def main():
    print("=" * 60)
    print("ТЕСТИРОВАНИЕ ГИБРИДНОГО ПОДХОДА")
    print("python-docx + прямая работа с XML")
    print("=" * 60)
    
    # Тест 1: XML-редактор отдельно
    xml_improvement = test_xml_editor_standalone()
    
    # Тест 2: Гибридный подход
    hybrid_success = test_hybrid_approach()
    
    # Итоги
    print("\n" + "=" * 60)
    print("ИТОГИ ТЕСТИРОВАНИЯ")
    print("=" * 60)
    
    print(f"XML-редактор: улучшение {xml_improvement:.1f}%")
    print(f"Гибридный подход: {'✓ УСПЕХ' if hybrid_success else '⚠ Требует доработки'}")
    
    return hybrid_success


if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
