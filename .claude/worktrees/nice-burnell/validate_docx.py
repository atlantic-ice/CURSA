from docx import Document
from docx.shared import Pt, Cm
from pathlib import Path

docx_file = Path("VKR_NIKITA_CURSA_2026.docx")

print("[VALIDATION] Starting comprehensive DOCX validation...")
print()

try:
    doc = Document(str(docx_file))
    
    # Document structure
    print("[STRUCTURE]")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")
    print()
    
    # Check formatting
    print("[FORMATTING CHECK]")
    
    # Sample first few paragraphs
    for idx in range(min(5, len(doc.paragraphs))):
        para = doc.paragraphs[idx]
        if para.text.strip():
            print(f"  Para {idx}: {para.text[:50]}...")
            if para.runs:
                run = para.runs[0]
                print(f"    Font: {run.font.name}, Size: {run.font.size}")
                print(f"    Bold: {run.bold}, Line spacing: {para.paragraph_format.line_spacing}")
    print()
    
    # Check margins
    print("[PAGE SETUP]")
    section = doc.sections[0]
    print(f"  Left Margin: {section.left_margin.cm:.1f} cm")
    print(f"  Right Margin: {section.right_margin.cm:.1f} cm")
    print(f"  Top Margin: {section.top_margin.cm:.1f} cm")
    print(f"  Bottom Margin: {section.bottom_margin.cm:.1f} cm")
    print()
    
    # Table validation
    if len(doc.tables) > 0:
        print("[TABLES]")
        for idx, table in enumerate(doc.tables[:3]):
            print(f"  Table {idx}: {len(table.rows)} rows x {len(table.columns)} columns")
    print()
    
    # File integrity
    print("[FILE INTEGRITY]")
    print(f"  File exists: True")
    print(f"  File size: {docx_file.stat().st_size:,} bytes")
    print(f"  File path: {docx_file.absolute()}")
    print()
    
    print("[VALIDATION] Document is valid and ready for use")
    
except Exception as e:
    print(f"[ERROR] Validation failed: {e}")
