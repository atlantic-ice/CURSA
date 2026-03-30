#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Konverter Markdown v DOCX dlya VKR
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_markdown_to_docx(md_file: str, output_file: str):
    """Konvertiruyet Markdown file v DOCX"""
    
    doc = Document()
    
    # Nastroyki straniczy
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    
    # Stiliz Normal
    styles = doc.styles
    style = styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Cm(1.25)
    
    # CHitaem file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Zagolovok 1 (# )
        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Zagolovok 2 (## )
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        
        # Zagolovok 3 (### )
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            p.paragraph_format.space_before = Pt(8)
        
        # Zhhirnyy tekst (**text**)
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                elif part:
                    run = p.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
        
        # Tabliza (nachinayetsya s |)
        elif line.strip().startswith('|'):
            # Sobirayem tablizu
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1
            
            # Parsim tablizu
            rows = []
            for tbl_line in table_lines:
                if '---' in tbl_line or ':--:' in tbl_line:
                    continue
                cells = [c.strip() for c in tbl_line.strip('|').split('|')]
                rows.append(cells)
            
            if len(rows) > 1:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        cell = tbl.rows[row_idx].cells[col_idx]
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                                if row_idx == 0:
                                    run.bold = True
        
        # Pustaya stroka
        elif line.strip() == '':
            pass
        
        # Obyichnyy tekst
        else:
            p = doc.add_paragraph()
            line = line.replace('`', '')
            run = p.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        
        i += 1
    
    # Sohranyaem
    doc.save(output_file)
    print(f"[OK] Sohraneno: {output_file}")


if __name__ == "__main__":
    script_dir = Path(__file__).parent
    
    # Konvertiruyem polnuyu VKR
    input_file = script_dir / "thesis" / "VKR_full.md"
    output_file = script_dir / "thesis" / "VKR_output.docx"
    
    if input_file.exists():
        parse_markdown_to_docx(str(input_file), str(output_file))
    else:
        print(f"[ERROR] File not found: {input_file}")
        
        # Konvertiruyem po chastyam
        parts = [
            ("title_page.md", "VKR_title.docx"),
            ("VKR_text.md", "VKR_part1.docx"),
            ("VKR_chapter2.md", "VKR_chapter2.docx"),
            ("VKR_chapter3.md", "VKR_chapter3.docx"),
            ("VKR_conclusion.md", "VKR_conclusion.docx"),
        ]
        
        for part, out in parts:
            input_p = script_dir / "thesis" / part
            output_p = script_dir / "thesis" / out
            if input_p.exists():
                parse_markdown_to_docx(str(input_p), str(output_p))
